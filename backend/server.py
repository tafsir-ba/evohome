from fastapi import FastAPI, APIRouter, HTTPException, Depends, Response, Request, UploadFile, File, Form, WebSocket, WebSocketDisconnect
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import secrets
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional, Literal, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import bcrypt
import jwt
import httpx
from io import BytesIO
import tempfile
import json
import re

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER

# AI extraction imports - using OpenAI directly
import openai

# PDF to image conversion
import fitz  # PyMuPDF

# Stripe imports - using stripe SDK directly
import stripe

# Import centralized auth module - ALL auth now goes through core.auth
# Import startup configuration validation - MUST be first
from core.config import validate_config, get_config

# Validate configuration at import time - fail fast if misconfigured
app_config = validate_config()

from core.auth import (
    init_auth, 
    create_access_token, 
    create_refresh_token,
    verify_token,
    extract_token,
    invalidate_token,
    is_token_invalidated,
    JWT_EXPIRY_HOURS,
    get_current_user,
    get_current_agent,
    get_current_buyer
)

# Import centralized access control module
from core.access_control import (
    init_access_control,
    can_access_project,
    can_access_client,
    can_access_vault_doc,
    can_access_document,
    get_accessible_project_ids,
    get_accessible_client_ids,
    is_agent,
    is_buyer,
    get_is_demo
)

# Import rate limiting
from core.rate_limit import rate_limit_check, check_rate_limit

# Import error monitoring
from core.monitoring import (
    capture_exception,
    capture_auth_failure,
    capture_payment_error,
    capture_email_error,
    capture_ai_error,
    capture_websocket_error,
    capture_document_error,
    ErrorContext
)

# Import API response models for strict contracts
from core.responses import (
    AuthSessionResponse,
    AuthLoginResponse,
    AuthRefreshResponse,
    AuthLogoutResponse,
    DocumentResponse,
    VaultDocumentResponse,
    NotificationResponse,
    ActivityResponse,
    ActivitiesListResponse,
    SuccessResponse
)

ROOT_DIR = Path(__file__).parent
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

def secure_filename(filename: str) -> str:
    """Sanitize filename to prevent path traversal attacks"""
    # Keep only safe characters
    safe_chars = re.sub(r'[^\w\s\-\.]', '', filename)
    # Replace spaces with underscores
    safe_chars = safe_chars.replace(' ', '_')
    # Limit length
    return safe_chars[:100] if safe_chars else 'unnamed_file'

load_dotenv(ROOT_DIR / '.env')

# MongoDB connection - using validated config
mongo_url = app_config.MONGO_URL
client = AsyncIOMotorClient(mongo_url)
db = client[app_config.DB_NAME]

# Initialize auth module with database
init_auth(db)

# Initialize access control module with database
init_access_control(db)

# Legacy JWT config - using validated config (no silent fallback)
JWT_SECRET = app_config.JWT_SECRET
JWT_ALGORITHM = 'HS256'
JWT_EXPIRY_DAYS = 7

app = FastAPI(title="UpgradeFlow API")
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ==================== PYDANTIC MODELS ====================

class UserBase(BaseModel):
    user_id: str
    email: str
    name: str
    role: str
    picture: Optional[str] = None
    is_demo: bool = False
    created_at: datetime

class AgentRegister(BaseModel):
    email: EmailStr
    password: str
    name: str

class AgentLogin(BaseModel):
    email: EmailStr
    password: str

class BuyerRegister(BaseModel):
    email: EmailStr
    password: str
    name: str
    invitation_code: Optional[str] = None  # For invited buyers

class BuyerLogin(BaseModel):
    email: EmailStr
    password: str

class ClientCreate(BaseModel):
    name: str
    email: EmailStr
    phone: Optional[str] = None
    project_id: str
    unit_id: Optional[str] = None  # References project_units.unit_id

class ClientUpdate(BaseModel):
    name: Optional[str] = None
    email: Optional[EmailStr] = None
    phone: Optional[str] = None
    project_id: Optional[str] = None
    unit_id: Optional[str] = None  # References project_units.unit_id
    force_unit_reassign: Optional[bool] = False  # Allow reassigning a unit from another client

class Client(BaseModel):
    client_id: str
    agent_id: str
    buyer_id: Optional[str] = None
    name: str
    email: str
    phone: Optional[str] = None
    project_id: str
    unit_id: Optional[str] = None
    unit_reference: Optional[str] = "General"  # Display name of the unit - optional with default
    status: Optional[str] = "active"
    is_demo: bool = False
    created_at: Optional[datetime] = None

class ProjectCreate(BaseModel):
    name: str
    address: Optional[str] = None
    description: Optional[str] = None
    total_units: Optional[int] = None
    construction_start: Optional[str] = None
    estimated_completion: Optional[str] = None

class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    address: Optional[str] = None
    description: Optional[str] = None
    total_units: Optional[int] = None
    construction_start: Optional[str] = None
    estimated_completion: Optional[str] = None

class Project(BaseModel):
    project_id: str
    agent_id: str
    name: str
    address: Optional[str] = None
    description: Optional[str] = None
    total_units: Optional[int] = None
    construction_start: Optional[str] = None
    estimated_completion: Optional[str] = None
    client_count: Optional[int] = 0
    unit_count: Optional[int] = 0  # Actual units created
    is_demo: bool = False
    created_at: datetime

# ==================== TEAM MEMBER MODELS ====================

class TeamMemberCreate(BaseModel):
    company_name: str
    contact_name: str
    role: str  # e.g., "Plumber", "Electrician", "Architect"
    email: Optional[str] = None
    phone: Optional[str] = None
    website: Optional[str] = None
    address: Optional[str] = None
    notes: Optional[str] = None

class TeamMemberUpdate(BaseModel):
    company_name: Optional[str] = None
    contact_name: Optional[str] = None
    role: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    website: Optional[str] = None
    address: Optional[str] = None
    notes: Optional[str] = None

class TeamMember(BaseModel):
    member_id: str
    project_id: str
    agent_id: str
    company_name: str
    contact_name: str
    role: str
    email: Optional[str] = None
    phone: Optional[str] = None
    website: Optional[str] = None
    address: Optional[str] = None
    notes: Optional[str] = None
    is_demo: bool = False
    created_at: str

class ExtractedContact(BaseModel):
    """Contact extracted from a document via AI"""
    company_name: Optional[str] = None
    contact_name: Optional[str] = None
    role: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    website: Optional[str] = None
    address: Optional[str] = None
    confidence: Optional[float] = None  # 0-1 confidence score

# ==================== UNIFIED DOCUMENT MODEL ====================
# Single model for BOTH quotes and invoices. No separate logic.

class DocumentLineItem(BaseModel):
    description: str
    quantity: int = 1
    unit_price: float
    total: float

class DocumentUpdate(BaseModel):
    """Update document with AI-extracted or manually edited data"""
    title: Optional[str] = None
    amount: Optional[float] = None
    items: Optional[List[DocumentLineItem]] = None
    supplier_name: Optional[str] = None
    notes: Optional[str] = None
    summary: Optional[str] = None
    hero_image_url: Optional[str] = None
    # Editable references
    client_id: Optional[str] = None
    project_id: Optional[str] = None
    unit_id: Optional[str] = None

class DocumentAction(BaseModel):
    """Buyer/Agent action on document"""
    action: str  # approve, reject, request_change, confirm_payment, convert_to_invoice
    comment: Optional[str] = None

class Document(BaseModel):
    """UNIFIED model for both Quotes and Invoices"""
    id: str = Field(..., alias="document_id")
    type: Literal["quote", "invoice"]
    status: str
    title: str
    amount: float
    items: List[DocumentLineItem]
    pdf_url: Optional[str] = None
    created_at: datetime
    updated_at: datetime
    buyer_id: Optional[str] = None
    agent_id: str
    client_id: str
    project_id: str
    unit_reference: str
    parent_document_id: Optional[str] = None  # For quote -> invoice linking
    # Versioning
    version: int = 1
    version_history: Optional[list] = None  # [{version, pdf_path, updated_at, updated_by}]
    # Additional fields
    document_number: str
    currency: str = "CHF"
    supplier_name: Optional[str] = None
    notes: Optional[str] = None
    summary: Optional[str] = None  # Short description for card display
    hero_image_url: Optional[str] = None  # Banner image URL
    hero_image_path: Optional[str] = None  # Local storage path
    change_request_comment: Optional[str] = None
    pdf_filename: Optional[str] = None
    pdf_path: Optional[str] = None
    ai_extraction_confidence: Optional[str] = None
    due_date: Optional[str] = None
    paid_date: Optional[str] = None
    is_demo: bool = False

# ==================== STATE MACHINE ====================
# QUOTE: Draft -> Sent -> (Change Requested -> Sent) -> Approved | Rejected
# INVOICE: Draft -> Sent -> Paid

VALID_TRANSITIONS = {
    "quote": {
        "Draft": ["Sent"],
        "Sent": ["Approved", "Rejected", "Change Requested"],
        "Change Requested": ["Sent"],
        "Approved": [],  # Terminal - can be converted to invoice
        "Rejected": [],  # Terminal
    },
    "invoice": {
        "Draft": ["Sent"],
        "Sent": ["Paid", "Change Requested"],
        "Change Requested": ["Sent"],
        "Paid": [],  # Terminal
    }
}

def validate_transition(doc_type: str, current_status: str, new_status: str) -> bool:
    """Validate if a status transition is allowed"""
    allowed = VALID_TRANSITIONS.get(doc_type, {}).get(current_status, [])
    return new_status in allowed

# ==================== PROJECT STAGE MODELS ====================

class ProjectStageCreate(BaseModel):
    name: str
    description: Optional[str] = None
    order: int
    planned_start: str
    planned_end: str
    dependencies: Optional[List[str]] = None

class ProjectStageUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    order: Optional[int] = None
    planned_start: Optional[str] = None
    planned_end: Optional[str] = None
    actual_start: Optional[str] = None
    actual_end: Optional[str] = None
    status: Optional[str] = None
    progress_percent: Optional[int] = None
    notes: Optional[str] = None

class ProjectStage(BaseModel):
    stage_id: str
    project_id: str
    agent_id: str
    name: str
    description: Optional[str] = None
    order: int
    planned_start: str
    planned_end: str
    actual_start: Optional[str] = None
    actual_end: Optional[str] = None
    status: str
    progress_percent: int = 0
    notes: Optional[str] = None
    dependencies: List[str] = []
    is_demo: bool = False
    created_at: datetime
    updated_at: datetime

# ==================== ACTIVITY FEED MODELS ====================

class ActivityCreate(BaseModel):
    type: Literal["message", "image", "file", "status"]  # file = pdf/docs with change request support
    title: Optional[str] = None
    content: Optional[str] = None  # Nullable when file-only
    project_id: str
    unit_id: Optional[str] = None
    client_ids: List[str]  # Recipients
    file_type: Optional[str] = None  # pdf, image, other (for type=file)

class ActivityReplyCreate(BaseModel):
    content: str

class Activity(BaseModel):
    activity_id: str
    type: Literal["message", "image", "file", "status", "pdf"]  # pdf kept for backwards compat
    title: Optional[str] = None
    content: Optional[str] = None
    file_url: Optional[str] = None
    file_name: Optional[str] = None
    file_size: Optional[int] = None
    file_type: Optional[str] = None  # pdf, image, other
    author_id: str
    author_role: Literal["agent", "buyer"]
    author_name: Optional[str] = None
    project_id: str
    project_name: Optional[str] = None
    unit_id: Optional[str] = None
    unit_reference: Optional[str] = None
    is_demo: bool = False
    created_at: str
    updated_at: Optional[str] = None
    recipients: Optional[List[dict]] = None
    replies: Optional[List[dict]] = None
    reply_count: int = 0
    # For file type - change request support
    status: Optional[str] = None  # None, "change_requested"
    change_comment: Optional[str] = None

class ActivityReply(BaseModel):
    reply_id: str
    activity_id: str
    author_id: str
    author_role: Literal["agent", "buyer"]
    author_name: Optional[str] = None
    content: str
    created_at: str

# ==================== TIMELINE/WORKFLOW MODELS ====================

class TimelineTemplateStepCreate(BaseModel):
    title: str
    description: Optional[str] = None
    order_index: int

class TimelineTemplateCreate(BaseModel):
    name: str
    steps: List[TimelineTemplateStepCreate]

class TimelineTemplate(BaseModel):
    template_id: str
    agent_id: str
    name: str
    is_demo: bool = False
    created_at: str

class TimelineTemplateStep(BaseModel):
    step_id: str
    template_id: str
    title: str
    description: Optional[str] = None
    order_index: int

class TimelineStepUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    status: Optional[Literal["pending", "in_progress", "completed", "approved"]] = None
    planned_date: Optional[str] = None
    order_index: Optional[int] = None
    # Extended fields (from project_stages migration)
    planned_start: Optional[str] = None
    planned_end: Optional[str] = None
    actual_start: Optional[str] = None
    progress_percent: Optional[int] = None
    notes: Optional[str] = None

class TimelineStepDocumentCreate(BaseModel):
    activity_id: str

class TimelineStepNoteCreate(BaseModel):
    content: str

class TimelineStep(BaseModel):
    step_id: str
    project_timeline_id: str
    title: str
    description: Optional[str] = None
    status: Literal["pending", "in_progress", "completed", "approved"]
    order_index: int
    planned_date: Optional[str] = None
    completed_at: Optional[str] = None
    created_at: str
    updated_at: str
    # Extended fields (from project_stages migration)
    planned_start: Optional[str] = None
    planned_end: Optional[str] = None
    actual_start: Optional[str] = None
    progress_percent: Optional[int] = 0
    notes: Optional[str] = None
    dependencies: Optional[List[str]] = []
    # Enriched data
    documents: Optional[List[dict]] = None
    internal_notes: Optional[List[dict]] = None

class ProjectTimeline(BaseModel):
    timeline_id: str
    project_id: str
    template_id: Optional[str] = None
    template_name: Optional[str] = None
    is_demo: bool = False
    created_at: str
    steps: Optional[List[TimelineStep]] = None

# ==================== NOTIFICATION MODELS ====================

class Notification(BaseModel):
    notification_id: str
    user_id: str
    title: str
    message: str
    notification_type: str
    link: Optional[str] = None
    metadata: Optional[dict] = None
    is_read: bool = False
    is_demo: bool = False
    created_at: datetime

# ==================== SUBSCRIPTION/BILLING MODELS ====================

# Subscription Plans Configuration (server-side only - never from frontend)
# NOTE: Limits are based on total UNITS across all projects, not project count
SUBSCRIPTION_PLANS = {
    "free": {
        "name": "Free",
        "price": 0.0,
        "currency": "CHF",
        "property_limit": 2,  # Max 2 units total
        "features": [
            "Up to 2 units",
            "Client communication",
            "Document management",
            "Basic timeline tracking"
        ]
    },
    "starter": {
        "name": "Starter",
        "price": 29.0,
        "currency": "CHF",
        "property_limit": 10,  # Max 10 units total
        "features": [
            "Manage up to 10 units",
            "Full client tracking & communication",
            "Quote & invoice management",
            "Email support"
        ]
    },
    "pro": {
        "name": "Pro",
        "price": 79.0,
        "currency": "CHF",
        "property_limit": 50,  # Max 50 units total
        "features": [
            "Scale to 50 units",
            "Priority support",
            "Advanced workflow templates",
            "Team collaboration",
            "Custom branding & logo"
        ]
    },
    "enterprise": {
        "name": "Enterprise",
        "price": None,  # Custom pricing
        "currency": "CHF",
        "property_limit": None,  # Unlimited
        "features": [
            "Unlimited units",
            "Custom workflows",
            "Dedicated account manager",
            "API access & integrations"
        ]
    }
}

class SubscriptionStatus(BaseModel):
    plan_id: str
    plan_name: str
    property_limit: Optional[int]
    property_usage: int
    can_create_property: bool
    stripe_subscription_id: Optional[str] = None
    stripe_customer_id: Optional[str] = None
    subscription_status: Optional[str] = None  # active, canceled, past_due
    current_period_end: Optional[str] = None

class CreateCheckoutRequest(BaseModel):
    plan_id: str
    origin_url: str

class CheckoutStatusRequest(BaseModel):
    session_id: str

# ==================== COMPOSITE RESPONSE MODELS (Phase 2) ====================
# These models define the canonical payload for each major screen.
# Backend is the single source of truth. Frontend renders only.

class ProjectSummary(BaseModel):
    """Minimal project info for lists"""
    project_id: str
    name: str
    address: Optional[str] = None
    status: Optional[str] = None
    created_at: Optional[str] = None

class ClientSummary(BaseModel):
    """Minimal client info for context"""
    client_id: str
    name: str
    email: Optional[str] = None
    project_id: Optional[str] = None
    unit_id: Optional[str] = None

class UnitSummary(BaseModel):
    """Minimal unit info"""
    unit_id: str
    reference: str
    type: Optional[str] = None
    client_id: Optional[str] = None

class TimelineStepSummary(BaseModel):
    """Timeline step for display"""
    step_id: str
    title: str
    description: Optional[str] = None
    status: str
    order_index: int
    planned_start: Optional[str] = None
    planned_end: Optional[str] = None
    progress_percent: int = 0

class RecentWorkItem(BaseModel):
    """Recent activity item"""
    id: str
    type: str  # client, project, document
    title: str
    subtitle: Optional[str] = None
    path: str
    timestamp: Optional[str] = None

class DashboardResponse(BaseModel):
    """Canonical payload for AgentHomePage (Command Center)"""
    projects: List[ProjectSummary]
    selected_project: Optional[ProjectSummary] = None
    recent_work: List[RecentWorkItem] = []
    
class ProjectContextResponse(BaseModel):
    """Canonical payload for project context (clients + units)"""
    project: ProjectSummary
    clients: List[ClientSummary] = []
    units: List[UnitSummary] = []

class ProjectTimelineResponse(BaseModel):
    """Canonical payload for AgentTimeline"""
    project: ProjectSummary
    timeline_id: Optional[str] = None
    steps: List[TimelineStepSummary] = []
    progress_percent: int = 0

class ProjectTeamResponse(BaseModel):
    """Canonical payload for AgentTeam"""
    project: ProjectSummary
    team_members: List[dict] = []  # Using dict for flexibility

class ProjectWorkflowResponse(BaseModel):
    """Canonical payload for AgentWorkflow"""
    project: ProjectSummary
    timeline_id: Optional[str] = None
    steps: List[TimelineStepSummary] = []
    activities: List[dict] = []
    templates: List[dict] = []

# ==================== EMAIL SERVICE ====================

import resend
import asyncio

RESEND_API_KEY = os.environ.get('RESEND_API_KEY', '')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', '')  # Must be configured in production
FRONTEND_URL = os.environ.get('FRONTEND_URL') or os.environ.get('APP_URL', '')  # APP_URL set by Emergent deployment

if RESEND_API_KEY:
    resend.api_key = RESEND_API_KEY

async def send_email_async(to_email: str, subject: str, html_content: str, request: Request = None) -> dict:
    """Send email asynchronously using Resend with error monitoring"""
    if not RESEND_API_KEY:
        logger.warning(f"RESEND_API_KEY not configured. Would send email to {to_email}: {subject}")
        return {"status": "skipped", "reason": "No API key configured"}
    
    if not SENDER_EMAIL:
        logger.warning(f"SENDER_EMAIL not configured. Cannot send email to {to_email}")
        return {"status": "skipped", "reason": "No sender email configured"}
    
    params = {
        "from": SENDER_EMAIL,
        "to": [to_email],
        "subject": subject,
        "html": html_content
    }
    
    try:
        result = await asyncio.to_thread(resend.Emails.send, params)
        logger.info(f"Email sent to {to_email}: {subject}")
        return {"status": "success", "email_id": result.get("id") if isinstance(result, dict) else str(result)}
    except Exception as e:
        # Capture email errors for monitoring
        capture_email_error(e, recipient=to_email, template=subject[:50], request=request)
        logger.error(f"Failed to send email to {to_email}: {str(e)}")
        # Return graceful failure - don't crash the calling function
        return {"status": "error", "error": str(e)}

# ==================== EMAIL NOTIFICATION TEMPLATES ====================

def get_email_template(template_type: str, data: dict) -> tuple[str, str]:
    """Generate email subject and HTML content based on template type"""
    
    # Use the global FRONTEND_URL configured at startup
    frontend_url = FRONTEND_URL.rstrip('/')
    
    # Agent profile info (company name, etc.) from data
    agent_company = data.get('company_name') or data.get('agent_name') or 'Evohome'
    agent_name = data.get('agent_name') or 'Your Agent'
    agent_email = data.get('agent_email') or ''
    agent_phone = data.get('agent_phone') or ''
    
    # Build agent signature block
    signature_parts = [f"<strong>{agent_name}</strong>"]
    if agent_company and agent_company != agent_name:
        signature_parts.append(agent_company)
    if agent_email:
        signature_parts.append(f'<a href="mailto:{agent_email}" style="color: #2563EB;">{agent_email}</a>')
    if agent_phone:
        signature_parts.append(agent_phone)
    agent_signature = "<br>".join(signature_parts)
    
    # CTA button style - highly visible, cross-client compatible
    cta_button_style = (
        "display: inline-block; "
        "background-color: #2563EB; "
        "color: #FFFFFF !important; "
        "padding: 14px 28px; "
        "text-decoration: none; "
        "border-radius: 6px; "
        "font-weight: bold; "
        "font-size: 16px; "
        "mso-padding-alt: 14px 28px; "  # Outlook fix
        "text-align: center;"
    )
    
    if template_type == "document_sent":
        # When agent sends a quote/invoice to buyer
        doc_type = data.get('doc_type', 'document').capitalize()
        cta_text = f"View {doc_type}" if doc_type else "View in Platform"
        subject = f"New {doc_type} from {agent_company}: {data.get('title', 'Document')}"
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; margin: 0; padding: 0; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background-color: #2563EB; color: #FFFFFF; padding: 24px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h1 style="margin: 0; color: #FFFFFF; font-size: 24px;">{doc_type} Received</h1>
                </div>
                <div style="background-color: #FFFFFF; padding: 32px; border: 1px solid #e5e7eb; border-top: none;">
                    <p style="margin-top: 0;">Hi {data.get('buyer_name', 'there')},</p>
                    <p><strong>{agent_name}</strong> from <strong>{agent_company}</strong> has sent you a new {doc_type.lower()}.</p>
                    
                    <div style="background-color: #f9fafb; padding: 20px; border-radius: 8px; margin: 24px 0; border: 1px solid #e5e7eb;">
                        <h3 style="margin-top: 0; color: #1a1a1a;">{data.get('title', 'Document')}</h3>
                        <p style="color: #666666; margin-bottom: 12px;">{data.get('summary', '')}</p>
                        <p style="font-size: 28px; font-weight: bold; color: #2563EB; margin: 0;">{data.get('currency', 'CHF')} {data.get('amount', 0):,.2f}</p>
                        <p style="color: #666666; font-size: 14px; margin-top: 8px; margin-bottom: 0;">Project: {data.get('project_name', 'N/A')} • Unit: {data.get('unit_reference', 'N/A')}</p>
                    </div>
                    
                    <p>Please review and take action:</p>
                    
                    <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="margin: 24px auto;">
                        <tr>
                            <td style="border-radius: 6px; background-color: #2563EB;">
                                <a href="{frontend_url}/buyer/dashboard" target="_blank" style="{cta_button_style}">{cta_text}</a>
                            </td>
                        </tr>
                    </table>
                    
                    <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 24px 0;">
                    <p style="font-size: 14px; color: #666666; margin-bottom: 0;">
                        {agent_signature}
                    </p>
                </div>
                <div style="padding: 16px; text-align: center; color: #9ca3af; font-size: 12px;">
                    <p style="margin: 0;">Powered by Evohome</p>
                </div>
            </div>
        </body>
        </html>
        """
        
    elif template_type == "quote_approved":
        subject = f"Quote Approved: {data.get('title', 'Document')}"
        cta_text = "View Quote"
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; margin: 0; padding: 0; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background-color: #16a34a; color: #FFFFFF; padding: 24px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h1 style="margin: 0; color: #FFFFFF; font-size: 24px;">✓ Quote Approved</h1>
                </div>
                <div style="background-color: #FFFFFF; padding: 32px; border: 1px solid #e5e7eb; border-top: none;">
                    <p style="margin-top: 0;">Great news!</p>
                    <p><strong>{data.get('buyer_name', 'Your client')}</strong> has approved the quote <strong>{data.get('document_number', '')}</strong>.</p>
                    
                    <div style="background-color: #f9fafb; padding: 20px; border-radius: 8px; margin: 24px 0; border: 1px solid #e5e7eb;">
                        <h3 style="margin-top: 0; color: #1a1a1a;">{data.get('title', 'Quote')}</h3>
                        <p style="font-size: 28px; font-weight: bold; color: #2563EB; margin: 8px 0;">{data.get('currency', 'CHF')} {data.get('amount', 0):,.2f}</p>
                        <span style="display: inline-block; padding: 4px 12px; border-radius: 20px; font-size: 12px; font-weight: 600; background-color: #dcfce7; color: #166534;">Approved</span>
                    </div>
                    
                    <p>You can now convert this quote to an invoice.</p>
                    
                    <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="margin: 24px auto;">
                        <tr>
                            <td style="border-radius: 6px; background-color: #2563EB;">
                                <a href="{frontend_url}/agent/documents/{data.get('document_id', '')}" target="_blank" style="{cta_button_style}">{cta_text}</a>
                            </td>
                        </tr>
                    </table>
                </div>
                <div style="padding: 16px; text-align: center; color: #9ca3af; font-size: 12px;">
                    <p style="margin: 0;">Evohome - Real Estate Management</p>
                </div>
            </div>
        </body>
        </html>
        """
        
    elif template_type == "change_requested":
        subject = f"Changes Requested: {data.get('title', 'Document')}"
        cta_text = "View & Revise"
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; margin: 0; padding: 0; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background-color: #f59e0b; color: #FFFFFF; padding: 24px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h1 style="margin: 0; color: #FFFFFF; font-size: 24px;">Changes Requested</h1>
                </div>
                <div style="background-color: #FFFFFF; padding: 32px; border: 1px solid #e5e7eb; border-top: none;">
                    <p style="margin-top: 0;"><strong>{data.get('buyer_name', 'Your client')}</strong> has requested changes to <strong>{data.get('document_number', 'the document')}</strong>.</p>
                    
                    <div style="background-color: #fef3c7; padding: 20px; border-radius: 8px; margin: 24px 0; border-left: 4px solid #f59e0b;">
                        <h4 style="margin-top: 0; color: #92400e;">Client's Comment:</h4>
                        <p style="font-style: italic; color: #78350f; margin-bottom: 0;">"{data.get('comment', 'No comment provided')}"</p>
                    </div>
                    
                    <p>Please review the feedback and upload a revised document.</p>
                    
                    <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="margin: 24px auto;">
                        <tr>
                            <td style="border-radius: 6px; background-color: #2563EB;">
                                <a href="{frontend_url}/agent/documents/{data.get('document_id', '')}" target="_blank" style="{cta_button_style}">{cta_text}</a>
                            </td>
                        </tr>
                    </table>
                </div>
                <div style="padding: 16px; text-align: center; color: #9ca3af; font-size: 12px;">
                    <p style="margin: 0;">Evohome - Real Estate Management</p>
                </div>
            </div>
        </body>
        </html>
        """
        
    elif template_type == "payment_confirmed":
        subject = f"Payment Confirmed: {data.get('title', 'Invoice')}"
        cta_text = "View Invoice"
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; margin: 0; padding: 0; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background-color: #16a34a; color: #FFFFFF; padding: 24px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h1 style="margin: 0; color: #FFFFFF; font-size: 24px;">💰 Payment Received</h1>
                </div>
                <div style="background-color: #FFFFFF; padding: 32px; border: 1px solid #e5e7eb; border-top: none;">
                    <p style="margin-top: 0;">Payment has been confirmed for invoice <strong>{data.get('document_number', '')}</strong>.</p>
                    
                    <div style="background-color: #f9fafb; padding: 20px; border-radius: 8px; margin: 24px 0; border: 1px solid #e5e7eb;">
                        <h3 style="margin-top: 0; color: #1a1a1a;">{data.get('title', 'Invoice')}</h3>
                        <p style="font-size: 28px; font-weight: bold; color: #16a34a; margin: 8px 0;">{data.get('currency', 'CHF')} {data.get('amount', 0):,.2f}</p>
                        <span style="display: inline-block; padding: 4px 12px; border-radius: 20px; font-size: 12px; font-weight: 600; background-color: #dcfce7; color: #166534;">Paid</span>
                    </div>
                    
                    <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="margin: 24px auto;">
                        <tr>
                            <td style="border-radius: 6px; background-color: #2563EB;">
                                <a href="{frontend_url}/agent/documents/{data.get('document_id', '')}" target="_blank" style="{cta_button_style}">{cta_text}</a>
                            </td>
                        </tr>
                    </table>
                </div>
                <div style="padding: 16px; text-align: center; color: #9ca3af; font-size: 12px;">
                    <p style="margin: 0;">Evohome - Real Estate Management</p>
                </div>
            </div>
        </body>
        </html>
        """
        
    elif template_type == "new_message":
        subject = f"New Message from {data.get('sender_name', 'Your Agent')}"
        cta_text = "View Message"
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; margin: 0; padding: 0; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background-color: #2563EB; color: #FFFFFF; padding: 24px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h1 style="margin: 0; color: #FFFFFF; font-size: 24px;">New Message</h1>
                </div>
                <div style="background-color: #FFFFFF; padding: 32px; border: 1px solid #e5e7eb; border-top: none;">
                    <p style="margin-top: 0;"><strong>{data.get('sender_name', 'Someone')}</strong> sent you a message:</p>
                    
                    <div style="background-color: #f9fafb; padding: 20px; border-radius: 8px; margin: 24px 0; border-left: 4px solid #2563EB;">
                        <p style="margin: 0; color: #333333;">{data.get('message_preview', '')[:200]}...</p>
                    </div>
                    
                    <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="margin: 24px auto;">
                        <tr>
                            <td style="border-radius: 6px; background-color: #2563EB;">
                                <a href="{frontend_url}/{data.get('link', 'buyer/dashboard')}" target="_blank" style="{cta_button_style}">{cta_text}</a>
                            </td>
                        </tr>
                    </table>
                </div>
                <div style="padding: 16px; text-align: center; color: #9ca3af; font-size: 12px;">
                    <p style="margin: 0;">Evohome - Real Estate Management</p>
                </div>
            </div>
        </body>
        </html>
        """
    elif template_type == "feed_update":
        # When agent posts a feed update to buyer
        subject = f"New Update from {data.get('agent_name', 'Your Agent')} - {data.get('project_name', 'Your Project')}"
        cta_text = "View Update"
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; margin: 0; padding: 0; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background-color: #2563EB; color: #FFFFFF; padding: 24px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h1 style="margin: 0; color: #FFFFFF; font-size: 24px;">New Update</h1>
                </div>
                <div style="background-color: #FFFFFF; padding: 32px; border: 1px solid #e5e7eb; border-top: none;">
                    <p style="margin-top: 0;">Hi {data.get('buyer_name', 'there')},</p>
                    <p>Your agent has posted a new update for <strong>{data.get('project_name', 'your project')}</strong>.</p>
                    
                    <div style="background-color: #f9fafb; padding: 20px; border-radius: 8px; margin: 24px 0; border-left: 4px solid #2563EB;">
                        <p style="margin: 0; color: #333333;">{data.get('message_preview', 'View the full update in your dashboard.')}</p>
                    </div>
                    
                    <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="margin: 24px auto;">
                        <tr>
                            <td style="border-radius: 6px; background-color: #2563EB;">
                                <a href="{frontend_url}/{data.get('link', 'buyer/dashboard')}" target="_blank" style="{cta_button_style}">{cta_text}</a>
                            </td>
                        </tr>
                    </table>
                    
                    <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 24px 0;">
                    <p style="font-size: 14px; color: #666666; margin-bottom: 0;">
                        {agent_signature}
                    </p>
                </div>
                <div style="padding: 16px; text-align: center; color: #9ca3af; font-size: 12px;">
                    <p style="margin: 0;">Evohome - Real Estate Management</p>
                </div>
            </div>
        </body>
        </html>
        """
    elif template_type == "milestone_completed":
        # When a construction milestone is completed
        subject = f"Milestone Reached: {data.get('milestone_name', 'Construction Update')}"
        cta_text = "View Progress"
        progress_percent = data.get('progress_percent', 0)
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; margin: 0; padding: 0; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background-color: #16a34a; color: #FFFFFF; padding: 24px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h1 style="margin: 0; color: #FFFFFF; font-size: 24px;">Milestone Completed</h1>
                </div>
                <div style="background-color: #FFFFFF; padding: 32px; border: 1px solid #e5e7eb; border-top: none;">
                    <p style="margin-top: 0;">Hi {data.get('buyer_name', 'there')},</p>
                    <p>Great news! A construction milestone for your property has been completed.</p>
                    
                    <div style="background-color: #f0fdf4; padding: 20px; border-radius: 8px; margin: 24px 0; border: 1px solid #bbf7d0;">
                        <h3 style="margin-top: 0; color: #166534;">{data.get('milestone_name', 'Milestone')}</h3>
                        <p style="color: #666666; margin-bottom: 12px;">{data.get('milestone_description', '')}</p>
                        <p style="font-size: 14px; color: #166534; margin: 0;">
                            <strong>Project:</strong> {data.get('project_name', 'N/A')} &bull; <strong>Unit:</strong> {data.get('unit_reference', 'N/A')}
                        </p>
                    </div>
                    
                    <div style="margin: 24px 0;">
                        <p style="margin-bottom: 8px; font-weight: 600;">Overall Progress: {progress_percent}%</p>
                        <div style="background-color: #e5e7eb; border-radius: 9999px; height: 12px; overflow: hidden;">
                            <div style="background-color: #16a34a; height: 100%; width: {progress_percent}%; border-radius: 9999px;"></div>
                        </div>
                    </div>
                    
                    <p>Log in to view the full construction timeline and details.</p>
                    
                    <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="margin: 24px auto;">
                        <tr>
                            <td style="border-radius: 6px; background-color: #2563EB;">
                                <a href="{frontend_url}/buyer/dashboard" target="_blank" style="{cta_button_style}">{cta_text}</a>
                            </td>
                        </tr>
                    </table>
                    
                    <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 24px 0;">
                    <p style="font-size: 14px; color: #666666; margin-bottom: 0;">
                        {agent_signature}
                    </p>
                </div>
                <div style="padding: 16px; text-align: center; color: #9ca3af; font-size: 12px;">
                    <p style="margin: 0;">Evohome - Real Estate Management</p>
                </div>
            </div>
        </body>
        </html>
        """
    else:
        subject = "Notification from Evohome"
        cta_text = "View in Platform"
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; margin: 0; padding: 0; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background-color: #2563EB; color: #FFFFFF; padding: 24px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h1 style="margin: 0; color: #FFFFFF; font-size: 24px;">Notification</h1>
                </div>
                <div style="background-color: #FFFFFF; padding: 32px; border: 1px solid #e5e7eb; border-top: none;">
                    <p style="margin-top: 0;">{data.get('message', 'You have a new notification.')}</p>
                    
                    <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="margin: 24px auto;">
                        <tr>
                            <td style="border-radius: 6px; background-color: #2563EB;">
                                <a href="{frontend_url}" target="_blank" style="{cta_button_style}">{cta_text}</a>
                            </td>
                        </tr>
                    </table>
                </div>
                <div style="padding: 16px; text-align: center; color: #9ca3af; font-size: 12px;">
                    <p style="margin: 0;">Evohome - Real Estate Management</p>
                </div>
            </div>
        </body>
        </html>
        """
    
    return subject, html

async def send_notification_email(template_type: str, to_email: str, data: dict) -> dict:
    """Send a notification email using a template"""
    subject, html = get_email_template(template_type, data)
    return await send_email_async(to_email, subject, html)

async def create_notification(user_id: str, title: str, message: str, notification_type: str, 
                             link: str = None, metadata: dict = None, is_demo: bool = False) -> str:
    """Create a notification in the database"""
    notification_id = f"notif_{uuid.uuid4().hex[:12]}"
    notification_doc = {
        "notification_id": notification_id,
        "user_id": user_id,
        "title": title,
        "message": message,
        "notification_type": notification_type,
        "link": link,
        "metadata": metadata or {},
        "is_read": False,
        "is_demo": is_demo,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.notifications.insert_one(notification_doc)
    return notification_id

# ==================== WEBSOCKET MANAGER ====================

class ConnectionManager:
    """Manages WebSocket connections for real-time updates"""
    
    def __init__(self):
        # Maps user_id to list of WebSocket connections
        self.active_connections: dict[str, list[WebSocket]] = {}
    
    async def connect(self, websocket: WebSocket, user_id: str):
        await websocket.accept()
        if user_id not in self.active_connections:
            self.active_connections[user_id] = []
        self.active_connections[user_id].append(websocket)
        logger.info(f"WebSocket connected for user: {user_id}")
    
    def disconnect(self, websocket: WebSocket, user_id: str):
        if user_id in self.active_connections:
            if websocket in self.active_connections[user_id]:
                self.active_connections[user_id].remove(websocket)
            if not self.active_connections[user_id]:
                del self.active_connections[user_id]
        logger.info(f"WebSocket disconnected for user: {user_id}")
    
    async def send_to_user(self, user_id: str, message: dict):
        """Send a message to all connections of a specific user"""
        if user_id in self.active_connections:
            disconnected = []
            for connection in self.active_connections[user_id]:
                try:
                    await connection.send_json(message)
                except Exception as e:
                    logger.error(f"Failed to send WebSocket message: {e}")
                    disconnected.append(connection)
            # Clean up disconnected
            for conn in disconnected:
                self.active_connections[user_id].remove(conn)
    
    async def broadcast_to_users(self, user_ids: list[str], message: dict):
        """Send a message to multiple users"""
        for user_id in user_ids:
            await self.send_to_user(user_id, message)

ws_manager = ConnectionManager()

async def notify_realtime(user_ids: list[str], event_type: str, data: dict):
    """Helper to send real-time notifications via WebSocket"""
    message = {
        "type": event_type,
        "data": data,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }
    await ws_manager.broadcast_to_users(user_ids, message)

async def send_milestone_notification(step: dict, project: dict, timeline: dict, user: dict, is_demo: bool):
    """Send notifications when a construction milestone is completed"""
    try:
        # Get all buyers linked to this project's units
        units = await db.units.find(
            {"project_id": project['project_id']},
            {"_id": 0, "unit_id": 1, "reference": 1}
        ).to_list(100)
        
        unit_ids = [u['unit_id'] for u in units]
        unit_refs = {u['unit_id']: u.get('reference', 'Unit') for u in units}
        
        # Find clients linked to these units
        clients = await db.clients.find(
            {"unit_id": {"$in": unit_ids}},
            {"_id": 0, "client_id": 1, "buyer_id": 1, "name": 1, "email": 1, "unit_id": 1}
        ).to_list(100)
        
        if not clients:
            logger.info(f"No clients to notify for milestone completion: {step.get('title')}")
            return
        
        # Get agent settings for email signature
        agent_settings = await db.agent_settings.find_one(
            {"agent_id": user['user_id']},
            {"_id": 0}
        ) or {}
        
        # Calculate overall progress
        all_steps = await db.timeline_steps.find(
            {"project_timeline_id": step['project_timeline_id']},
            {"_id": 0, "status": 1}
        ).to_list(100)
        
        completed_count = sum(1 for s in all_steps if s['status'] in ['completed', 'approved'])
        total_count = len(all_steps)
        progress_percent = round((completed_count / total_count) * 100) if total_count > 0 else 0
        
        # Send notification to each client
        for client in clients:
            buyer_id = client.get('buyer_id')
            if not buyer_id:
                continue
            
            unit_ref = unit_refs.get(client.get('unit_id'), 'Your Unit')
            
            # Create in-app notification
            await create_notification(
                user_id=buyer_id,
                title=f"Milestone Reached: {step.get('title', 'Construction Update')}",
                message=f"The '{step.get('title')}' phase has been completed for {unit_ref}. Overall progress: {progress_percent}%",
                notification_type="milestone_completed",
                link="/buyer/dashboard",
                is_demo=is_demo,
                metadata={
                    "step_id": step.get('step_id'),
                    "project_id": project.get('project_id'),
                    "progress_percent": progress_percent
                }
            )
            
            # Send email notification
            if client.get('email'):
                email_data = {
                    "buyer_name": client.get('name', 'there'),
                    "milestone_name": step.get('title', 'Construction Phase'),
                    "milestone_description": step.get('description', ''),
                    "project_name": project.get('name', 'Your Project'),
                    "unit_reference": unit_ref,
                    "progress_percent": progress_percent,
                    "agent_name": agent_settings.get('profile', {}).get('display_name') or user.get('name', 'Your Agent'),
                    "company_name": agent_settings.get('company_name', ''),
                    "agent_email": agent_settings.get('profile', {}).get('contact_email', ''),
                    "agent_phone": agent_settings.get('profile', {}).get('contact_phone', '')
                }
                
                try:
                    await send_notification_email("milestone_completed", client['email'], email_data)
                    logger.info(f"Sent milestone email to {client['email']} for step {step.get('step_id')}")
                except Exception as e:
                    logger.error(f"Failed to send milestone email to {client['email']}: {e}")
            
            # Send real-time WebSocket notification
            await notify_realtime(
                [buyer_id],
                "milestone_completed",
                {
                    "step_id": step.get('step_id'),
                    "step_title": step.get('title'),
                    "progress_percent": progress_percent
                }
            )
        
        logger.info(f"Sent milestone notifications to {len(clients)} clients for step: {step.get('title')}")
        
    except Exception as e:
        logger.error(f"Failed to send milestone notifications: {e}")

# ==================== SWISS QR CODE HELPER ====================

from qrbill import QRBill
import base64

# Default company details (fallback)
DEFAULT_IBAN = "CH9300762011623852957"
DEFAULT_COMPANY_NAME = "Evohome SA"
DEFAULT_COMPANY_ADDRESS = "Rue du Rhône 1"
DEFAULT_COMPANY_PCODE = "1204"
DEFAULT_COMPANY_CITY = "Genève"
DEFAULT_COMPANY_COUNTRY = "CH"

def generate_swiss_qr_code(
    amount: float, 
    reference: str, 
    buyer_name: str = None,
    iban: str = None,
    creditor_name: str = None,
    creditor_address: str = None,
    creditor_pcode: str = None,
    creditor_city: str = None
) -> bytes:
    """Generate Swiss QR bill code as SVG bytes"""
    try:
        # Use provided values or defaults
        account_iban = iban or DEFAULT_IBAN
        creditor = {
            'name': creditor_name or DEFAULT_COMPANY_NAME,
            'street': creditor_address or DEFAULT_COMPANY_ADDRESS,
            'pcode': creditor_pcode or DEFAULT_COMPANY_PCODE,
            'city': creditor_city or DEFAULT_COMPANY_CITY,
            'country': DEFAULT_COMPANY_COUNTRY,
        }
        
        # Build the QR bill with NON reference type (unstructured)
        # since our invoice numbers don't conform to QRR/SCOR formats
        bill = QRBill(
            account=account_iban,
            creditor=creditor,
            amount=str(round(amount, 2)),
            currency='CHF',
            # Use unstructured message for reference instead of structured reference_number
            additional_information=f"Invoice {reference}" if reference else None
        )
        
        with tempfile.NamedTemporaryFile(suffix='.svg', delete=False) as tmp:
            bill.as_svg(tmp.name)
            with open(tmp.name, 'rb') as f:
                return f.read()
    except Exception as e:
        logger.error(f"Failed to generate QR code: {e}")
        return None

def generate_swiss_qr_code_base64(
    amount: float, 
    reference: str, 
    buyer_name: str = None,
    iban: str = None,
    creditor_name: str = None,
    creditor_address: str = None,
    creditor_pcode: str = None,
    creditor_city: str = None
) -> str:
    """Generate Swiss QR bill code as base64 SVG string for frontend display"""
    svg_bytes = generate_swiss_qr_code(
        amount, reference, buyer_name,
        iban, creditor_name, creditor_address, creditor_pcode, creditor_city
    )
    if svg_bytes:
        return base64.b64encode(svg_bytes).decode('utf-8')
    return None

# ==================== AUTH HELPERS ====================
# NOTE: Primary auth functions are now in core/auth.py
# These legacy functions are kept ONLY for backward compatibility during migration

def create_jwt_token(user_id: str, role: str, is_demo: bool = False) -> str:
    """
    DEPRECATED: Use create_access_token from core.auth instead.
    Kept for backward compatibility - redirects to centralized auth.
    NOTE: is_demo parameter is IGNORED - is_demo is looked up from DB on each request.
    """
    return create_access_token(user_id, role)

def verify_jwt_token(token: str) -> dict:
    """
    DEPRECATED: Use verify_token from core.auth instead.
    Kept for backward compatibility in session endpoint.
    """
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# NOTE: get_current_user, get_current_agent, get_current_buyer 
# are now imported from core.auth - DO NOT redefine here

# ==================== WEBSOCKET ENDPOINT ====================

@api_router.websocket("/ws/{user_id}")
async def websocket_endpoint(websocket: WebSocket, user_id: str):
    """
    WebSocket endpoint for real-time updates.
    
    SECURITY: Validates JWT token from query parameter before accepting connection.
    Token must match the user_id in the URL path.
    """
    # Extract token from query params
    token = websocket.query_params.get("token")
    
    if not token:
        await websocket.close(code=4001, reason="Authentication required")
        capture_websocket_error(Exception("No token provided"), user_id=user_id, action="connect")
        return
    
    # Verify token and extract user_id
    try:
        payload = verify_token(token, expected_type='access')
        token_user_id = payload.get('user_id')
        
        # Ensure token user_id matches URL user_id - prevents impersonation
        if token_user_id != user_id:
            await websocket.close(code=4003, reason="User ID mismatch")
            capture_websocket_error(
                Exception(f"Token user_id {token_user_id} does not match URL user_id {user_id}"),
                user_id=user_id, 
                action="connect"
            )
            return
            
    except HTTPException as e:
        await websocket.close(code=4001, reason=str(e.detail))
        capture_websocket_error(e, user_id=user_id, action="connect")
        return
    except Exception as e:
        await websocket.close(code=4001, reason="Invalid token")
        capture_websocket_error(e, user_id=user_id, action="connect")
        return
    
    # Verify user exists in database
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user:
        await websocket.close(code=4004, reason="User not found")
        capture_websocket_error(Exception("User not found"), user_id=user_id, action="connect")
        return
    
    await ws_manager.connect(websocket, user_id)
    try:
        while True:
            # Keep connection alive, listen for pings
            data = await websocket.receive_text()
            if data == "ping":
                await websocket.send_text("pong")
    except WebSocketDisconnect:
        ws_manager.disconnect(websocket, user_id)
    except Exception as e:
        capture_websocket_error(e, user_id=user_id, action="message")
        ws_manager.disconnect(websocket, user_id)

# ==================== AUTH ENDPOINTS ====================

@api_router.post("/auth/register")
async def register_agent(data: AgentRegister, response: Response, request: Request):
    """Register a new agent with email/password, or link password to existing OAuth account"""
    # Rate limit: 3 registrations per minute per IP
    rate_limit_check(request, "auth_register")
    
    existing = await db.users.find_one({"email": data.email, "role": "agent"}, {"_id": 0})
    
    if existing:
        # Check if this is an OAuth-only account (no password set)
        if not existing.get('password_hash'):
            # Link password to existing OAuth account
            hashed_password = bcrypt.hashpw(data.password.encode('utf-8'), bcrypt.gensalt())
            await db.users.update_one(
                {"user_id": existing['user_id']},
                {"$set": {
                    "password_hash": hashed_password.decode('utf-8'),
                    "name": data.name if data.name else existing['name']  # Update name if provided
                }}
            )
            
            token = create_jwt_token(existing['user_id'], "agent", existing.get('is_demo', False))
            
            response.set_cookie(
                key="session_token",
                value=token,
                httponly=True,
                secure=True,
                samesite="none",
                path="/",
                max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
            )
            
            return {
                "user_id": existing['user_id'],
                "email": data.email,
                "name": data.name if data.name else existing['name'],
                "role": "agent",
                "is_demo": existing.get('is_demo', False),
                "token": token,
                "account_linked": True,  # Indicates password was added to existing OAuth account
                "message": "Password successfully added to your Google account. You can now login with email/password."
            }
        else:
            # Account already has a password
            raise HTTPException(
                status_code=400, 
                detail="An account with this email already exists. Please login instead, or use 'Forgot Password' to reset."
            )
    
    # New account - create with email/password
    hashed_password = bcrypt.hashpw(data.password.encode('utf-8'), bcrypt.gensalt())
    user_id = f"agent_{uuid.uuid4().hex[:12]}"
    
    user_doc = {
        "user_id": user_id,
        "email": data.email,
        "name": data.name,
        "password_hash": hashed_password.decode('utf-8'),
        "role": "agent",
        "picture": None,
        "is_demo": False,
        "subscription_plan": "free",
        "subscription_status": "active",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    token = create_jwt_token(user_id, "agent", False)
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    return {
        "user_id": user_id,
        "email": data.email,
        "name": data.name,
        "role": "agent",
        "is_demo": False,
        "token": token
    }

@api_router.post("/auth/login")
async def login_agent(data: AgentLogin, response: Response, request: Request):
    """Login agent with email/password"""
    # Rate limit: 5 attempts per minute per IP
    rate_limit_check(request, "auth_login")
    
    user = await db.users.find_one({"email": data.email, "role": "agent"}, {"_id": 0})
    if not user:
        capture_auth_failure("invalid_credentials", email=data.email, request=request)
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check if user has a password set (OAuth-only accounts won't have one)
    if not user.get('password_hash'):
        capture_auth_failure("oauth_only_account", email=data.email, request=request)
        raise HTTPException(
            status_code=401, 
            detail="This account was created with Google. Please login with Google, or create a password by clicking 'Create Account' with the same email."
        )
    
    if not bcrypt.checkpw(data.password.encode('utf-8'), user['password_hash'].encode('utf-8')):
        capture_auth_failure("wrong_password", email=data.email, request=request)
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_jwt_token(user['user_id'], "agent", user.get('is_demo', False))
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    return {
        "user_id": user['user_id'],
        "email": user['email'],
        "name": user['name'],
        "role": "agent",
        "is_demo": user.get('is_demo', False),
        "token": token
    }

# ==================== BUYER EMAIL/PASSWORD AUTH ====================

@api_router.post("/auth/buyer/register")
async def register_buyer(data: BuyerRegister, response: Response):
    """Register a new buyer with email/password, or link password to existing OAuth account"""
    existing = await db.users.find_one({"email": data.email, "role": "buyer"}, {"_id": 0})
    
    if existing:
        # Check if this is an OAuth-only account (no password set)
        if not existing.get('password_hash'):
            # Link password to existing OAuth account
            hashed_password = bcrypt.hashpw(data.password.encode('utf-8'), bcrypt.gensalt())
            await db.users.update_one(
                {"user_id": existing['user_id']},
                {"$set": {
                    "password_hash": hashed_password.decode('utf-8'),
                    "name": data.name if data.name else existing['name']
                }}
            )
            
            token = create_jwt_token(existing['user_id'], "buyer", existing.get('is_demo', False))
            
            response.set_cookie(
                key="session_token",
                value=token,
                httponly=True,
                secure=True,
                samesite="none",
                path="/",
                max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
            )
            
            return {
                "user_id": existing['user_id'],
                "email": data.email,
                "name": data.name if data.name else existing['name'],
                "role": "buyer",
                "is_demo": existing.get('is_demo', False),
                "token": token,
                "account_linked": True,
                "message": "Password successfully added to your Google account. You can now login with email/password."
            }
        else:
            # Account already has a password
            raise HTTPException(
                status_code=400, 
                detail="An account with this email already exists. Please login instead, or use 'Forgot Password' to reset."
            )
    
    # New account - create with email/password
    hashed_password = bcrypt.hashpw(data.password.encode('utf-8'), bcrypt.gensalt())
    user_id = f"buyer_{uuid.uuid4().hex[:12]}"
    
    user_doc = {
        "user_id": user_id,
        "email": data.email,
        "name": data.name,
        "password_hash": hashed_password.decode('utf-8'),
        "role": "buyer",
        "picture": None,
        "is_demo": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    # Auto-link buyer to any existing client records with matching email
    await db.clients.update_many(
        {"email": data.email, "buyer_id": None, "is_demo": False},
        {"$set": {"buyer_id": user_id}}
    )
    await db.clients.update_many(
        {"email": data.email, "buyer_id": {"$exists": False}, "is_demo": False},
        {"$set": {"buyer_id": user_id}}
    )
    
    # Also handle invitation code if provided
    if data.invitation_code:
        client = await db.clients.find_one({"invitation_code": data.invitation_code}, {"_id": 0})
        if client and not client.get('buyer_id'):
            await db.clients.update_one(
                {"client_id": client['client_id']},
                {"$set": {"buyer_id": user_id}}
            )
    
    logger.info(f"Buyer registered: {data.email} (user_id: {user_id})")
    
    token = create_jwt_token(user_id, "buyer", False)
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    return {
        "user_id": user_id,
        "email": data.email,
        "name": data.name,
        "role": "buyer",
        "is_demo": False,
        "token": token
    }

@api_router.post("/auth/buyer/login")
async def login_buyer(data: BuyerLogin, response: Response, request: Request):
    """Login buyer with email/password"""
    # Rate limit: 5 attempts per minute per IP
    rate_limit_check(request, "auth_login")
    
    user = await db.users.find_one({"email": data.email, "role": "buyer"}, {"_id": 0})
    if not user:
        capture_auth_failure("invalid_credentials", email=data.email, request=request)
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check if user has a password set (OAuth-only accounts won't have one)
    if not user.get('password_hash'):
        capture_auth_failure("oauth_only_account", email=data.email, request=request)
        raise HTTPException(
            status_code=401, 
            detail="This account was created with Google. Please login with Google, or create a password by clicking 'Create Account' with the same email."
        )
    
    if not bcrypt.checkpw(data.password.encode('utf-8'), user['password_hash'].encode('utf-8')):
        capture_auth_failure("wrong_password", email=data.email, request=request)
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_jwt_token(user['user_id'], "buyer", user.get('is_demo', False))
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    return {
        "user_id": user['user_id'],
        "email": user['email'],
        "name": user['name'],
        "role": "buyer",
        "is_demo": user.get('is_demo', False),
        "token": token
    }

@api_router.post("/auth/session")
async def exchange_session(request: Request, response: Response):
    """Exchange Emergent OAuth session_id for app session - with role enforcement"""
    body = await request.json()
    session_id = body.get('session_id')
    intended_role = body.get('intended_role')  # Role user selected at login
    
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id required")
    
    # OAuth backend URL - configurable for different environments
    OAUTH_BACKEND_URL = os.environ.get('OAUTH_BACKEND_URL', 'https://demobackend.emergentagent.com')
    
    async with httpx.AsyncClient() as client_http:
        try:
            resp = await client_http.get(
                f"{OAUTH_BACKEND_URL}/auth/v1/env/oauth/session-data",
                headers={"X-Session-ID": session_id}
            )
            if resp.status_code != 200:
                raise HTTPException(status_code=401, detail="Invalid session")
            
            oauth_data = resp.json()
        except Exception as e:
            logger.error(f"OAuth exchange failed: {e}")
            raise HTTPException(status_code=401, detail="Authentication failed")
    
    email = oauth_data.get('email')
    
    # Check if user exists in our system
    existing_agent = await db.users.find_one({"email": email, "role": "agent"}, {"_id": 0})
    existing_buyer = await db.users.find_one({"email": email, "role": "buyer"}, {"_id": 0})
    client_link = await db.clients.find_one({"email": email}, {"_id": 0})
    
    # Determine if the account is authorized for the intended role
    if intended_role == 'agent':
        # Agent login requested
        if existing_agent:
            # Known agent - allow
            user_id = existing_agent['user_id']
            role = "agent"
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {
                    "name": oauth_data.get('name', existing_agent['name']),
                    "picture": oauth_data.get('picture')
                }}
            )
        elif existing_buyer:
            # This email is registered as buyer, not agent
            raise HTTPException(
                status_code=403, 
                detail="This account is registered as a buyer. Please use 'Login as Buyer' instead."
            )
        else:
            # New user trying to login as agent - reject (agents must be pre-registered)
            raise HTTPException(
                status_code=403, 
                detail="This email is not authorized as an agent. Contact your administrator to register."
            )
    
    elif intended_role == 'buyer':
        # Buyer login requested
        if existing_buyer:
            # Known buyer - allow
            user_id = existing_buyer['user_id']
            role = "buyer"
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {
                    "name": oauth_data.get('name', existing_buyer['name']),
                    "picture": oauth_data.get('picture')
                }}
            )
        elif existing_agent:
            # This email is registered as agent, not buyer
            raise HTTPException(
                status_code=403, 
                detail="This account is registered as an agent. Please use 'Login as Agent' instead."
            )
        elif client_link:
            # New user but their email is linked to a client - create as buyer
            user_id = f"buyer_{uuid.uuid4().hex[:12]}"
            role = "buyer"
            user_doc = {
                "user_id": user_id,
                "email": email,
                "name": oauth_data.get('name', 'User'),
                "picture": oauth_data.get('picture'),
                "role": role,
                "is_demo": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(user_doc)
            await db.clients.update_one(
                {"client_id": client_link['client_id']},
                {"$set": {"buyer_id": user_id}}
            )
        else:
            # New user with no client link - create as buyer (they can be linked later)
            user_id = f"buyer_{uuid.uuid4().hex[:12]}"
            role = "buyer"
            user_doc = {
                "user_id": user_id,
                "email": email,
                "name": oauth_data.get('name', 'User'),
                "picture": oauth_data.get('picture'),
                "role": role,
                "is_demo": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(user_doc)
    
    else:
        # No intended role specified - fallback to legacy behavior (determine by account)
        if existing_agent:
            user_id = existing_agent['user_id']
            role = "agent"
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {
                    "name": oauth_data.get('name', existing_agent['name']),
                    "picture": oauth_data.get('picture')
                }}
            )
        elif existing_buyer:
            user_id = existing_buyer['user_id']
            role = "buyer"
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {
                    "name": oauth_data.get('name', existing_buyer['name']),
                    "picture": oauth_data.get('picture')
                }}
            )
        elif client_link:
            user_id = f"buyer_{uuid.uuid4().hex[:12]}"
            role = "buyer"
            user_doc = {
                "user_id": user_id,
                "email": email,
                "name": oauth_data.get('name', 'User'),
                "picture": oauth_data.get('picture'),
                "role": role,
                "is_demo": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(user_doc)
            await db.clients.update_one(
                {"client_id": client_link['client_id']},
                {"$set": {"buyer_id": user_id}}
            )
        else:
            user_id = f"buyer_{uuid.uuid4().hex[:12]}"
            role = "buyer"
            user_doc = {
                "user_id": user_id,
                "email": email,
                "name": oauth_data.get('name', 'User'),
                "picture": oauth_data.get('picture'),
                "role": role,
                "is_demo": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(user_doc)
    
    token = create_jwt_token(user_id, role, False)
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
    
    return {**user, "token": token}

# Google OAuth configuration
GOOGLE_CLIENT_ID = "1053944720664-ulkunc0n9qhvfq4cdntro2ks0lsru0p6.apps.googleusercontent.com"
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET', '')

@api_router.post("/auth/google/callback")
async def google_oauth_callback(request: Request, response: Response):
    """Handle Google OAuth callback - exchange code for tokens and create/login user"""
    body = await request.json()
    code = body.get('code')
    state = body.get('state', '{}')
    redirect_uri = body.get('redirect_uri')
    
    if not code:
        raise HTTPException(status_code=400, detail="Authorization code required")
    
    if not GOOGLE_CLIENT_SECRET:
        raise HTTPException(status_code=500, detail="Google OAuth not configured - missing client secret")
    
    # Parse state to get intended role
    try:
        state_data = json.loads(state)
        intended_role = state_data.get('role', 'buyer')
    except (json.JSONDecodeError, ValueError):
        intended_role = 'buyer'
    
    # Exchange code for tokens
    async with httpx.AsyncClient() as client_http:
        try:
            token_response = await client_http.post(
                "https://oauth2.googleapis.com/token",
                data={
                    "code": code,
                    "client_id": GOOGLE_CLIENT_ID,
                    "client_secret": GOOGLE_CLIENT_SECRET,
                    "redirect_uri": redirect_uri,
                    "grant_type": "authorization_code"
                }
            )
            
            if token_response.status_code != 200:
                logger.error(f"Google token exchange failed: {token_response.text}")
                raise HTTPException(status_code=401, detail="Failed to exchange authorization code")
            
            tokens = token_response.json()
            access_token = tokens.get('access_token')
            
            # Get user info from Google
            userinfo_response = await client_http.get(
                "https://www.googleapis.com/oauth2/v2/userinfo",
                headers={"Authorization": f"Bearer {access_token}"}
            )
            
            if userinfo_response.status_code != 200:
                raise HTTPException(status_code=401, detail="Failed to get user info from Google")
            
            google_user = userinfo_response.json()
            
        except httpx.RequestError as e:
            logger.error(f"Google OAuth request failed: {e}")
            raise HTTPException(status_code=401, detail="Authentication failed")
    
    email = google_user.get('email')
    name = google_user.get('name', 'User')
    picture = google_user.get('picture')
    
    if not email:
        raise HTTPException(status_code=400, detail="Email not provided by Google")
    
    # Check if user exists
    existing_agent = await db.users.find_one({"email": email, "role": "agent"}, {"_id": 0})
    existing_buyer = await db.users.find_one({"email": email, "role": "buyer"}, {"_id": 0})
    client_link = await db.clients.find_one({"email": email}, {"_id": 0})
    
    # Handle role-based login/registration
    if intended_role == 'agent':
        if existing_agent:
            user_id = existing_agent['user_id']
            role = "agent"
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {"name": name, "picture": picture}}
            )
        elif existing_buyer:
            raise HTTPException(
                status_code=403, 
                detail="This account is registered as a buyer. Please use 'Login as Buyer' instead."
            )
        else:
            # Auto-register new agents via Google OAuth
            user_id = f"agent_{uuid.uuid4().hex[:12]}"
            role = "agent"
            user_doc = {
                "user_id": user_id,
                "email": email,
                "name": name,
                "picture": picture,
                "role": role,
                "is_demo": False,
                "subscription_plan": "free",
                "subscription_status": "active",
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(user_doc)
            logger.info(f"New agent registered via Google OAuth: {email}")
    else:
        # Buyer login/registration
        if existing_buyer:
            user_id = existing_buyer['user_id']
            role = "buyer"
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {"name": name, "picture": picture}}
            )
        elif existing_agent:
            raise HTTPException(
                status_code=403, 
                detail="This account is registered as an agent. Please use 'Login as Agent' instead."
            )
        elif client_link:
            user_id = f"buyer_{uuid.uuid4().hex[:12]}"
            role = "buyer"
            user_doc = {
                "user_id": user_id,
                "email": email,
                "name": name,
                "picture": picture,
                "role": role,
                "is_demo": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(user_doc)
            await db.clients.update_one(
                {"client_id": client_link['client_id']},
                {"$set": {"buyer_id": user_id}}
            )
        else:
            user_id = f"buyer_{uuid.uuid4().hex[:12]}"
            role = "buyer"
            user_doc = {
                "user_id": user_id,
                "email": email,
                "name": name,
                "picture": picture,
                "role": role,
                "is_demo": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(user_doc)
    
    # Create JWT token
    token = create_jwt_token(user_id, role, False)
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
    
    return {**user, "token": token}

@api_router.get("/auth/me")
async def get_me(user: dict = Depends(get_current_user)):
    """Get current authenticated user"""
    return {k: v for k, v in user.items() if k != 'password_hash'}

# Password Reset Endpoints
class ForgotPasswordRequest(BaseModel):
    email: str
    role: str  # 'agent' or 'buyer'

class ResetPasswordRequest(BaseModel):
    token: str
    new_password: str

class SetPasswordRequest(BaseModel):
    email: str
    password: str
    role: str  # 'agent' or 'buyer'

class CheckEmailRequest(BaseModel):
    email: str
    role: str  # 'agent' or 'buyer'

@api_router.post("/auth/check-email")
async def check_email_status(request: CheckEmailRequest):
    """
    Check if an email is registered and what auth methods are available.
    Useful for frontend to show appropriate login/register options.
    """
    email = request.email.lower().strip()
    role = request.role
    
    if role not in ['agent', 'buyer']:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    user = await db.users.find_one({"email": email, "role": role}, {"_id": 0})
    
    if not user:
        return {
            "exists": False,
            "has_password": False,
            "has_google": False,
            "message": "No account found. You can create a new account."
        }
    
    has_password = bool(user.get('password_hash'))
    has_google = user.get('picture') is not None  # Google OAuth provides picture
    
    if has_password and has_google:
        return {
            "exists": True,
            "has_password": True,
            "has_google": True,
            "message": "You can login with email/password or Google."
        }
    elif has_password:
        return {
            "exists": True,
            "has_password": True,
            "has_google": False,
            "message": "Login with your email and password."
        }
    else:
        return {
            "exists": True,
            "has_password": False,
            "has_google": True,
            "message": "This account was created with Google. Login with Google, or set a password to enable email login."
        }

@api_router.post("/auth/set-password")
async def set_password_for_oauth_user(request: SetPasswordRequest, response: Response):
    """
    Set password for an OAuth-only account.
    This allows users who signed up via Google to also login with email/password.
    """
    email = request.email.lower().strip()
    role = request.role
    
    if role not in ['agent', 'buyer']:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    if len(request.password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    
    # Find the user
    user = await db.users.find_one({"email": email, "role": role}, {"_id": 0})
    
    if not user:
        raise HTTPException(status_code=404, detail="No account found with this email")
    
    if user.get('password_hash'):
        raise HTTPException(
            status_code=400, 
            detail="This account already has a password. Use 'Forgot Password' to reset it."
        )
    
    # Set the password
    password_hash = bcrypt.hashpw(request.password.encode(), bcrypt.gensalt()).decode()
    
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {"$set": {"password_hash": password_hash}}
    )
    
    logger.info(f"Password set for OAuth user: {user['user_id']}")
    
    # Generate token and set cookie for auto-login
    token = create_jwt_token(user['user_id'], role, user.get('is_demo', False))
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    return {
        "message": "Password set successfully. You can now login with email/password.",
        "user_id": user['user_id'],
        "email": email,
        "name": user['name'],
        "role": role,
        "token": token
    }

@api_router.post("/auth/forgot-password")
async def forgot_password(request: ForgotPasswordRequest, req: Request):
    """Send password reset email"""
    # Rate limit: 3 attempts per 5 minutes per IP
    rate_limit_check(req, "auth_password_reset")
    
    email = request.email.lower().strip()
    role = request.role
    
    if role not in ['agent', 'buyer']:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    # Find user with matching email and role
    user = await db.users.find_one({
        "email": email,
        "role": role,
        "password_hash": {"$exists": True}  # Only for email/password users
    })
    
    # Always return success to prevent email enumeration
    if not user:
        logger.info(f"Password reset requested for non-existent email: {email}")
        return {"message": "If an account exists with this email, you will receive a reset link."}
    
    # Generate reset token
    reset_token = secrets.token_urlsafe(32)
    reset_expiry = datetime.now(timezone.utc) + timedelta(hours=1)
    
    # Store reset token
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {"$set": {
            "password_reset_token": reset_token,
            "password_reset_expires": reset_expiry.isoformat()
        }}
    )
    
    # Get frontend URL for reset link
    frontend_url = os.environ.get('REACT_APP_BACKEND_URL', 'https://evohome.ch').replace('/api', '')
    reset_link = f"{frontend_url}/reset-password?token={reset_token}"
    
    # Send email
    subject = "Reset your Evohome password"
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <h2 style="color: #2563EB;">Reset Your Password</h2>
        <p>Hi {user.get('name', 'there')},</p>
        <p>We received a request to reset your Evohome password. Click the button below to create a new password:</p>
        <p style="text-align: center; margin: 30px 0;">
            <a href="{reset_link}" style="background-color: #2563EB; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block;">
                Reset Password
            </a>
        </p>
        <p style="color: #666; font-size: 14px;">This link will expire in 1 hour.</p>
        <p style="color: #666; font-size: 14px;">If you didn't request this, you can safely ignore this email.</p>
        <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;">
        <p style="color: #999; font-size: 12px;">Evohome - Real Estate Management</p>
    </div>
    """
    
    await send_email_async(email, subject, html_content)
    
    return {"message": "If an account exists with this email, you will receive a reset link."}

@api_router.post("/auth/reset-password")
async def reset_password(request: ResetPasswordRequest):
    """Reset password using token"""
    # Find user with valid reset token
    user = await db.users.find_one({
        "password_reset_token": request.token
    })
    
    if not user:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Check if token is expired
    expiry = user.get("password_reset_expires")
    if expiry:
        expiry_dt = datetime.fromisoformat(expiry.replace('Z', '+00:00'))
        if datetime.now(timezone.utc) > expiry_dt:
            raise HTTPException(status_code=400, detail="Reset token has expired")
    
    # Validate new password
    if len(request.new_password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    
    # Hash new password and update
    password_hash = bcrypt.hashpw(request.new_password.encode(), bcrypt.gensalt()).decode()
    
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {
            "$set": {"password_hash": password_hash},
            "$unset": {"password_reset_token": "", "password_reset_expires": ""}
        }
    )
    
    logger.info(f"Password reset successful for user: {user['user_id']}")
    
    return {"message": "Password reset successful. You can now log in with your new password."}

@api_router.post("/auth/demo/{role}")
async def demo_login(role: str, response: Response, buyer_num: int = 1):
    """Login as demo user"""
    if role not in ['buyer', 'agent']:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    if role == 'buyer':
        buyer_id = f"demo_buyer_00{buyer_num}" if buyer_num in [1, 2] else "demo_buyer_001"
        demo_user = await db.users.find_one({"user_id": buyer_id, "is_demo": True, "role": "buyer"}, {"_id": 0})
    else:
        demo_user = await db.users.find_one({"is_demo": True, "role": role}, {"_id": 0})
    
    if not demo_user:
        raise HTTPException(status_code=404, detail="Demo not initialized. Please seed demo data first.")
    
    token = create_jwt_token(demo_user['user_id'], role, True)
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    return {
        "user_id": demo_user['user_id'],
        "email": demo_user['email'],
        "name": demo_user['name'],
        "role": role,
        "is_demo": True,
        "token": token
    }


@api_router.post("/auth/refresh")
async def refresh_token(request: Request, response: Response):
    """
    Refresh access token using current valid token.
    Call this before token expires to maintain session.
    """
    token = extract_token(request)
    if not token:
        raise HTTPException(status_code=401, detail="No token provided")
    
    # Check if token is invalidated (logged out)
    if is_token_invalidated(token):
        raise HTTPException(status_code=401, detail="Session has been invalidated. Please login again.")
    
    try:
        # Verify current token (allow expired for refresh within grace period)
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM], options={"verify_exp": False})
        
        # Check if token is within refresh window (expired less than 24 hours ago)
        exp = datetime.fromtimestamp(payload['exp'], tz=timezone.utc)
        now = datetime.now(timezone.utc)
        if now > exp + timedelta(hours=24):
            raise HTTPException(status_code=401, detail="Token too old to refresh. Please login again.")
        
        user_id = payload['user_id']
        
        # Get current user data from database (ensures is_demo is current)
        user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        
        # Create new access token
        new_token = create_access_token(user_id, user['role'])
        
        response.set_cookie(
            key="session_token",
            value=new_token,
            httponly=True,
            secure=True,
            samesite="none",
            path="/",
            max_age=JWT_EXPIRY_HOURS * 60 * 60
        )
        
        return {
            "success": True,
            "token": new_token,
            "expires_in": JWT_EXPIRY_HOURS * 3600
        }
        
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


@api_router.get("/auth/session", response_model=AuthSessionResponse)
async def check_session(request: Request):
    """
    Check if current session is valid.
    Returns user info if authenticated, 401 if not.
    Use this for frontend auth state sync.
    """
    token = extract_token(request)
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    if is_token_invalidated(token):
        raise HTTPException(status_code=401, detail="Session invalidated")
    
    try:
        # Use local verify function for backward compatibility
        payload = verify_jwt_token(token)
        user = await db.users.find_one({"user_id": payload['user_id']}, {"_id": 0})
        
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        
        return {
            "authenticated": True,
            "user": {
                "user_id": user['user_id'],
                "email": user['email'],
                "name": user['name'],
                "role": user['role'],
                "is_demo": user.get('is_demo', False)
            }
        }
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid session")


@api_router.post("/auth/logout", response_model=AuthLogoutResponse)
async def logout(request: Request, response: Response):
    """
    Logout current user.
    Invalidates token and clears cookie.
    """
    token = extract_token(request)
    
    if token:
        # Invalidate the token
        invalidate_token(token)
    
    # Clear the cookie
    response.delete_cookie(
        key="session_token",
        path="/",
        secure=True,
        httponly=True,
        samesite="none"
    )
    
    return {"success": True, "message": "Logged out successfully"}


# ==================== PROJECT ENDPOINTS ====================

@api_router.get("/projects", response_model=List[Project])
async def get_projects(user: dict = Depends(get_current_user)):
    """
    Get projects for current user.
    - Agents: see all their projects (filtered by is_demo to maintain data isolation)
    - Buyers: see only the project(s) they are associated with via client record
    """
    is_demo = user.get('is_demo', False)
    
    if user['role'] == 'agent':
        # Agent sees their projects matching is_demo context
        query = {"agent_id": user['user_id'], "is_demo": is_demo}
    elif user['role'] == 'buyer':
        # Buyer sees only projects they are associated with via client record
        client = await db.clients.find_one(
            {"buyer_id": user['user_id'], "is_demo": is_demo},
            {"_id": 0, "project_id": 1}
        )
        if not client or not client.get('project_id'):
            return []  # Buyer has no associated project
        query = {"project_id": client['project_id'], "is_demo": is_demo}
    else:
        return []  # Unknown role gets nothing
    
    projects = await db.projects.find(query, {"_id": 0}).to_list(100)
    
    # Add client count and unit count to each project
    for project in projects:
        client_count = await db.clients.count_documents({"project_id": project['project_id'], "is_demo": is_demo})
        unit_count = await db.project_units.count_documents({"project_id": project['project_id']})
        project['client_count'] = client_count
        project['unit_count'] = unit_count
    
    return projects

@api_router.post("/projects", response_model=Project)
async def create_project(data: ProjectCreate, user: dict = Depends(get_current_agent)):
    """Create a new project (agent only) - enforces plan-based unit limits"""
    is_demo = user.get('is_demo', False)
    
    # Unit gating - check subscription limits (skip for demo users)
    # Note: Creating a project adds 1 unit by default (more units can be added later)
    if not is_demo:
        subscription_data = await get_agent_subscription_data(user)
        if not subscription_data['can_create_unit']:
            plan_name = subscription_data['plan_name']
            limit = subscription_data['unit_limit']
            current = subscription_data['unit_usage']
            raise HTTPException(
                status_code=403, 
                detail=f"Unit limit reached. Your {plan_name} plan allows {limit} units ({current} used). Please upgrade to add more units."
            )
    
    project_id = f"proj_{uuid.uuid4().hex[:12]}"
    
    # NOTE: is_demo NOT stored on project - ownership scoped by agent_id
    project_doc = {
        "project_id": project_id,
        "agent_id": user['user_id'],
        "name": data.name,
        "address": data.address,
        "description": data.description,
        "total_units": data.total_units,
        "construction_start": data.construction_start,
        "estimated_completion": data.estimated_completion,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.projects.insert_one(project_doc)
    project = await db.projects.find_one({"project_id": project_id}, {"_id": 0})
    # Count linked clients and units (scoped by project_id)
    client_count = await db.clients.count_documents({"project_id": project_id})
    unit_count = await db.project_units.count_documents({"project_id": project_id})
    project['client_count'] = client_count
    project['unit_count'] = unit_count
    return project

@api_router.put("/projects/{project_id}", response_model=Project)
async def update_project(project_id: str, data: ProjectUpdate, user: dict = Depends(get_current_agent)):
    """Update a project (agent only)"""
    # Ownership check by agent_id only
    query = {"project_id": project_id, "agent_id": user['user_id']}
    
    project = await db.projects.find_one(query, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    update_data = {k: v for k, v in data.model_dump().items() if v is not None}
    if update_data:
        await db.projects.update_one(query, {"$set": update_data})
    
    updated_project = await db.projects.find_one(query, {"_id": 0})
    client_count = await db.clients.count_documents({"project_id": project_id})
    unit_count = await db.project_units.count_documents({"project_id": project_id})
    updated_project['client_count'] = client_count
    updated_project['unit_count'] = unit_count
    return updated_project

@api_router.delete("/projects/{project_id}")
async def delete_project(project_id: str, user: dict = Depends(get_current_agent)):
    """Delete a project (agent only)"""
    is_demo = user.get('is_demo', False)
    query = {"project_id": project_id, "agent_id": user['user_id']}
    
    project = await db.projects.find_one(query, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Check if project has linked clients
    client_count = await db.clients.count_documents({"project_id": project_id})
    if client_count > 0:
        raise HTTPException(status_code=400, detail=f"Cannot delete project with {client_count} linked client(s). Remove clients first.")
    
    await db.projects.delete_one(query)
    return {"message": "Project deleted"}

# ==================== PROJECT UNITS ENDPOINTS ====================

@api_router.get("/projects/{project_id}/units")
async def get_project_units(project_id: str, user: dict = Depends(get_current_agent)):
    """Get all units for a project with assignment status"""
    is_demo = user.get('is_demo', False)
    query = {"project_id": project_id, "agent_id": user['user_id']}
    
    project = await db.projects.find_one(query, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    units = await db.project_units.find(
        {"project_id": project_id},
        {"_id": 0}
    ).to_list(500)
    
    # Enrich with client assignment info
    for unit in units:
        # Check if any client is assigned to this unit
        assigned_client = await db.clients.find_one(
            {"unit_id": unit['unit_id']},
            {"_id": 0, "client_id": 1, "name": 1}
        )
        if assigned_client:
            unit['assigned_client_id'] = assigned_client['client_id']
            unit['assigned_client_name'] = assigned_client['name']
            unit['is_available'] = False
        else:
            unit['assigned_client_id'] = None
            unit['assigned_client_name'] = None
            unit['is_available'] = True
        
        # Legacy client_id field for backward compatibility
        if unit.get('client_id'):
            client = await db.clients.find_one({"client_id": unit['client_id']}, {"_id": 0, "name": 1})
            unit['client_name'] = client['name'] if client else None
    
    return units

@api_router.post("/projects/{project_id}/units")
async def create_project_unit(project_id: str, data: dict, user: dict = Depends(get_current_agent)):
    """Add a unit to a project - enforces plan-based unit limits"""
    is_demo = user.get('is_demo', False)
    query = {"project_id": project_id, "agent_id": user['user_id']}
    
    project = await db.projects.find_one(query, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Unit gating - check subscription limits (skip for demo users)
    if not is_demo:
        subscription_data = await get_agent_subscription_data(user)
        if not subscription_data['can_create_unit']:
            plan_name = subscription_data['plan_name']
            limit = subscription_data['unit_limit']
            current = subscription_data['unit_usage']
            raise HTTPException(
                status_code=403, 
                detail=f"Unit limit reached. Your {plan_name} plan allows {limit} units ({current} used). Please upgrade to add more units."
            )
    
    unit_reference = data.get('unit_reference', '').strip()
    if not unit_reference:
        raise HTTPException(status_code=400, detail="Unit reference is required")
    
    # Check for duplicate
    existing = await db.project_units.find_one({
        "project_id": project_id,
        "unit_reference": unit_reference,
        "is_demo": is_demo
    })
    if existing:
        raise HTTPException(status_code=400, detail="Unit reference already exists in this project")
    
    unit_id = f"unit_{uuid.uuid4().hex[:12]}"
    unit_doc = {
        "unit_id": unit_id,
        "project_id": project_id,
        "unit_reference": unit_reference,
        "client_id": None,
        "is_demo": is_demo,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.project_units.insert_one(unit_doc)
    return await db.project_units.find_one({"unit_id": unit_id}, {"_id": 0})

@api_router.delete("/projects/{project_id}/units/{unit_id}")
async def delete_project_unit(project_id: str, unit_id: str, user: dict = Depends(get_current_agent)):
    """Remove a unit from a project"""
    is_demo = user.get('is_demo', False)
    
    # Verify project ownership
    project = await db.projects.find_one({
        "project_id": project_id,
        "agent_id": user['user_id'],
        "is_demo": is_demo
    }, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    unit = await db.project_units.find_one({
        "unit_id": unit_id,
        "project_id": project_id,
        "is_demo": is_demo
    }, {"_id": 0})
    if not unit:
        raise HTTPException(status_code=404, detail="Unit not found")
    
    if unit.get('client_id'):
        raise HTTPException(status_code=400, detail="Cannot delete unit that is assigned to a client")
    
    await db.project_units.delete_one({"unit_id": unit_id})
    return {"message": "Unit deleted"}

# ==================== TEAM MEMBER ENDPOINTS ====================

@api_router.get("/projects/{project_id}/team")
async def get_team_members(project_id: str, user: dict = Depends(get_current_user)):
    """Get team members for a project - accessible by both agent and linked buyers"""
    is_demo = user.get('is_demo', False)
    
    if user['role'] == 'agent':
        # Agent must own the project
        project = await db.projects.find_one(
            {"project_id": project_id, "agent_id": user['user_id']},
            {"_id": 0}
        )
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
    else:
        # Buyer must be linked to a client in this project
        client = await db.clients.find_one(
            {"buyer_id": user['user_id'], "project_id": project_id},
            {"_id": 0}
        )
        if not client:
            raise HTTPException(status_code=403, detail="Not authorized to view this project's team")
    
    members = await db.team_members.find(
        {"project_id": project_id},
        {"_id": 0}
    ).sort("name", 1).to_list(100)
    
    return members

@api_router.post("/projects/{project_id}/team")
async def create_team_member(project_id: str, data: TeamMemberCreate, user: dict = Depends(get_current_agent)):
    """Add a team member to a project (agent only)"""
    is_demo = user.get('is_demo', False)
    
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    member_id = f"member_{uuid.uuid4().hex[:12]}"
    member_doc = {
        "member_id": member_id,
        "project_id": project_id,
        "agent_id": user['user_id'],
        "company_name": data.company_name,
        "contact_name": data.contact_name,
        "role": data.role,
        "email": data.email,
        "phone": data.phone,
        "website": data.website,
        "address": data.address,
        "notes": data.notes,
        "is_demo": is_demo,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.team_members.insert_one(member_doc)
    return await db.team_members.find_one({"member_id": member_id}, {"_id": 0})

@api_router.put("/projects/{project_id}/team/{member_id}")
async def update_team_member(project_id: str, member_id: str, data: TeamMemberUpdate, user: dict = Depends(get_current_agent)):
    """Update a team member (agent only)"""
    is_demo = user.get('is_demo', False)
    
    member = await db.team_members.find_one({
        "member_id": member_id,
        "project_id": project_id,
        "agent_id": user['user_id'],
        "is_demo": is_demo
    }, {"_id": 0})
    
    if not member:
        raise HTTPException(status_code=404, detail="Team member not found")
    
    update_data = {k: v for k, v in data.model_dump().items() if v is not None}
    if update_data:
        await db.team_members.update_one(
            {"member_id": member_id},
            {"$set": update_data}
        )
    
    return await db.team_members.find_one({"member_id": member_id}, {"_id": 0})

@api_router.delete("/projects/{project_id}/team/{member_id}")
async def delete_team_member(project_id: str, member_id: str, user: dict = Depends(get_current_agent)):
    """Delete a team member (agent only)"""
    is_demo = user.get('is_demo', False)
    
    member = await db.team_members.find_one({
        "member_id": member_id,
        "project_id": project_id,
        "agent_id": user['user_id'],
        "is_demo": is_demo
    }, {"_id": 0})
    
    if not member:
        raise HTTPException(status_code=404, detail="Team member not found")
    
    await db.team_members.delete_one({"member_id": member_id})
    return {"message": "Team member deleted"}

# ==================== GLOBAL TEAM DIRECTORY & AI EXTRACTION ====================

@api_router.get("/team/directory")
async def get_team_directory(
    search: Optional[str] = None,
    project_id: Optional[str] = None,
    limit: int = 50,
    user: dict = Depends(get_current_agent)
):
    """
    Get all team members across all projects for the agent.
    This serves as the global supplier/contact directory.
    """
    is_demo = user.get('is_demo', False)
    query = {"agent_id": user['user_id']}
    
    if project_id:
        query["project_id"] = project_id
    
    members = await db.team_members.find(query, {"_id": 0}).to_list(500)
    
    # Deduplicate by company_name + contact_name
    seen = set()
    unique_members = []
    for member in members:
        key = (member.get('company_name', '').lower(), member.get('contact_name', '').lower())
        if key not in seen:
            seen.add(key)
            unique_members.append(member)
    
    # Apply search filter
    if search:
        search_lower = search.lower()
        unique_members = [
            m for m in unique_members 
            if search_lower in (m.get('company_name', '') or '').lower()
            or search_lower in (m.get('contact_name', '') or '').lower()
            or search_lower in (m.get('role', '') or '').lower()
        ]
    
    # Sort by most recently used (created_at descending)
    unique_members.sort(key=lambda x: x.get('created_at', ''), reverse=True)
    
    return unique_members[:limit]

@api_router.post("/team/extract-contacts")
async def extract_contacts_from_document(
    file: UploadFile = File(...),
    user: dict = Depends(get_current_agent)
):
    """
    Extract contact information from uploaded document using AI.
    Supports PDF, Word, Excel, and image files.
    Returns structured list of contacts for review before import.
    """
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=503, detail="AI extraction not available - OpenAI API key not configured")
    
    # Validate file type
    allowed_extensions = ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.png', '.jpg', '.jpeg', '.webp']
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
        )
    
    # Save uploaded file temporarily
    temp_dir = "/tmp/contact_extraction"
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, f"{uuid.uuid4().hex}_{file.filename}")
    
    try:
        content = await file.read()
        with open(temp_path, "wb") as f:
            f.write(content)
        
        # Extract text/content based on file type
        extracted_text = ""
        images_base64 = []
        
        if file_ext == '.pdf':
            # Convert PDF to images for vision API
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(temp_path)
                for page_num in range(min(doc.page_count, 5)):  # Limit to first 5 pages
                    page = doc[page_num]
                    # Render page to image
                    mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    images_base64.append(base64.b64encode(img_data).decode('utf-8'))
                doc.close()
            except Exception as e:
                logger.error(f"PDF processing error: {e}")
                raise HTTPException(status_code=500, detail="Failed to process PDF")
        
        elif file_ext in ['.png', '.jpg', '.jpeg', '.webp']:
            with open(temp_path, "rb") as f:
                images_base64.append(base64.b64encode(f.read()).decode('utf-8'))
        
        elif file_ext in ['.doc', '.docx', '.xls', '.xlsx']:
            # For Office documents, we'll use a simpler text extraction
            # In production, you'd use python-docx or openpyxl
            try:
                if file_ext == '.docx':
                    from docx import Document
                    doc = Document(temp_path)
                    extracted_text = "\n".join([para.text for para in doc.paragraphs])
                elif file_ext in ['.xlsx', '.xls']:
                    import openpyxl
                    wb = openpyxl.load_workbook(temp_path, data_only=True)
                    texts = []
                    for sheet in wb.worksheets[:3]:  # First 3 sheets
                        for row in sheet.iter_rows(max_row=100, values_only=True):
                            row_text = " | ".join([str(cell) for cell in row if cell])
                            if row_text.strip():
                                texts.append(row_text)
                    extracted_text = "\n".join(texts)
            except ImportError:
                # Fallback: treat as binary and let AI try to extract
                logger.warning(f"Office document libraries not available for {file_ext}")
                raise HTTPException(status_code=400, detail="Office document processing not available")
        
        # Use OpenAI to extract contacts
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        extraction_prompt = """Analyze this document and extract all contact information for companies, suppliers, or individuals.

For each contact found, extract:
- company_name: The company or business name
- contact_name: The person's full name (if available)
- role: Their job title or function (e.g., "Plumber", "Architect", "Sales Manager")
- email: Email address
- phone: Phone number(s)
- website: Website URL
- address: Full address

Rules:
1. Only extract real contacts that appear to be suppliers, contractors, or business partners
2. Ignore generic text that isn't a contact
3. Deduplicate contacts - if the same company appears multiple times, merge into one complete record
4. Accept partial data - not all fields need to be present
5. For role, try to infer from context if not explicitly stated

Return a JSON array of contacts. Example:
[
  {
    "company_name": "SaniTech SA",
    "contact_name": "Pierre Dupont",
    "role": "Plumber",
    "email": "contact@sanitech.ch",
    "phone": "+41 79 123 45 67",
    "website": "https://sanitech.ch",
    "address": "Rue du Lac 15, 1200 Genève"
  }
]

If no contacts are found, return an empty array: []"""

        messages = [{"role": "system", "content": extraction_prompt}]
        
        if images_base64:
            # Use vision model for images/PDFs
            content_parts = [{"type": "text", "text": "Extract all contact information from this document:"}]
            for img_b64 in images_base64[:5]:  # Limit images
                content_parts.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{img_b64}", "detail": "high"}
                })
            messages.append({"role": "user", "content": content_parts})
        else:
            # Text-only extraction
            messages.append({
                "role": "user", 
                "content": f"Extract all contact information from this document:\n\n{extracted_text[:10000]}"
            })
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=4000,
            temperature=0.1
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # Parse JSON from response
        import json
        try:
            # Try to extract JSON from response (it might be wrapped in markdown)
            if "```json" in response_text:
                json_str = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_str = response_text.split("```")[1].split("```")[0].strip()
            else:
                json_str = response_text
            
            contacts = json.loads(json_str)
            
            if not isinstance(contacts, list):
                contacts = [contacts] if contacts else []
            
            # Validate and clean contacts
            valid_contacts = []
            for contact in contacts:
                if isinstance(contact, dict) and (contact.get('company_name') or contact.get('contact_name')):
                    valid_contacts.append({
                        "company_name": contact.get('company_name', ''),
                        "contact_name": contact.get('contact_name', ''),
                        "role": contact.get('role', ''),
                        "email": contact.get('email', ''),
                        "phone": contact.get('phone', ''),
                        "website": contact.get('website', ''),
                        "address": contact.get('address', ''),
                        "confidence": 0.85  # Default confidence
                    })
            
            return {
                "contacts": valid_contacts,
                "source_filename": file.filename,
                "extraction_count": len(valid_contacts)
            }
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response: {e}\nResponse: {response_text[:500]}")
            return {
                "contacts": [],
                "source_filename": file.filename,
                "extraction_count": 0,
                "error": "Could not extract contacts from document"
            }
    
    finally:
        # Cleanup temp file
        if os.path.exists(temp_path):
            os.remove(temp_path)

class BulkContactsRequest(BaseModel):
    contacts: List[dict]

@api_router.post("/projects/{project_id}/team/bulk")
async def create_team_members_bulk(
    project_id: str, 
    request: BulkContactsRequest,
    user: dict = Depends(get_current_agent)
):
    """
    Bulk import multiple team members to a project.
    Used after AI extraction review.
    """
    is_demo = user.get('is_demo', False)
    
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    created_members = []
    skipped_count = 0
    
    for contact in request.contacts:
        # Skip if no company name or contact name
        if not contact.get('company_name') and not contact.get('contact_name'):
            skipped_count += 1
            continue
        
        # Check for duplicates in this project
        existing = await db.team_members.find_one({
            "project_id": project_id,
            "agent_id": user['user_id'],
            "company_name": contact.get('company_name', ''),
            "contact_name": contact.get('contact_name', ''),
            "is_demo": is_demo
        })
        if existing:
            skipped_count += 1
            continue
        
        member_id = f"member_{uuid.uuid4().hex[:12]}"
        member_doc = {
            "member_id": member_id,
            "project_id": project_id,
            "agent_id": user['user_id'],
            "company_name": contact.get('company_name', ''),
            "contact_name": contact.get('contact_name', ''),
            "role": contact.get('role', 'Supplier'),
            "email": contact.get('email'),
            "phone": contact.get('phone'),
            "website": contact.get('website'),
            "address": contact.get('address'),
            "notes": contact.get('notes'),
            "is_demo": is_demo,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.team_members.insert_one(member_doc)
        created_members.append(member_doc)
    
    return {
        "created": len(created_members),
        "skipped": skipped_count,
        "members": [{k: v for k, v in m.items() if k != '_id'} for m in created_members]
    }

# ==================== CLIENT ENDPOINTS ====================

@api_router.get("/clients", response_model=List[Client])
async def get_clients(project_id: Optional[str] = None, user: dict = Depends(get_current_agent)):
    """Get all clients for agent, optionally filtered by project"""
    query = {"agent_id": user['user_id'], "is_demo": user.get('is_demo', False)}
    if project_id:
        query["project_id"] = project_id
    clients = await db.clients.find(query, {"_id": 0}).to_list(100)
    return clients

@api_router.get("/clients/{client_id}", response_model=Client)
async def get_client(client_id: str, user: dict = Depends(get_current_agent)):
    """Get single client"""
    query = {"client_id": client_id, "agent_id": user['user_id'], "is_demo": user.get('is_demo', False)}
    client = await db.clients.find_one(query, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return client

@api_router.get("/clients/{client_id}/preview")
async def get_client_preview_data(client_id: str, user: dict = Depends(get_current_agent)):
    """
    Get client data for "View as Client" preview.
    Returns client info, their activities, documents, and project details.
    Agent uses this to see what the client sees.
    """
    is_demo = user.get('is_demo', False)
    
    client = await db.clients.find_one(
        {"client_id": client_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    project = await db.projects.find_one(
        {"project_id": client['project_id']},
        {"_id": 0}
    )
    
    # Get activities for this client
    recipient_records = await db.activity_recipients.find(
        {"client_id": client_id},
        {"_id": 0, "activity_id": 1}
    ).to_list(1000)
    
    activity_ids = [r['activity_id'] for r in recipient_records]
    
    activities = []
    if activity_ids:
        activities_raw = await db.activities.find(
            {"activity_id": {"$in": activity_ids}},
            {"_id": 0}
        ).sort("created_at", -1).limit(20).to_list(20)
        
        for act in activities_raw:
            activities.append(await enrich_activity(act))
    
    # Get documents for this client
    documents = await db.documents.find(
        {"client_id": client_id, "is_demo": is_demo, "status": {"$ne": "Draft"}},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    
    # Get construction stages for the project
    stages = await db.project_stages.find(
        {"project_id": client['project_id']},
        {"_id": 0}
    ).sort("order", 1).to_list(50)
    
    # Get team members for the project
    team = await db.team_members.find(
        {"project_id": client['project_id']},
        {"_id": 0}
    ).sort("name", 1).to_list(100)
    
    return {
        "client": client,
        "project": project,
        "activities": activities,
        "documents": documents,
        "stages": stages,
        "team": team,
        "is_preview": True
    }

@api_router.post("/clients", response_model=Client)
async def create_client(data: ClientCreate, user: dict = Depends(get_current_agent)):
    """Create a new client linked to a project"""
    is_demo = user.get('is_demo', False)
    
    project = await db.projects.find_one(
        {"project_id": data.project_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Resolve unit_id to unit_reference
    unit_reference = "General"
    unit_id = None
    if data.unit_id:
        unit = await db.project_units.find_one({
            "unit_id": data.unit_id,
            "project_id": data.project_id,
            "is_demo": is_demo
        }, {"_id": 0})
        if unit:
            unit_id = data.unit_id
            unit_reference = unit.get('unit_reference', unit.get('name', 'Unit'))
        else:
            raise HTTPException(status_code=400, detail="Invalid unit for this project")
        
        # Check if unit is already assigned to another client
        existing_client = await db.clients.find_one({
            "unit_id": data.unit_id,
            "is_demo": is_demo
        }, {"_id": 0, "name": 1})
        if existing_client:
            raise HTTPException(
                status_code=409, 
                detail=f"Unit already assigned to client: {existing_client.get('name', 'Unknown')}"
            )
    
    client_id = f"client_{uuid.uuid4().hex[:12]}"
    buyer = await db.users.find_one({"email": data.email, "role": "buyer"}, {"_id": 0})
    
    client_doc = {
        "client_id": client_id,
        "agent_id": user['user_id'],
        "buyer_id": buyer['user_id'] if buyer else None,
        "name": data.name,
        "email": data.email,
        "phone": data.phone,
        "project_id": data.project_id,
        "unit_id": unit_id,
        "unit_reference": unit_reference,
        "is_demo": is_demo,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.clients.insert_one(client_doc)
    return await db.clients.find_one({"client_id": client_id}, {"_id": 0})

@api_router.put("/clients/{client_id}", response_model=Client)
async def update_client(client_id: str, data: ClientUpdate, user: dict = Depends(get_current_agent)):
    """Update a client"""
    is_demo = user.get('is_demo', False)
    query = {"client_id": client_id, "agent_id": user['user_id']}
    client = await db.clients.find_one(query, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    update_data = {}
    
    # Handle basic field updates
    if data.name is not None:
        update_data['name'] = data.name
    if data.email is not None:
        update_data['email'] = data.email
    if data.phone is not None:
        update_data['phone'] = data.phone
    
    # Handle project change
    if data.project_id is not None and data.project_id != client.get('project_id'):
        project = await db.projects.find_one({
            "project_id": data.project_id,
            "agent_id": user['user_id'],
            "is_demo": is_demo
        })
        if not project:
            raise HTTPException(status_code=400, detail="Invalid project")
        update_data['project_id'] = data.project_id
        # Reset unit when project changes
        update_data['unit_id'] = None
        update_data['unit_reference'] = "General"
    
    # Handle unit change
    if data.unit_id is not None:
        target_project_id = data.project_id or client.get('project_id')
        if data.unit_id == "":
            update_data['unit_id'] = None
            update_data['unit_reference'] = "General"
        else:
            unit = await db.project_units.find_one({
                "unit_id": data.unit_id,
                "project_id": target_project_id,
                "is_demo": is_demo
            }, {"_id": 0})
            if not unit:
                raise HTTPException(status_code=400, detail="Invalid unit for this project")
            
            # Check if unit is already assigned to another client
            existing_client = await db.clients.find_one({
                "unit_id": data.unit_id,
                "client_id": {"$ne": client_id},
                "is_demo": is_demo
            }, {"_id": 0, "name": 1, "client_id": 1})
            
            if existing_client:
                if data.force_unit_reassign:
                    # Remove unit from the other client
                    await db.clients.update_one(
                        {"client_id": existing_client['client_id']},
                        {"$set": {"unit_id": None, "unit_reference": "General"}}
                    )
                    logger.info(f"Reassigned unit {data.unit_id} from {existing_client['client_id']} to {client_id}")
                else:
                    raise HTTPException(
                        status_code=409, 
                        detail=f"Unit already assigned to client: {existing_client.get('name', 'Unknown')}"
                    )
            
            update_data['unit_id'] = data.unit_id
            update_data['unit_reference'] = unit.get('unit_reference', unit.get('name', 'Unit'))
    
    if update_data:
        await db.clients.update_one(query, {"$set": update_data})
    
    return await db.clients.find_one(query, {"_id": 0})

@api_router.delete("/clients/{client_id}")
async def delete_client(client_id: str, user: dict = Depends(get_current_agent)):
    """Delete a client"""
    query = {"client_id": client_id, "agent_id": user['user_id'], "is_demo": user.get('is_demo', False)}
    result = await db.clients.delete_one(query)
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Client not found")
    return {"message": "Client deleted"}

# ==================== AI PDF EXTRACTION ====================

OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')

async def extract_document_from_pdf(pdf_path: str, filename: str) -> dict:
    """Extract document data from PDF using AI (OpenAI GPT-4o with vision)"""
    if not OPENAI_API_KEY:
        logger.warning("OPENAI_API_KEY not set, returning fallback extraction")
        return {
            "title": filename.replace('.pdf', '').replace('_', ' ').title(),
            "amount": None,
            "items": [],
            "supplier_name": None,
            "description": "AI extraction unavailable - please enter details manually",
            "confidence": "low",
            "extraction_failed": True
        }
    
    try:
        # Initialize OpenAI client
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        system_prompt = """You are a document extraction assistant for a real estate construction/renovation project management system.

Your task is to extract structured data from quotes, invoices, and other construction-related documents.

EXTRACTION RULES:

1. TOTAL AMOUNT (CRITICAL):
   - Find the FINAL TOTAL at the END of the document
   - This is usually labeled: "Total", "Grand Total", "Total TTC", "Montant Total", "Gesamtbetrag", "Totale"
   - If there are multiple totals (subtotal, tax, grand total), use the GRAND TOTAL
   - Look at the BOTTOM of the document - the final number is what matters
   - Return as a number WITHOUT currency symbols

2. SUMMARY/DESCRIPTION:
   - Write a clean, professional single paragraph (2-4 sentences)
   - Summarize WHAT work/items are covered in this document
   - Include key details: type of work, scope, location if mentioned
   - This will be shown to the client, so be clear and professional

3. LINE ITEMS:
   - Extract individual items with description, quantity, unit price, and line total
   - If items are grouped, extract the group totals
   - If no clear line items, create ONE line item with the document description and total amount

4. SUPPLIER INFO:
   - Extract company name from header/letterhead
   - This is the vendor/contractor providing the quote/invoice

Return ONLY a JSON object with this structure:
{
  "title": "short descriptive title (e.g., 'Bathroom Renovation Quote')",
  "amount": final_total_number_or_null,
  "items": [{"description": "string", "quantity": 1, "unit_price": number, "total": number}],
  "supplier_name": "company name or null",
  "description": "professional single-paragraph summary of the document scope",
  "confidence": "high/medium/low"
}"""
        
        # Convert PDF to images using PyMuPDF
        import base64
        doc = fitz.open(pdf_path)
        image_contents = []
        
        # Convert first 3 pages max (to avoid token limits)
        for page_num in range(min(3, len(doc))):
            page = doc[page_num]
            # Render page to image at 150 DPI for good quality
            mat = fitz.Matrix(150/72, 150/72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img_base64 = base64.b64encode(img_bytes).decode('utf-8')
            image_contents.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/png;base64,{img_base64}"
                }
            })
        doc.close()
        
        if not image_contents:
            raise ValueError("Could not extract any images from PDF")
        
        # Build message content with text prompt and images
        message_content = [
            {
                "type": "text",
                "text": "Extract document information from this PDF (shown as images). Focus on finding the FINAL TOTAL AMOUNT at the end of the document. Write a professional summary. Return only valid JSON."
            }
        ] + image_contents
        
        # Use GPT-4o with vision for image analysis
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": message_content}
            ],
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content
        
        logger.info(f"AI extraction response: {response_text[:200]}...")
        
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            extracted = json.loads(json_match.group())
        else:
            raise ValueError("No JSON found in response")
        
        result = {
            "title": extracted.get("title") or filename.replace('.pdf', '').replace('_', ' ').title(),
            "amount": extracted.get("amount"),
            "items": extracted.get("items", []),
            "supplier_name": extracted.get("supplier_name"),
            "description": extracted.get("description", ""),
            "confidence": extracted.get("confidence", "medium"),
            "extraction_failed": False
        }
        
        # Validate and clean line items
        valid_items = []
        for item in result["items"]:
            if isinstance(item, dict) and item.get("description"):
                valid_items.append({
                    "description": str(item.get("description", "")),
                    "quantity": int(item.get("quantity", 1)),
                    "unit_price": float(item.get("unit_price", 0)),
                    "total": float(item.get("total", item.get("unit_price", 0)))
                })
        
        # If we have an amount but no line items, create a single line item
        if result["amount"] and not valid_items:
            valid_items.append({
                "description": result["description"] or result["title"],
                "quantity": 1,
                "unit_price": float(result["amount"]),
                "total": float(result["amount"])
            })
        
        result["items"] = valid_items
        
        logger.info(f"PDF extraction successful: {result.get('title')}, amount: {result.get('amount')}, confidence: {result.get('confidence')}")
        return result
        
    except Exception as e:
        # Capture AI extraction errors for monitoring
        capture_ai_error(e, operation="pdf_extraction", document_id=filename)
        logger.error(f"PDF extraction failed: {str(e)}")
        # Don't expose internal paths in error messages
        error_msg = str(e)
        if '/tmp/' in error_msg or '/app/' in error_msg:
            error_msg = "Could not process document. Please ensure it is a valid PDF."
        return {
            "title": filename.replace('.pdf', '').replace('_', ' ').title() if filename else "Document",
            "amount": None,
            "items": [],
            "supplier_name": None,
            "description": f"Extraction failed: {error_msg[:80]}" if len(error_msg) <= 80 else "Extraction failed. Please enter details manually.",
            "confidence": "low",
            "extraction_failed": True
        }

# ==================== UNIFIED DOCUMENT ENDPOINTS ====================
# Single source of truth: db.documents collection

@api_router.post("/documents/upload")
async def upload_document(
    request: Request,
    client_id: str = Form(...),
    file: UploadFile = File(...),
    doc_type: str = "quote",
    user: dict = Depends(get_current_agent)
):
    """
    Upload a PDF for extraction ONLY - does NOT create a document.
    Returns extracted data for user preview/editing.
    Call POST /documents/create to actually create the document after user confirmation.
    
    This follows the draft-first architecture:
    input → extraction → preview → edit → confirm → final object
    """
    # Rate limit: 10 AI extractions per minute per user
    rate_limit_check(request, "ai_extraction")
    
    is_demo = user.get('is_demo', False)
    
    # Validate document type
    if doc_type not in ['quote', 'invoice']:
        doc_type = 'quote'
    
    # Verify client exists
    client = await db.clients.find_one(
        {"client_id": client_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Validate file
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    # File size limit: 10MB max for PDF uploads
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413, 
            detail=f"File too large. Maximum size is 10MB (received {len(content) / 1024 / 1024:.1f}MB)"
        )
    
    # Save file temporarily for extraction
    file_id = uuid.uuid4().hex[:12]
    filename = f"{file_id}_{file.filename}"
    file_path = UPLOAD_DIR / filename
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Run AI extraction
    extraction = await extract_document_from_pdf(str(file_path), file.filename)
    
    # Calculate total from line items if not extracted
    items_total = sum(item.get('total', 0) for item in extraction.get('items', []))
    amount = extraction.get('amount') or items_total or 0
    
    # Return extraction preview (NO database write!)
    # The frontend will display this for user review, then call /documents/create
    preview_data = {
        "preview_id": file_id,  # Used to reference the uploaded file later
        "type": doc_type,
        "client_id": client_id,
        "project_id": client['project_id'],
        "unit_reference": client['unit_reference'],
        "title": extraction.get('title', 'Untitled Document'),
        "amount": amount,
        "items": extraction.get('items', []),
        "supplier_name": extraction.get('supplier_name'),
        "summary": extraction.get('description', ''),
        "pdf_filename": file.filename,
        "pdf_path": str(file_path),
        "ai_extraction_confidence": extraction.get('confidence', 'low'),
        "extraction_warning": extraction.get('extraction_failed', False) or extraction.get('amount') is None,
        "is_preview": True  # Flag to indicate this is preview data, not a saved document
    }
    
    return preview_data


@api_router.post("/documents/create")
async def create_document_from_preview(
    request: Request,
    user: dict = Depends(get_current_agent)
):
    """
    Create a document from previewed extraction data.
    This is called AFTER the user reviews and confirms the extraction.
    
    Expected body:
    {
        "preview_id": "...",  # From the upload preview
        "type": "quote" | "invoice",
        "client_id": "...",
        "title": "...",
        "amount": 1234.56,
        "items": [...],
        "supplier_name": "...",
        "summary": "...",
        "notes": "...",
        "due_date": "2026-04-01",
        "pdf_path": "..."  # Path to the uploaded PDF
    }
    """
    body = await request.json()
    is_demo = user.get('is_demo', False)
    
    doc_type = body.get('type', 'quote')
    if doc_type not in ['quote', 'invoice']:
        doc_type = 'quote'
    
    client_id = body.get('client_id')
    if not client_id:
        raise HTTPException(status_code=400, detail="client_id is required")
    
    # Verify client exists
    client = await db.clients.find_one(
        {"client_id": client_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Generate document number
    if doc_type == 'invoice':
        count = await db.documents.count_documents({"agent_id": user['user_id'], "type": "invoice"})
        doc_number = f"INV-{datetime.now().year}-{str(count + 1).zfill(4)}"
    else:
        count = await db.documents.count_documents({"agent_id": user['user_id'], "type": "quote"})
        doc_number = f"QT-{datetime.now().year}-{str(count + 1).zfill(4)}"
    
    doc_id = f"doc_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    # Parse due date
    due_date = None
    if body.get('due_date'):
        try:
            due_date = datetime.fromisoformat(body['due_date'].replace('Z', '+00:00')).isoformat()
        except (ValueError, AttributeError):
            due_date = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
    elif doc_type == 'invoice':
        due_date = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
    
    doc = {
        "document_id": doc_id,
        "document_number": doc_number,
        "type": doc_type,  # Primary field
        "document_type": doc_type,  # Alias for compatibility
        "status": "Draft",
        "agent_id": user['user_id'],  # Required - ownership scoping
        "client_id": client_id,
        "buyer_id": client.get('buyer_id'),
        "project_id": client['project_id'],
        "unit_reference": client['unit_reference'],
        "title": body.get('title', 'Untitled Document'),
        "amount": float(body.get('amount', 0)),  # Primary field
        "total_amount": float(body.get('amount', 0)),  # Alias for compatibility
        "items": body.get('items', []),
        "currency": "CHF",
        "supplier_name": body.get('supplier_name'),
        "notes": body.get('notes'),
        "summary": body.get('summary', ''),
        "hero_image_url": None,
        "hero_image_path": None,
        "change_request_comment": None,
        "pdf_filename": body.get('pdf_filename'),
        "pdf_path": body.get('pdf_path'),
        "ai_extraction_confidence": body.get('ai_extraction_confidence', 'low'),
        "parent_document_id": None,
        "due_date": due_date,
        "paid_date": None,
        "is_demo": is_demo,
        "created_at": now,
        "updated_at": now
    }
    
    await db.documents.insert_one(doc)
    
    result = await db.documents.find_one({"document_id": doc_id}, {"_id": 0})
    
    logger.info(f"Document created: {doc_id} ({doc_type}) by user {user['user_id']}")
    
    return result

@api_router.post("/documents/{document_id}/reupload")
async def reupload_document_pdf(
    document_id: str,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_agent)
):
    """Upload a revised PDF for an existing document with version tracking"""
    is_demo = user.get('is_demo', False)
    query = {"document_id": document_id, "agent_id": user['user_id']}
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Allow revision in Draft, Sent, or Change Requested status
    if doc['status'] not in ['Draft', 'Sent', 'Change Requested']:
        raise HTTPException(status_code=400, detail="Cannot revise document in current status")
    
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    now = datetime.now(timezone.utc).isoformat()
    current_version = doc.get('version', 1)
    
    # Archive current version in version history
    version_history = doc.get('version_history', []) or []
    if doc.get('pdf_path'):
        version_history.append({
            'version': current_version,
            'pdf_path': doc['pdf_path'],
            'pdf_filename': doc.get('pdf_filename'),
            'title': doc.get('title'),
            'amount': doc.get('amount'),
            'archived_at': now,
            'archived_by': user['user_id']
        })
    
    # Save new file
    file_id = uuid.uuid4().hex[:12]
    filename = f"{file_id}_{file.filename}"
    file_path = UPLOAD_DIR / filename
    
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Run AI extraction on new PDF
    extraction = await extract_document_from_pdf(str(file_path), file.filename)
    
    items_total = sum(item.get('total', 0) for item in extraction.get('items', []))
    amount = extraction.get('amount') or items_total or doc['amount']
    
    update_data = {
        "title": extraction.get('title') or doc['title'],
        "amount": amount,
        "items": extraction.get('items') if extraction.get('items') else doc['items'],
        "supplier_name": extraction.get('supplier_name') or doc.get('supplier_name'),
        "pdf_filename": file.filename,
        "pdf_path": str(file_path),
        "ai_extraction_confidence": extraction.get('confidence', 'low'),
        "updated_at": now,
        "version": current_version + 1,
        "version_history": version_history,
        # Reset to Draft if it was Change Requested
        "status": "Draft" if doc['status'] == 'Change Requested' else doc['status'],
        # Clear change request comment after revision
        "change_request_comment": None
    }
    
    await db.documents.update_one(query, {"$set": update_data})
    
    result = await db.documents.find_one(query, {"_id": 0})
    result['extraction_warning'] = extraction.get('extraction_failed', False) or extraction.get('amount') is None
    
    return result

@api_router.put("/documents/{document_id}")
async def update_document(document_id: str, data: DocumentUpdate, user: dict = Depends(get_current_agent)):
    """Update document with edited extraction data or reference changes"""
    is_demo = user.get('is_demo', False)
    query = {"document_id": document_id, "agent_id": user['user_id']}
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Cannot edit finalized documents (Approved, Rejected, Paid)
    if doc['status'] in ['Approved', 'Rejected', 'Paid']:
        raise HTTPException(status_code=400, detail=f"Cannot edit document with status '{doc['status']}'. Create a new version instead.")
    
    update_data = {}
    
    if data.title is not None:
        update_data['title'] = data.title
    if data.amount is not None:
        update_data['amount'] = data.amount
    if data.supplier_name is not None:
        update_data['supplier_name'] = data.supplier_name
    if data.notes is not None:
        update_data['notes'] = data.notes
    if data.summary is not None:
        update_data['summary'] = data.summary
    if data.hero_image_url is not None:
        update_data['hero_image_url'] = data.hero_image_url
    if data.items is not None:
        update_data['items'] = [item.model_dump() for item in data.items]
        if data.amount is None:
            update_data['amount'] = sum(item.total for item in data.items)
    
    # Handle reference changes (client, project, unit)
    if data.project_id is not None and data.project_id != doc.get('project_id'):
        # Validate project exists and belongs to agent
        project = await db.projects.find_one({
            "project_id": data.project_id, 
            "agent_id": user['user_id'], 
            "is_demo": is_demo
        })
        if not project:
            raise HTTPException(status_code=400, detail="Invalid project")
        update_data['project_id'] = data.project_id
    
    if data.client_id is not None and data.client_id != doc.get('client_id'):
        # Validate client exists and belongs to agent
        client = await db.clients.find_one({
            "client_id": data.client_id,
            "agent_id": user['user_id'],
            "is_demo": is_demo
        })
        if not client:
            raise HTTPException(status_code=400, detail="Invalid client")
        update_data['client_id'] = data.client_id
    
    if data.unit_id is not None:
        if data.unit_id == "":
            # Clear unit assignment
            update_data['unit_id'] = None
            update_data['unit_reference'] = "General"
        else:
            # Validate unit exists and belongs to the project
            target_project_id = data.project_id or doc.get('project_id')
            unit = await db.project_units.find_one({
                "unit_id": data.unit_id,
                "project_id": target_project_id,
                "is_demo": is_demo
            })
            if not unit:
                raise HTTPException(status_code=400, detail="Invalid unit for this project")
            update_data['unit_id'] = data.unit_id
            update_data['unit_reference'] = unit.get('unit_reference', unit.get('name', 'Unit'))
    
    update_data['updated_at'] = datetime.now(timezone.utc).isoformat()
    
    await db.documents.update_one(query, {"$set": update_data})
    return await db.documents.find_one(query, {"_id": 0})

@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str, force: bool = False, user: dict = Depends(get_current_agent)):
    """Delete a document. 
    By default only allowed for Draft status.
    Use force=true to delete documents in any status (for agents only).
    """
    is_demo = user.get('is_demo', False)
    query = {"document_id": document_id, "agent_id": user['user_id']}
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Only allow deletion of drafts unless force=true
    if doc['status'] != 'Draft' and not force:
        raise HTTPException(
            status_code=400, 
            detail="Can only delete documents in Draft status. Use force=true to delete sent documents, or archive them instead."
        )
    
    # Delete the document
    await db.documents.delete_one(query)
    
    # Clean up associated files if any
    if doc.get('pdf_path') and os.path.exists(doc['pdf_path']):
        try:
            os.remove(doc['pdf_path'])
        except Exception as e:
            logger.warning(f"Failed to delete PDF file: {e}")
    
    if doc.get('hero_image_url') and doc['hero_image_url'].startswith('/uploads/'):
        try:
            file_path = os.path.join(UPLOAD_DIR, doc['hero_image_url'].replace('/uploads/', ''))
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            logger.warning(f"Failed to delete hero image: {e}")
    
    logger.info(f"Document {document_id} deleted by {user['user_id']} (force={force})")
    
    return {"message": "Document deleted successfully"}

@api_router.post("/documents/{document_id}/revert-to-draft")
async def revert_document_to_draft(document_id: str, user: dict = Depends(get_current_agent)):
    """Revert a document back to Draft status.
    Allowed for: Sent, Change Requested statuses.
    Not allowed for: Approved, Rejected, Paid.
    """
    is_demo = user.get('is_demo', False)
    query = {"document_id": document_id, "agent_id": user['user_id']}
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Cannot revert finalized documents
    if doc['status'] in ['Approved', 'Rejected', 'Paid']:
        raise HTTPException(
            status_code=400, 
            detail=f"Cannot revert document with status '{doc['status']}'. This status is final."
        )
    
    if doc['status'] == 'Draft':
        return {"message": "Document is already in Draft status", "status": "Draft"}
    
    # Update status to Draft
    await db.documents.update_one(
        query,
        {"$set": {"status": "Draft", "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    logger.info(f"Document {document_id} reverted to Draft by {user['user_id']}")
    
    return {"message": "Document reverted to Draft", "status": "Draft"}

@api_router.get("/documents", response_model=List[DocumentResponse])
async def get_documents(
    doc_type: Optional[str] = None,
    status: Optional[str] = None,
    user: dict = Depends(get_current_user)
):
    """Get documents for current user - ownership scoped, schema validated"""
    query = {}
    
    if user['role'] == 'agent':
        query["agent_id"] = user['user_id']
    else:
        # Buyer: get documents via client relationship
        clients = await db.clients.find(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1}
        ).to_list(100)
        client_ids = [c['client_id'] for c in clients]
        query["client_id"] = {"$in": client_ids}
    
    if doc_type:
        query["type"] = doc_type
    if status:
        query["status"] = status
    
    docs = await db.documents.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return docs

@api_router.get("/documents/{document_id}")
async def get_document(document_id: str, user: dict = Depends(get_current_user)):
    """Get single document"""
    is_demo = user.get('is_demo', False)
    query = {"document_id": document_id}
    
    if user['role'] == 'agent':
        query["agent_id"] = user['user_id']
    else:
        clients = await db.clients.find(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1}
        ).to_list(100)
        client_ids = [c['client_id'] for c in clients]
        query["client_id"] = {"$in": client_ids}
    
    doc = await db.documents.find_one(query, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Include client and project info
    client = await db.clients.find_one({"client_id": doc['client_id']}, {"_id": 0})
    project = await db.projects.find_one({"project_id": doc['project_id']}, {"_id": 0})
    
    return {**doc, "client": client, "project": project}

@api_router.get("/documents/{document_id}/source-pdf")
async def get_document_source_pdf(document_id: str, user: dict = Depends(get_current_user)):
    """Get the original uploaded PDF for a document"""
    is_demo = user.get('is_demo', False)
    
    if user.get('role') == 'agent':
        query = {"document_id": document_id, "agent_id": user['user_id']}
    else:
        clients = await db.clients.find(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1}
        ).to_list(100)
        client_ids = [c['client_id'] for c in clients]
        query = {"document_id": document_id, "client_id": {"$in": client_ids}}
    
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not doc.get('pdf_path') or not os.path.exists(doc['pdf_path']):
        raise HTTPException(status_code=404, detail="Source PDF not found")
    
    return FileResponse(
        doc['pdf_path'],
        media_type="application/pdf",
        filename=doc.get('pdf_filename', f"document_{document_id}.pdf")
    )

@api_router.post("/documents/{document_id}/hero-image")
async def upload_hero_image(
    document_id: str,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_agent)
):
    """Upload a hero/banner image for a document"""
    is_demo = user.get('is_demo', False)
    query = {"document_id": document_id, "agent_id": user['user_id']}
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Validate image file
    allowed_types = ['image/jpeg', 'image/jpg', 'image/png', 'image/webp']
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Only JPEG, PNG, and WebP images are allowed")
    
    # Delete old hero image if exists
    if doc.get('hero_image_path') and os.path.exists(doc['hero_image_path']):
        try:
            os.remove(doc['hero_image_path'])
        except OSError:
            pass
    
    # Save new image
    file_ext = file.filename.split('.')[-1] if '.' in file.filename else 'jpg'
    image_filename = f"hero_{document_id}_{uuid.uuid4().hex[:8]}.{file_ext}"
    image_path = UPLOAD_DIR / image_filename
    
    content = await file.read()
    with open(image_path, "wb") as f:
        f.write(content)
    
    # Generate URL for the image
    hero_image_url = f"/api/documents/{document_id}/hero-image"
    
    await db.documents.update_one(query, {
        "$set": {
            "hero_image_url": hero_image_url,
            "hero_image_path": str(image_path),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
    })
    
    return {"hero_image_url": hero_image_url}

@api_router.get("/documents/{document_id}/hero-image")
async def get_hero_image(document_id: str, user: dict = Depends(get_current_user)):
    """Get the hero image for a document"""
    is_demo = user.get('is_demo', False)
    
    if user.get('role') == 'agent':
        query = {"document_id": document_id, "agent_id": user['user_id']}
    else:
        clients = await db.clients.find(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1}
        ).to_list(100)
        client_ids = [c['client_id'] for c in clients]
        query = {"document_id": document_id, "client_id": {"$in": client_ids}}
    
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not doc.get('hero_image_path') or not os.path.exists(doc['hero_image_path']):
        raise HTTPException(status_code=404, detail="Hero image not found")
    
    # Determine media type
    file_ext = doc['hero_image_path'].split('.')[-1].lower()
    media_types = {'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png', 'webp': 'image/webp'}
    media_type = media_types.get(file_ext, 'image/jpeg')
    
    return FileResponse(doc['hero_image_path'], media_type=media_type)

@api_router.delete("/documents/{document_id}/hero-image")
async def delete_hero_image(document_id: str, user: dict = Depends(get_current_agent)):
    """Delete the hero image from a document"""
    is_demo = user.get('is_demo', False)
    query = {"document_id": document_id, "agent_id": user['user_id']}
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file if exists
    if doc.get('hero_image_path') and os.path.exists(doc['hero_image_path']):
        try:
            os.remove(doc['hero_image_path'])
        except OSError:
            pass
    
    await db.documents.update_one(query, {
        "$set": {
            "hero_image_url": None,
            "hero_image_path": None,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
    })
    
    return {"message": "Hero image deleted"}

@api_router.get("/documents/{document_id}/qr-code")
async def get_document_qr_code(document_id: str, user: dict = Depends(get_current_user)):
    """Get Swiss QR payment code for an invoice"""
    is_demo = user.get('is_demo', False)
    
    if user.get('role') == 'agent':
        query = {"document_id": document_id, "agent_id": user['user_id']}
    else:
        clients = await db.clients.find(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1}
        ).to_list(100)
        client_ids = [c['client_id'] for c in clients]
        query = {"document_id": document_id, "client_id": {"$in": client_ids}}
    
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if doc['type'] != 'invoice':
        raise HTTPException(status_code=400, detail="QR code only available for invoices")
    
    # Get buyer name
    client = await db.clients.find_one({"client_id": doc['client_id']}, {"_id": 0})
    buyer_name = client.get('name', 'Client') if client else 'Client'
    
    # Get agent settings for IBAN
    agent = await db.users.find_one({"user_id": doc['agent_id']}, {"_id": 0})
    agent_settings = agent.get('settings', {}) if agent else {}
    billing_info = agent_settings.get('billing', {})
    
    # Generate QR code as base64 using agent's billing info or defaults
    qr_base64 = generate_swiss_qr_code_base64(
        amount=doc['amount'],
        reference=doc['document_number'],
        buyer_name=buyer_name,
        iban=billing_info.get('iban'),
        creditor_name=billing_info.get('company_name') or agent_settings.get('company_name'),
        creditor_address=billing_info.get('address'),
        creditor_pcode=billing_info.get('postal_code'),
        creditor_city=billing_info.get('city')
    )
    
    if not qr_base64:
        raise HTTPException(status_code=500, detail="Failed to generate QR code")
    
    # Get the IBAN used (from agent settings or default)
    iban_used = billing_info.get('iban') or DEFAULT_IBAN
    company_used = billing_info.get('company_name') or agent_settings.get('company_name') or DEFAULT_COMPANY_NAME
    
    return {
        "qr_code_svg_base64": qr_base64,
        "amount": doc['amount'],
        "currency": doc['currency'],
        "document_number": doc['document_number'],
        "payment_info": {
            "beneficiary": company_used,
            "iban": iban_used,
            "reference": doc['document_number']
        }
    }

@api_router.post("/documents/{document_id}/send")
async def send_document(document_id: str, user: dict = Depends(get_current_agent)):
    """
    Send document to buyer (transition from Draft/Change Requested to Sent).
    Returns detailed status including email delivery result.
    """
    is_demo = user.get('is_demo', False)
    query = {"document_id": document_id, "agent_id": user['user_id']}
    doc = await db.documents.find_one(query, {"_id": 0})
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Validate client exists
    client = await db.clients.find_one({"client_id": doc.get('client_id')}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=400, detail="Client not found. Please assign a client to this document first.")
    
    # Validate client has email
    if not client.get('email'):
        raise HTTPException(status_code=400, detail="Client has no email address. Please update client details first.")
    
    # Enforce state machine
    if not validate_transition(doc['type'], doc['status'], 'Sent'):
        raise HTTPException(status_code=400, detail=f"Cannot send document from status: {doc['status']}")
    
    now = datetime.now(timezone.utc).isoformat()
    
    await db.documents.update_one(query, {
        "$set": {
            "status": "Sent",
            "sent_at": now,
            "change_request_comment": None,
            "updated_at": now
        }
    })
    
    # Get related info for notifications
    project = await db.projects.find_one({"project_id": doc.get('project_id')}, {"_id": 0})
    agent_settings = user.get('settings', {})
    agent_profile = user.get('profile', {})
    
    # Track delivery status
    delivery_status = {
        "notification_created": False,
        "websocket_sent": False,
        "email_sent": False,
        "email_error": None
    }
    
    # Create notification for buyer if exists
    buyer_id = doc.get('buyer_id') or client.get('buyer_id')
    if buyer_id:
        try:
            await create_notification(
                user_id=buyer_id,
                title=f"New {doc['type'].title()} Received",
                message=f"You have a new {doc['type']} for {doc['title']} - CHF {doc['amount']:,.2f}",
                notification_type="document_sent",
                link="/buyer",
                is_demo=is_demo
            )
            delivery_status["notification_created"] = True
        except Exception as e:
            logger.warning(f"Failed to create notification: {e}")
        
        # Send real-time update via WebSocket
        try:
            await notify_realtime(
                [buyer_id],
                "document_sent",
                {
                    "document_id": document_id,
                    "type": doc['type'],
                    "title": doc['title'],
                    "amount": doc['amount']
                }
            )
            delivery_status["websocket_sent"] = True
        except Exception as e:
            logger.warning(f"Failed to send WebSocket notification: {e}")
    
    # Send email notification to buyer
    if client.get('email'):
        try:
            email_result = await send_notification_email(
                "document_sent",
                client['email'],
                {
                    "doc_type": doc['type'],
                    "buyer_name": client.get('name', 'there'),
                    "agent_name": agent_profile.get('display_name') or user.get('name', 'Your agent'),
                    "company_name": agent_settings.get('company_name', 'the agency'),
                    "agent_email": agent_profile.get('contact_email', ''),
                    "agent_phone": agent_profile.get('contact_phone', ''),
                    "title": doc.get('title', 'Document'),
                    "summary": doc.get('summary', ''),
                    "currency": doc.get('currency', 'CHF'),
                    "amount": doc.get('amount', 0),
                    "project_name": project.get('name', 'N/A') if project else 'N/A',
                    "unit_reference": doc.get('unit_reference', 'N/A')
                }
            )
            # Check if email was sent - status is "success" on success, "error" or "skipped" on failure
            if isinstance(email_result, dict):
                delivery_status["email_sent"] = email_result.get("status") == "success"
                if not delivery_status["email_sent"]:
                    delivery_status["email_error"] = email_result.get("error") or email_result.get("reason", "Unknown error")
            else:
                delivery_status["email_sent"] = True
        except Exception as e:
            logger.warning(f"Failed to send email notification: {e}")
            delivery_status["email_error"] = str(e)
    
    # Return detailed response
    return {
        "message": "Document sent successfully",
        "status": "Sent",
        "document_id": document_id,
        "recipient": {
            "name": client.get('name'),
            "email": client.get('email')
        },
        "delivery": delivery_status,
        "warnings": [] if delivery_status["email_sent"] else [
            f"Email may not have been delivered: {delivery_status['email_error'] or 'Unknown reason'}"
        ]
    }

@api_router.post("/documents/{document_id}/action")
async def document_action(document_id: str, action_data: DocumentAction, user: dict = Depends(get_current_user)):
    """Perform action on document (buyer or agent)"""
    is_demo = user.get('is_demo', False)
    action = action_data.action
    
    # Get document
    if user['role'] == 'agent':
        query = {"document_id": document_id, "agent_id": user['user_id']}
    else:
        clients = await db.clients.find(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1}
        ).to_list(100)
        client_ids = [c['client_id'] for c in clients]
        query = {"document_id": document_id, "client_id": {"$in": client_ids}}
    
    doc = await db.documents.find_one(query, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Handle different actions based on document type and user role
    if action == 'approve' and doc['type'] == 'quote' and user['role'] == 'buyer':
        if not validate_transition('quote', doc['status'], 'Approved'):
            raise HTTPException(status_code=400, detail="Cannot approve quote in current status")
        
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {"status": "Approved", "updated_at": now}}
        )
        
        # Get agent info for email
        agent = await db.users.find_one({"user_id": doc.get('agent_id')}, {"_id": 0, "email": 1, "name": 1})
        
        # Notify agent
        if doc.get('agent_id'):
            await create_notification(
                user_id=doc['agent_id'],
                title="Quote Approved",
                message=f"{user.get('name', 'Buyer')} approved quote {doc['document_number']}",
                notification_type="quote_approved",
                link=f"/agent/documents/{document_id}",
                is_demo=is_demo
            )
            # Send real-time update via WebSocket
            await notify_realtime(
                [doc['agent_id']],
                "quote_approved",
                {
                    "document_id": document_id,
                    "title": doc['title'],
                    "buyer_name": user.get('name', 'Buyer')
                }
            )
            
            # Send email notification to agent (instant - critical event)
            if agent and agent.get('email'):
                await send_notification_email(
                    "quote_approved",
                    agent['email'],
                    {
                        "buyer_name": user.get('name', 'Your client'),
                        "document_number": doc.get('document_number', ''),
                        "document_id": document_id,
                        "title": doc.get('title', 'Quote'),
                        "currency": doc.get('currency', 'CHF'),
                        "amount": doc.get('amount', 0)
                    }
                )
        
        return {"message": "Quote approved", "status": "Approved"}
    
    elif action == 'reject' and doc['type'] == 'quote' and user['role'] == 'buyer':
        if not validate_transition('quote', doc['status'], 'Rejected'):
            raise HTTPException(status_code=400, detail="Cannot reject quote in current status")
        
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {"status": "Rejected", "updated_at": now}}
        )
        
        return {"message": "Quote rejected", "status": "Rejected"}
    
    elif action == 'request_change' and doc['type'] == 'quote' and user['role'] == 'buyer':
        if not action_data.comment:
            raise HTTPException(status_code=400, detail="Comment required for change request")
        
        if not validate_transition('quote', doc['status'], 'Change Requested'):
            raise HTTPException(status_code=400, detail="Cannot request changes in current status")
        
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {
                "status": "Change Requested",
                "change_request_comment": action_data.comment,
                "updated_at": now
            }}
        )
        
        # Notify agent
        if doc.get('agent_id'):
            await create_notification(
                user_id=doc['agent_id'],
                title="Change Requested",
                message=f"Changes requested for quote {doc['document_number']}",
                notification_type="change_requested",
                link=f"/agent/documents/{document_id}",
                metadata={"comment": action_data.comment},
                is_demo=is_demo
            )
            
            # Send email notification to agent (instant - critical event)
            agent = await db.users.find_one({"user_id": doc.get('agent_id')}, {"_id": 0, "email": 1})
            if agent and agent.get('email'):
                await send_notification_email(
                    "change_requested",
                    agent['email'],
                    {
                        "buyer_name": user.get('name', 'Your client'),
                        "document_number": doc.get('document_number', ''),
                        "document_id": document_id,
                        "title": doc.get('title', 'Quote'),
                        "comment": action_data.comment
                    }
                )
        
        return {"message": "Change requested", "status": "Change Requested"}
    
    elif action == 'request_change' and doc['type'] == 'invoice' and user['role'] == 'buyer':
        if not action_data.comment:
            raise HTTPException(status_code=400, detail="Comment required for change request")
        
        if not validate_transition('invoice', doc['status'], 'Change Requested'):
            raise HTTPException(status_code=400, detail="Cannot request changes in current status")
        
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {
                "status": "Change Requested",
                "change_request_comment": action_data.comment,
                "updated_at": now
            }}
        )
        
        # Notify agent
        if doc.get('agent_id'):
            await create_notification(
                user_id=doc['agent_id'],
                title="Change Requested",
                message=f"Changes requested for invoice {doc['document_number']}",
                notification_type="change_requested",
                link=f"/agent/documents/{document_id}",
                metadata={"comment": action_data.comment},
                is_demo=is_demo
            )
            
            # Send email notification to agent (instant - critical event)
            agent = await db.users.find_one({"user_id": doc.get('agent_id')}, {"_id": 0, "email": 1})
            if agent and agent.get('email'):
                await send_notification_email(
                    "change_requested",
                    agent['email'],
                    {
                        "buyer_name": user.get('name', 'Your client'),
                        "document_number": doc.get('document_number', ''),
                        "document_id": document_id,
                        "title": doc.get('title', 'Invoice'),
                        "comment": action_data.comment
                    }
                )
        
        return {"message": "Change requested", "status": "Change Requested"}
    
    elif action == 'confirm_payment' and doc['type'] == 'invoice' and user['role'] == 'buyer':
        if not validate_transition('invoice', doc['status'], 'Paid'):
            raise HTTPException(status_code=400, detail="Cannot confirm payment in current status")
        
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {"status": "Paid", "paid_date": now, "updated_at": now}}
        )
        
        # Notify agent
        if doc.get('agent_id'):
            await create_notification(
                user_id=doc['agent_id'],
                title="Payment Confirmed",
                message=f"Payment confirmed for invoice {doc['document_number']}",
                notification_type="payment_confirmed",
                link=f"/agent/documents/{document_id}",
                is_demo=is_demo
            )
            
            # Send email notification to agent (instant - critical event)
            agent = await db.users.find_one({"user_id": doc.get('agent_id')}, {"_id": 0, "email": 1})
            if agent and agent.get('email'):
                await send_notification_email(
                    "payment_confirmed",
                    agent['email'],
                    {
                        "document_number": doc.get('document_number', ''),
                        "document_id": document_id,
                        "title": doc.get('title', 'Invoice'),
                        "currency": doc.get('currency', 'CHF'),
                        "amount": doc.get('amount', 0)
                    }
                )
        
        return {"message": "Payment confirmed", "status": "Paid"}
    
    elif action == 'convert_to_invoice' and doc['type'] == 'quote' and user['role'] == 'agent':
        if doc['status'] != 'Approved':
            raise HTTPException(status_code=400, detail="Only approved quotes can be converted to invoices")
        
        # ATOMIC: Create new invoice document from quote
        count = await db.documents.count_documents({"agent_id": user['user_id'], "type": "invoice"})
        invoice_number = f"INV-{datetime.now().year}-{str(count + 1).zfill(4)}"
        invoice_id = f"doc_{uuid.uuid4().hex[:12]}"
        
        invoice_doc = {
            "document_id": invoice_id,
            "document_number": invoice_number,
            "type": "invoice",
            "status": "Draft",
            "agent_id": doc['agent_id'],
            "client_id": doc['client_id'],
            "buyer_id": doc.get('buyer_id'),
            "project_id": doc['project_id'],
            "unit_reference": doc['unit_reference'],
            "title": doc['title'],  # Copy from quote
            "amount": doc['amount'],  # Copy from quote
            "items": doc['items'],  # Copy from quote
            "currency": doc['currency'],
            "supplier_name": doc.get('supplier_name'),
            "notes": doc.get('notes'),
            "change_request_comment": None,
            "pdf_filename": None,
            "pdf_path": None,
            "ai_extraction_confidence": None,
            "parent_document_id": document_id,  # Link to source quote
            "due_date": (datetime.now(timezone.utc) + timedelta(days=30)).isoformat(),
            "paid_date": None,
            "is_demo": is_demo,
            "created_at": now,
            "updated_at": now
        }
        
        await db.documents.insert_one(invoice_doc)
        
        # Notify buyer
        if doc.get('buyer_id'):
            await create_notification(
                user_id=doc['buyer_id'],
                title="Invoice Ready",
                message=f"Invoice {invoice_number} for CHF {doc['amount']:,.2f} is ready for payment",
                notification_type="invoice_created",
                link="/buyer",
                is_demo=is_demo
            )
        
        return {
            "message": "Invoice created",
            "invoice_id": invoice_id,
            "invoice_number": invoice_number
        }
    
    else:
        raise HTTPException(status_code=400, detail=f"Invalid action: {action}")

# ==================== TIMELINE ENDPOINT (SINGLE SOURCE) ====================

@api_router.get("/timeline")
async def get_timeline(user: dict = Depends(get_current_user)):
    """Get unified timeline - ALL documents sorted by created_at - ownership scoped"""
    if user['role'] == 'buyer':
        # Get buyer's clients (ownership scoped)
        clients = await db.clients.find(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1, "project_id": 1, "unit_reference": 1}
        ).to_list(100)
        client_ids = [c['client_id'] for c in clients]
        
        if not client_ids:
            return {"documents": [], "project_info": None}
        
        # Get all documents for buyer's clients (not Draft status for buyers)
        # Note: ownership scoped by client_id, not is_demo
        docs = await db.documents.find(
            {"client_id": {"$in": client_ids}, "status": {"$ne": "Draft"}},
            {"_id": 0}
        ).sort("created_at", -1).to_list(100)
        
        # Get project info
        project_id = clients[0]['project_id'] if clients else None
        project = await db.projects.find_one({"project_id": project_id}, {"_id": 0}) if project_id else None
        
        if project and clients:
            project['unit_reference'] = clients[0].get('unit_reference', '')
        
        # Format for frontend
        timeline_events = []
        for doc in docs:
            timeline_events.append({
                "id": doc['document_id'],
                "type": doc['type'],
                "title": doc['title'],
                "status": doc['status'],
                "amount": doc['amount'],
                "date": doc['updated_at'] or doc['created_at'],
                "dueDate": doc.get('due_date'),
                "items": doc.get('items', []),
                "changeComment": doc.get('change_request_comment'),
                "actionRequired": (doc['type'] == 'quote' and doc['status'] == 'Sent') or 
                                 (doc['type'] == 'invoice' and doc['status'] == 'Sent'),
                "hasSourcePdf": bool(doc.get('pdf_path')),
                "supplierName": doc.get('supplier_name'),
                "documentNumber": doc['document_number'],
                "parentDocumentId": doc.get('parent_document_id'),
                "summary": doc.get('summary', ''),
                "heroImageUrl": doc.get('hero_image_url'),
                "currency": doc.get('currency', 'CHF')
            })
        
        return {"documents": timeline_events, "project_info": project}
    
    else:
        # Agent view - all their documents
        docs = await db.documents.find(
            {"agent_id": user['user_id']},
            {"_id": 0}
        ).sort("created_at", -1).to_list(100)
        
        return {"documents": docs}

# ==================== PDF GENERATION ====================

@api_router.get("/documents/{document_id}/pdf")
async def get_document_pdf(document_id: str, user: dict = Depends(get_current_user)):
    """Generate and return PDF for document"""
    is_demo = user.get('is_demo', False)
    
    if user['role'] == 'agent':
        query = {"document_id": document_id, "agent_id": user['user_id']}
    else:
        clients = await db.clients.find(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1}
        ).to_list(100)
        client_ids = [c['client_id'] for c in clients]
        query = {"document_id": document_id, "client_id": {"$in": client_ids}}
    
    doc = await db.documents.find_one(query, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    client = await db.clients.find_one({"client_id": doc['client_id']}, {"_id": 0})
    project = await db.projects.find_one({"project_id": doc['project_id']}, {"_id": 0})
    
    # Generate PDF
    buffer = BytesIO()
    pdf_doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CompanyName', fontSize=18, fontName='Helvetica-Bold', spaceAfter=6))
    styles.add(ParagraphStyle(name='DocTitle', fontSize=24, fontName='Helvetica-Bold', spaceAfter=20))
    styles.add(ParagraphStyle(name='SectionTitle', fontSize=12, fontName='Helvetica-Bold', spaceBefore=16, spaceAfter=8))
    styles.add(ParagraphStyle(name='Normal12', fontSize=10, fontName='Helvetica', spaceAfter=4))
    
    elements = []
    
    # Header
    elements.append(Paragraph("UpgradeFlow", styles['CompanyName']))
    elements.append(Paragraph("Real Estate Post-Sale Management", styles['Normal12']))
    elements.append(Spacer(1, 10*mm))
    
    # Title
    doc_type_label = "INVOICE" if doc['type'] == 'invoice' else "QUOTE"
    elements.append(Paragraph(doc_type_label, styles['DocTitle']))
    
    # Info table
    created_date = doc['created_at']
    if isinstance(created_date, str):
        created_date = datetime.fromisoformat(created_date.replace('Z', '+00:00'))
    
    info_data = [
        [f"{doc_type_label[0]}# :", doc['document_number']],
        ["Date:", created_date.strftime("%d %B %Y")],
        ["Status:", doc['status']],
        ["Project:", project['name'] if project else 'N/A'],
        ["Unit:", doc['unit_reference']],
    ]
    
    if doc['type'] == 'invoice' and doc.get('due_date'):
        due_date = doc['due_date']
        if isinstance(due_date, str):
            due_date = datetime.fromisoformat(due_date.replace('Z', '+00:00'))
        info_data.append(["Due Date:", due_date.strftime("%d %B %Y") if isinstance(due_date, datetime) else due_date])
    
    info_table = Table(info_data, colWidths=[35*mm, 100*mm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 10*mm))
    
    # Client info
    elements.append(Paragraph("Bill To:", styles['SectionTitle']))
    if client:
        elements.append(Paragraph(client['name'], styles['Normal12']))
        elements.append(Paragraph(client['email'], styles['Normal12']))
    elements.append(Spacer(1, 8*mm))
    
    # Title and description
    elements.append(Paragraph(doc['title'], styles['SectionTitle']))
    elements.append(Spacer(1, 6*mm))
    
    # Line items
    if doc.get('items') and len(doc['items']) > 0:
        table_data = [["Description", "Qty", "Unit Price", "Total"]]
        for item in doc['items']:
            table_data.append([
                item['description'],
                str(item['quantity']),
                f"CHF {item['unit_price']:,.2f}",
                f"CHF {item['total']:,.2f}"
            ])
        
        # Add total row
        table_data.append(["", "", "Total:", f"CHF {doc['amount']:,.2f}"])
        
        items_table = Table(table_data, colWidths=[85*mm, 15*mm, 35*mm, 35*mm])
        items_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -2), 0.5, colors.HexColor('#e5e7eb')),
            ('FONTNAME', (-2, -1), (-1, -1), 'Helvetica-Bold'),
            ('LINEABOVE', (0, -1), (-1, -1), 1, colors.black),
        ]))
        elements.append(items_table)
    else:
        # No line items, just show total
        elements.append(Paragraph(f"Total: CHF {doc['amount']:,.2f}", styles['SectionTitle']))
    
    elements.append(Spacer(1, 10*mm))
    
    # Notes
    if doc.get('notes'):
        elements.append(Paragraph("Notes:", styles['SectionTitle']))
        elements.append(Paragraph(doc['notes'], styles['Normal12']))
    
    # Footer
    elements.append(Spacer(1, 20*mm))
    footer_style = ParagraphStyle(name='Footer', fontSize=8, textColor=colors.grey)
    elements.append(Paragraph("UpgradeFlow SA • Rue du Rhône 1, 1204 Genève • Switzerland", footer_style))
    
    pdf_doc.build(elements)
    buffer.seek(0)
    
    filename = f"{doc['type']}_{doc['document_number']}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ==================== NOTIFICATION ENDPOINTS ====================

@api_router.get("/notifications")
async def get_notifications(user: dict = Depends(get_current_user)):
    """Get notifications for current user (ownership-scoped)"""
    notifications = await db.notifications.find(
        {"user_id": user['user_id']},
        {"_id": 0}
    ).sort("created_at", -1).limit(50).to_list(50)
    
    unread_count = await db.notifications.count_documents(
        {"user_id": user['user_id'], "is_read": False}
    )
    
    return {"notifications": notifications, "unread_count": unread_count}

@api_router.put("/notifications/read-all")
@api_router.post("/notifications/mark-all-read")
async def mark_all_notifications_read(user: dict = Depends(get_current_user)):
    """Mark all notifications as read"""
    await db.notifications.update_many(
        {"user_id": user['user_id'], "is_read": False},
        {"$set": {"is_read": True}}
    )
    return {"message": "All notifications marked as read"}

@api_router.put("/notifications/{notification_id}/read")
@api_router.post("/notifications/{notification_id}/read")
async def mark_notification_read(notification_id: str, user: dict = Depends(get_current_user)):
    """Mark a single notification as read (called when notification is opened/clicked)"""
    result = await db.notifications.update_one(
        {"notification_id": notification_id, "user_id": user['user_id']},
        {"$set": {"is_read": True}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Notification not found")
    return {"message": "Notification marked as read"}

# ==================== PROJECT STAGES ENDPOINTS ====================

@api_router.get("/projects/{project_id}/stages")
async def get_project_stages(project_id: str, user: dict = Depends(get_current_user)):
    """Get all stages for a project
    
    SOURCE OF TRUTH: timeline_steps (via project_timelines)
    No fallback. No dual-read.
    """
    if user['role'] == 'agent':
        project = await db.projects.find_one(
            {"project_id": project_id, "agent_id": user['user_id']},
            {"_id": 0}
        )
    else:
        clients = await db.clients.find(
            {"buyer_id": user['user_id'], "project_id": project_id},
            {"_id": 0}
        ).to_list(10)
        if not clients:
            raise HTTPException(status_code=403, detail="No access to this project")
        project = await db.projects.find_one({"project_id": project_id}, {"_id": 0})
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Single source: timeline_steps
    timeline = await db.project_timelines.find_one(
        {"project_id": project_id},
        {"_id": 0}
    )
    
    if not timeline:
        # No timeline exists yet - return empty stages
        return {"project": project, "stages": []}
    
    timeline_id = timeline.get('timeline_id') or timeline.get('project_timeline_id')
    steps = await db.timeline_steps.find(
        {"project_timeline_id": timeline_id},
        {"_id": 0}
    ).sort("order_index", 1).to_list(50)
    
    # Map to ProjectStage format for API compatibility
    stages = []
    for step in steps:
        stages.append({
            "stage_id": step.get('step_id'),
            "project_id": project_id,
            "agent_id": project.get('agent_id'),
            "name": step.get('title'),
            "description": step.get('description'),
            "order": step.get('order_index', 0),
            "planned_start": step.get('planned_start') or step.get('planned_date'),
            "planned_end": step.get('planned_end'),
            "actual_start": step.get('actual_start'),
            "actual_end": step.get('completed_at'),
            "status": _map_timeline_status_to_stage(step.get('status', 'pending')),
            "progress_percent": step.get('progress_percent', 0),
            "notes": step.get('notes'),
            "dependencies": step.get('dependencies', []),
            "is_demo": step.get('is_demo', False),
            "created_at": step.get('created_at'),
            "updated_at": step.get('updated_at')
        })
    
    return {"project": project, "stages": stages}


def _map_timeline_status_to_stage(timeline_status: str) -> str:
    """Map TimelineStep status to ProjectStage status for API compatibility"""
    mapping = {
        'pending': 'upcoming',
        'in_progress': 'in_progress',
        'completed': 'completed',
        'approved': 'completed'
    }
    return mapping.get(timeline_status, 'upcoming')


def _map_stage_status_to_timeline(stage_status: str) -> str:
    """Map ProjectStage status to TimelineStep status"""
    mapping = {
        'upcoming': 'pending',
        'in_progress': 'in_progress',
        'completed': 'completed',
        'delayed': 'in_progress',
        'on_hold': 'pending'
    }
    return mapping.get(stage_status, 'pending')

@api_router.post("/projects/{project_id}/stages")
async def create_project_stage(project_id: str, data: ProjectStageCreate, user: dict = Depends(get_current_agent)):
    """Create a new project stage
    
    SOURCE OF TRUTH: timeline_steps (via project_timelines)
    Creates timeline if it doesn't exist.
    """
    is_demo = user.get('is_demo', False)
    
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id'], "is_demo": is_demo},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Ensure timeline exists for this project
    timeline = await db.project_timelines.find_one({"project_id": project_id, "is_demo": is_demo}, {"_id": 0})
    if not timeline:
        timeline_id = f"timeline_{uuid.uuid4().hex[:12]}"
        timeline = {
            "timeline_id": timeline_id,
            "project_timeline_id": timeline_id,
            "project_id": project_id,
            "template_id": None,
            "template_name": None,
            "is_demo": is_demo,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.project_timelines.insert_one(timeline)
    
    timeline_id = timeline.get('timeline_id') or timeline.get('project_timeline_id')
    step_id = f"step_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    # Write to timeline_steps (single source of truth)
    step_doc = {
        "step_id": step_id,
        "project_timeline_id": timeline_id,
        "title": data.name,
        "description": data.description,
        "order_index": data.order,
        "planned_date": data.planned_start,
        "planned_start": data.planned_start,
        "planned_end": data.planned_end,
        "actual_start": None,
        "completed_at": None,
        "status": "pending",
        "progress_percent": 0,
        "notes": None,
        "dependencies": data.dependencies or [],
        "is_demo": is_demo,
        "created_at": now,
        "updated_at": now
    }
    
    await db.timeline_steps.insert_one(step_doc)
    
    # Return in ProjectStage format for API compatibility
    return {
        "stage_id": step_id,
        "project_id": project_id,
        "agent_id": user['user_id'],
        "name": data.name,
        "description": data.description,
        "order": data.order,
        "planned_start": data.planned_start,
        "planned_end": data.planned_end,
        "actual_start": None,
        "actual_end": None,
        "status": "upcoming",
        "progress_percent": 0,
        "notes": None,
        "dependencies": data.dependencies or [],
        "is_demo": is_demo,
        "created_at": now,
        "updated_at": now
    }

@api_router.put("/projects/{project_id}/stages/{stage_id}")
async def update_project_stage(project_id: str, stage_id: str, data: ProjectStageUpdate, user: dict = Depends(get_current_agent)):
    """Update a project stage
    
    SOURCE OF TRUTH: timeline_steps
    """
    # Find the step in timeline_steps
    step = await db.timeline_steps.find_one({"step_id": stage_id}, {"_id": 0})
    if not step:
        raise HTTPException(status_code=404, detail="Stage not found")
    
    # Verify ownership via timeline -> project
    timeline = await db.project_timelines.find_one(
        {"$or": [
            {"timeline_id": step['project_timeline_id']},
            {"project_timeline_id": step['project_timeline_id']}
        ]},
        {"_id": 0}
    )
    if not timeline or timeline.get('project_id') != project_id:
        raise HTTPException(status_code=404, detail="Stage not found in this project")
    
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    now = datetime.now(timezone.utc).isoformat()
    update_data = {"updated_at": now}
    
    # Map ProjectStageUpdate fields to TimelineStep fields
    if data.name is not None:
        update_data['title'] = data.name
    if data.description is not None:
        update_data['description'] = data.description
    if data.order is not None:
        update_data['order_index'] = data.order
    if data.planned_start is not None:
        update_data['planned_start'] = data.planned_start
        update_data['planned_date'] = data.planned_start
    if data.planned_end is not None:
        update_data['planned_end'] = data.planned_end
    if data.actual_start is not None:
        update_data['actual_start'] = data.actual_start
    if data.actual_end is not None:
        update_data['completed_at'] = data.actual_end
    if data.status is not None:
        update_data['status'] = _map_stage_status_to_timeline(data.status)
    if data.progress_percent is not None:
        update_data['progress_percent'] = data.progress_percent
    if data.notes is not None:
        update_data['notes'] = data.notes
    
    await db.timeline_steps.update_one({"step_id": stage_id}, {"$set": update_data})
    
    # Return updated stage in ProjectStage format
    updated_step = await db.timeline_steps.find_one({"step_id": stage_id}, {"_id": 0})
    return {
        "stage_id": updated_step.get('step_id'),
        "project_id": project_id,
        "agent_id": user['user_id'],
        "name": updated_step.get('title'),
        "description": updated_step.get('description'),
        "order": updated_step.get('order_index', 0),
        "planned_start": updated_step.get('planned_start'),
        "planned_end": updated_step.get('planned_end'),
        "actual_start": updated_step.get('actual_start'),
        "actual_end": updated_step.get('completed_at'),
        "status": _map_timeline_status_to_stage(updated_step.get('status', 'pending')),
        "progress_percent": updated_step.get('progress_percent', 0),
        "notes": updated_step.get('notes'),
        "dependencies": updated_step.get('dependencies', []),
        "is_demo": updated_step.get('is_demo', False),
        "created_at": updated_step.get('created_at'),
        "updated_at": updated_step.get('updated_at')
    }

@api_router.delete("/projects/{project_id}/stages/{stage_id}")
async def delete_project_stage(project_id: str, stage_id: str, user: dict = Depends(get_current_agent)):
    """Delete a project stage
    
    SOURCE OF TRUTH: timeline_steps
    """
    # Find the step
    step = await db.timeline_steps.find_one({"step_id": stage_id}, {"_id": 0})
    if not step:
        raise HTTPException(status_code=404, detail="Stage not found")
    
    # Verify ownership via timeline -> project
    timeline = await db.project_timelines.find_one(
        {"$or": [
            {"timeline_id": step['project_timeline_id']},
            {"project_timeline_id": step['project_timeline_id']}
        ]},
        {"_id": 0}
    )
    if not timeline or timeline.get('project_id') != project_id:
        raise HTTPException(status_code=404, detail="Stage not found in this project")
    
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    result = await db.timeline_steps.delete_one({"step_id": stage_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Stage not found")
    
    return {"message": "Stage deleted"}

# ==================== COMPOSITE ENDPOINTS (Phase 2) ====================
# One canonical endpoint per major screen. Backend is single source of truth.
# No fallbacks. No reconstruction logic. Frontend renders only.

@api_router.get("/agent/dashboard", response_model=DashboardResponse)
async def get_agent_dashboard(user: dict = Depends(get_current_agent)):
    """
    Canonical endpoint for AgentHomePage (Command Center).
    Returns: projects, selected_project, recent_work
    Single source of truth. No client-side reconstruction.
    """
    # Get all projects for this agent
    projects_cursor = db.projects.find(
        {"agent_id": user['user_id']},
        {"_id": 0}
    ).sort("created_at", -1)
    projects_raw = await projects_cursor.to_list(100)
    
    projects = [
        ProjectSummary(
            project_id=p['project_id'],
            name=p['name'],
            address=p.get('address'),
            status=p.get('status'),
            created_at=p.get('created_at')
        ) for p in projects_raw
    ]
    
    # Selected project (first by default, or from user preference later)
    selected_project = projects[0] if projects else None
    
    # Recent work - last 6 activities/modifications
    recent_work = []
    
    # Get recent clients
    recent_clients = await db.clients.find(
        {"agent_id": user['user_id']},
        {"_id": 0}
    ).sort("updated_at", -1).limit(3).to_list(3)
    
    for c in recent_clients:
        recent_work.append(RecentWorkItem(
            id=c['client_id'],
            type='client',
            title=c['name'],
            subtitle=c.get('email'),
            path=f"/agent/clients/{c['client_id']}",
            timestamp=c.get('updated_at')
        ))
    
    # Get recent documents
    recent_docs = await db.documents.find(
        {"agent_id": user['user_id']},
        {"_id": 0}
    ).sort("updated_at", -1).limit(3).to_list(3)
    
    for d in recent_docs:
        doc_type = d.get('document_type') or d.get('type', 'document')
        recent_work.append(RecentWorkItem(
            id=d['document_id'],
            type='document',
            title=d.get('title') or f"{doc_type.title()} #{d.get('document_number', '')}",
            subtitle=d.get('client_name'),
            path=f"/agent/{doc_type}s/{d['document_id']}",
            timestamp=d.get('updated_at')
        ))
    
    # Sort by timestamp and limit to 6
    recent_work.sort(key=lambda x: x.timestamp or '', reverse=True)
    recent_work = recent_work[:6]
    
    return DashboardResponse(
        projects=projects,
        selected_project=selected_project,
        recent_work=recent_work
    )



class RecentWorkResponse(BaseModel):
    items: List[RecentWorkItem] = []

@api_router.get("/command/recent-work", response_model=RecentWorkResponse)
async def get_recent_work(user: dict = Depends(get_current_agent)):
    """
    Returns recent work items for the agent.
    Independent of project selection.
    """
    recent_work = []
    
    # Get recent clients
    recent_clients = await db.clients.find(
        {"agent_id": user['user_id']},
        {"_id": 0}
    ).sort("updated_at", -1).limit(3).to_list(3)
    
    for c in recent_clients:
        recent_work.append(RecentWorkItem(
            id=c['client_id'],
            type='client',
            title=c['name'],
            subtitle=c.get('email'),
            path=f"/agent/clients/{c['client_id']}",
            timestamp=c.get('updated_at')
        ))
    
    # Get recent documents
    recent_docs = await db.documents.find(
        {"agent_id": user['user_id']},
        {"_id": 0}
    ).sort("updated_at", -1).limit(3).to_list(3)
    
    for d in recent_docs:
        doc_type = d.get('document_type') or d.get('type', 'document')
        recent_work.append(RecentWorkItem(
            id=d['document_id'],
            type='document',
            title=d.get('title') or f"{doc_type.title()} #{d.get('document_number', '')}",
            subtitle=d.get('client_name'),
            path=f"/agent/{doc_type}s/{d['document_id']}",
            timestamp=d.get('updated_at')
        ))
    
    # Sort by timestamp and limit to 6
    recent_work.sort(key=lambda x: x.timestamp or '', reverse=True)
    recent_work = recent_work[:6]
    
    return RecentWorkResponse(items=recent_work)


@api_router.get("/projects/{project_id}/context", response_model=ProjectContextResponse)
async def get_project_context(project_id: str, user: dict = Depends(get_current_agent)):
    """
    Canonical endpoint for project context (clients + units).
    Called when project selection changes.
    Single source of truth. No fallbacks.
    """
    # Get project
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    project_summary = ProjectSummary(
        project_id=project['project_id'],
        name=project['name'],
        address=project.get('address'),
        status=project.get('status'),
        created_at=project.get('created_at')
    )
    
    # Get clients for this project
    clients_raw = await db.clients.find(
        {"project_id": project_id, "agent_id": user['user_id']},
        {"_id": 0}
    ).to_list(100)
    
    clients = [
        ClientSummary(
            client_id=c['client_id'],
            name=c['name'],
            email=c.get('email'),
            project_id=c.get('project_id'),
            unit_id=c.get('unit_id')
        ) for c in clients_raw
    ]
    
    # Get units for this project
    units_raw = await db.project_units.find(
        {"project_id": project_id},
        {"_id": 0}
    ).to_list(100)
    
    units = [
        UnitSummary(
            unit_id=u['unit_id'],
            reference=u.get('unit_reference', u.get('reference', '')),
            type=u.get('unit_type', u.get('type')),
            client_id=u.get('client_id')
        ) for u in units_raw
    ]
    
    return ProjectContextResponse(
        project=project_summary,
        clients=clients,
        units=units
    )


@api_router.get("/projects/{project_id}/timeline/full", response_model=ProjectTimelineResponse)
async def get_project_timeline_full(project_id: str, user: dict = Depends(get_current_user)):
    """
    Canonical endpoint for AgentTimeline.
    Returns: project, timeline_id, steps, progress
    Single source of truth: timeline_steps collection.
    """
    # Access control
    if user['role'] == 'agent':
        project = await db.projects.find_one(
            {"project_id": project_id, "agent_id": user['user_id']},
            {"_id": 0}
        )
    else:
        # Buyer access via client link
        client = await db.clients.find_one(
            {"buyer_id": user['user_id'], "project_id": project_id},
            {"_id": 0}
        )
        if not client:
            raise HTTPException(status_code=403, detail="No access to this project")
        project = await db.projects.find_one({"project_id": project_id}, {"_id": 0})
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    project_summary = ProjectSummary(
        project_id=project['project_id'],
        name=project['name'],
        address=project.get('address'),
        status=project.get('status'),
        created_at=project.get('created_at')
    )
    
    # Get timeline
    timeline = await db.project_timelines.find_one(
        {"project_id": project_id},
        {"_id": 0}
    )
    
    if not timeline:
        return ProjectTimelineResponse(
            project=project_summary,
            timeline_id=None,
            steps=[],
            progress_percent=0
        )
    
    timeline_id = timeline.get('timeline_id') or timeline.get('project_timeline_id')
    
    # Get steps from single source of truth
    steps_raw = await db.timeline_steps.find(
        {"project_timeline_id": timeline_id},
        {"_id": 0}
    ).sort("order_index", 1).to_list(50)
    
    steps = [
        TimelineStepSummary(
            step_id=s['step_id'],
            title=s['title'],
            description=s.get('description'),
            status=s.get('status', 'pending'),
            order_index=s.get('order_index', 0),
            planned_start=s.get('planned_start') or s.get('planned_date'),
            planned_end=s.get('planned_end'),
            progress_percent=s.get('progress_percent', 0)
        ) for s in steps_raw
    ]
    
    # Calculate overall progress
    if steps:
        total_progress = sum(s.progress_percent for s in steps)
        progress_percent = total_progress // len(steps)
    else:
        progress_percent = 0
    
    return ProjectTimelineResponse(
        project=project_summary,
        timeline_id=timeline_id,
        steps=steps,
        progress_percent=progress_percent
    )


@api_router.get("/projects/{project_id}/workflow/full", response_model=ProjectWorkflowResponse)
async def get_project_workflow_full(project_id: str, user: dict = Depends(get_current_agent)):
    """
    Canonical endpoint for AgentWorkflow.
    Returns: project, timeline, steps, activities, templates
    Single source of truth for all workflow data.
    """
    # Get project
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    project_summary = ProjectSummary(
        project_id=project['project_id'],
        name=project['name'],
        address=project.get('address'),
        status=project.get('status'),
        created_at=project.get('created_at')
    )
    
    # Get timeline and steps
    timeline = await db.project_timelines.find_one(
        {"project_id": project_id},
        {"_id": 0}
    )
    
    timeline_id = None
    steps = []
    
    if timeline:
        timeline_id = timeline.get('timeline_id') or timeline.get('project_timeline_id')
        
        steps_raw = await db.timeline_steps.find(
            {"project_timeline_id": timeline_id},
            {"_id": 0}
        ).sort("order_index", 1).to_list(50)
        
        steps = [
            TimelineStepSummary(
                step_id=s['step_id'],
                title=s['title'],
                description=s.get('description'),
                status=s.get('status', 'pending'),
                order_index=s.get('order_index', 0),
                planned_start=s.get('planned_start') or s.get('planned_date'),
                planned_end=s.get('planned_end'),
                progress_percent=s.get('progress_percent', 0)
            ) for s in steps_raw
        ]
    
    # Get recent activities for this project
    activities_raw = await db.activities.find(
        {"project_id": project_id},
        {"_id": 0}
    ).sort("created_at", -1).limit(20).to_list(20)
    
    # Get templates (agent-specific or global)
    templates_raw = await db.timeline_templates.find(
        {"$or": [{"agent_id": user['user_id']}, {"is_global": True}]},
        {"_id": 0}
    ).to_list(50)
    
    return ProjectWorkflowResponse(
        project=project_summary,
        timeline_id=timeline_id,
        steps=steps,
        activities=activities_raw,
        templates=templates_raw
    )


# ==================== ACTIVITY FEED ENDPOINTS (UNIFIED) ====================
# Single feed architecture - same data, role-based access filtering

ACTIVITY_UPLOAD_DIR = UPLOAD_DIR / "activities"
ACTIVITY_UPLOAD_DIR.mkdir(exist_ok=True)
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
ALLOWED_FILE_TYPES = {
    'image': ['image/jpeg', 'image/jpg', 'image/png', 'image/webp', 'image/gif'],
    'pdf': ['application/pdf']
}

async def get_buyer_unit_context(user: dict) -> dict:
    """Get the unit_id and project_id for a buyer based on their client link"""
    is_demo = user.get('is_demo', False)
    
    # Find the client linked to this buyer
    client = await db.clients.find_one(
        {"buyer_id": user['user_id']},
        {"_id": 0, "client_id": 1, "project_id": 1, "unit_reference": 1}
    )
    
    if not client:
        return None
    
    # Find the unit for this client
    unit = await db.project_units.find_one(
        {"client_id": client['client_id']},
        {"_id": 0, "unit_id": 1, "unit_reference": 1, "project_id": 1}
    )
    
    return {
        "client_id": client['client_id'],
        "project_id": client.get('project_id') or (unit.get('project_id') if unit else None),
        "unit_id": unit.get('unit_id') if unit else None,
        "unit_reference": unit.get('unit_reference') if unit else client.get('unit_reference')
    }

async def enrich_activity(activity: dict, include_replies: bool = False) -> dict:
    """Enrich activity with author name, project name, recipients, and optionally replies"""
    # Get author name
    author = await db.users.find_one({"user_id": activity['author_id']}, {"_id": 0, "name": 1})
    activity['author_name'] = author['name'] if author else 'Unknown'
    
    # Get project name
    project = await db.projects.find_one({"project_id": activity['project_id']}, {"_id": 0, "name": 1})
    activity['project_name'] = project['name'] if project else None
    
    # Get unit reference
    if activity.get('unit_id'):
        unit = await db.project_units.find_one({"unit_id": activity['unit_id']}, {"_id": 0, "unit_reference": 1})
        activity['unit_reference'] = unit['unit_reference'] if unit else None
    
    # Get recipients
    recipients = await db.activity_recipients.find(
        {"activity_id": activity['activity_id']},
        {"_id": 0}
    ).to_list(100)
    
    for recipient in recipients:
        client = await db.clients.find_one({"client_id": recipient['client_id']}, {"_id": 0, "name": 1})
        recipient['client_name'] = client['name'] if client else 'Unknown'
    
    activity['recipients'] = recipients
    
    # Get reply count
    reply_count = await db.activity_replies.count_documents({"activity_id": activity['activity_id']})
    activity['reply_count'] = reply_count
    
    # Get replies if requested
    if include_replies:
        replies = await db.activity_replies.find(
            {"activity_id": activity['activity_id']},
            {"_id": 0}
        ).sort("created_at", 1).to_list(100)
        
        for reply in replies:
            author = await db.users.find_one({"user_id": reply['author_id']}, {"_id": 0, "name": 1})
            reply['author_name'] = author['name'] if author else 'Unknown'
        
        activity['replies'] = replies
    
    return activity

@api_router.post("/activities")
async def create_activity(
    type: str = Form(...),
    project_id: str = Form(...),
    client_ids: str = Form(...),  # Comma-separated
    title: Optional[str] = Form(None),
    content: Optional[str] = Form(None),
    unit_id: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    user: dict = Depends(get_current_agent)
):
    """
    Create a new activity (AGENT ONLY).
    Single feed architecture - activities are stored once and filtered by role.
    """
    is_demo = user.get('is_demo', False)
    now = datetime.now(timezone.utc).isoformat()
    
    # Validate type
    if type not in ["message", "image", "pdf", "status"]:
        raise HTTPException(status_code=400, detail="Invalid activity type")
    
    # Parse client_ids
    recipient_ids = [cid.strip() for cid in client_ids.split(',') if cid.strip()]
    if not recipient_ids:
        raise HTTPException(status_code=400, detail="At least one client_id is required")
    
    # Validate project ownership
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Validate clients belong to agent
    for client_id in recipient_ids:
        client = await db.clients.find_one(
            {"client_id": client_id, "agent_id": user['user_id']},
            {"_id": 0}
        )
        if not client:
            raise HTTPException(status_code=400, detail=f"Client {client_id} not found or not owned by you")
    
    # Validate content or file is provided
    if type in ["message", "status"] and not content:
        raise HTTPException(status_code=400, detail="Content is required for message/status activities")
    
    if type in ["image", "pdf"] and not file:
        raise HTTPException(status_code=400, detail=f"File is required for {type} activities")
    
    # Handle file upload
    file_url = None
    file_name = None
    file_size = None
    
    if file:
        # Validate file size
        file_content = await file.read()
        file_size = len(file_content)
        
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail=f"File too large. Max size is {MAX_FILE_SIZE // (1024*1024)}MB")
        
        # Validate file type
        content_type = file.content_type or 'application/octet-stream'
        if type == "image" and content_type not in ALLOWED_FILE_TYPES['image']:
            raise HTTPException(status_code=400, detail="Invalid image type. Allowed: JPEG, PNG, WebP, GIF")
        if type == "pdf" and content_type not in ALLOWED_FILE_TYPES['pdf']:
            raise HTTPException(status_code=400, detail="Invalid file type. Only PDF allowed")
        
        # Save file
        file_id = uuid.uuid4().hex[:12]
        ext = Path(file.filename).suffix if file.filename else ('.jpg' if type == 'image' else '.pdf')
        filename = f"{file_id}_{file.filename or f'upload{ext}'}"
        file_path = ACTIVITY_UPLOAD_DIR / filename
        
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        file_url = f"/api/activities/files/{filename}"
        file_name = file.filename
    
    # Create activity
    activity_id = f"act_{uuid.uuid4().hex[:12]}"
    
    activity_doc = {
        "activity_id": activity_id,
        "type": type,
        "title": title,
        "content": content,
        "file_url": file_url,
        "file_name": file_name,
        "file_size": file_size,
        "author_id": user['user_id'],
        "author_role": "agent",
        "project_id": project_id,
        "unit_id": unit_id,
        "is_demo": is_demo,
        "created_at": now,
        "updated_at": now
    }
    
    await db.activities.insert_one(activity_doc)
    
    # Create recipient links and send notifications
    for client_id in recipient_ids:
        await db.activity_recipients.insert_one({
            "recipient_id": f"rcpt_{uuid.uuid4().hex[:8]}",
            "activity_id": activity_id,
            "client_id": client_id,
            "is_demo": is_demo,
            "created_at": now
        })
        
        # Send email notification to buyer if they have an account
        client = await db.clients.find_one({"client_id": client_id}, {"_id": 0})
        if client and client.get('buyer_id'):
            buyer = await db.users.find_one({"user_id": client['buyer_id']}, {"_id": 0, "email": 1, "name": 1})
            if buyer and buyer.get('email'):
                # Send email notification (non-blocking)
                try:
                    agent = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0, "name": 1, "contact_email": 1, "phone": 1, "company_name": 1})
                    email_data = {
                        "buyer_name": buyer.get('name', client.get('name', 'there')),
                        "agent_name": agent.get('name', 'Your Agent') if agent else 'Your Agent',
                        "agent_email": agent.get('contact_email', '') if agent else '',
                        "agent_phone": agent.get('phone', '') if agent else '',
                        "company_name": agent.get('company_name', '') if agent else '',
                        "project_name": project.get('name', 'Your Project'),
                        "message_preview": content[:200] if content else f"New {type} update",
                        "link": "buyer/dashboard"
                    }
                    await send_notification_email("feed_update", buyer['email'], email_data)
                    logger.info(f"Sent feed notification email to {buyer['email']}")
                except Exception as e:
                    logger.error(f"Failed to send feed notification email: {e}")
                
                # Create in-app notification
                try:
                    await create_notification(
                        user_id=client['buyer_id'],
                        title="New Update from Your Agent",
                        message=content[:100] if content else f"New {type} posted",
                        notification_type="feed_update",
                        link="/buyer/dashboard",
                        metadata={"activity_id": activity_id, "project_id": project_id},
                        is_demo=is_demo
                    )
                    
                    # Send real-time notification
                    await notify_realtime(
                        [client['buyer_id']],
                        "new_activity",
                        {"activity_id": activity_id, "type": type, "project_id": project_id}
                    )
                except Exception as e:
                    logger.error(f"Failed to create in-app notification: {e}")
    
    # Fetch and enrich the created activity
    activity = await db.activities.find_one({"activity_id": activity_id}, {"_id": 0})
    activity = await enrich_activity(activity)
    
    return activity

@api_router.get("/activities")
async def get_activities(
    project_id: Optional[str] = None,
    client_id: Optional[str] = None,
    unit_id: Optional[str] = None,
    type: Optional[str] = None,
    limit: int = 20,
    offset: int = 0,
    user: dict = Depends(get_current_user)
):
    """
    Get activities - ownership scoped by role.
    
    - Agent: Can filter by project_id, client_id, unit_id (ownership verified)
    - Buyer: Automatically filtered to their unit/client context only
    """
    if user['role'] == 'buyer':
        # STRICT BUYER ACCESS - auto-filter to their context
        buyer_context = await get_buyer_unit_context(user)
        
        if not buyer_context:
            # Buyer has no linked client - return empty
            return {"activities": [], "total": 0, "limit": limit, "offset": offset}
        
        # Find activities where this client is a recipient
        recipient_records = await db.activity_recipients.find(
            {"client_id": buyer_context['client_id']},
            {"_id": 0, "activity_id": 1}
        ).to_list(1000)
        
        activity_ids = [r['activity_id'] for r in recipient_records]
        
        if not activity_ids:
            return {"activities": [], "total": 0, "limit": limit, "offset": offset}
        
        query = {"activity_id": {"$in": activity_ids}}
        
    else:
        # AGENT ACCESS - can filter freely (ownership enforced below)
        query = {}
        
        # Verify agent owns the project if filtering by project
        if project_id:
            project = await db.projects.find_one(
                {"project_id": project_id, "agent_id": user['user_id']}
            )
            if not project:
                raise HTTPException(status_code=404, detail="Project not found")
            query["project_id"] = project_id
        
        # Filter by client_id (via recipients)
        if client_id:
            client = await db.clients.find_one(
                {"client_id": client_id, "agent_id": user['user_id']}
            )
            if not client:
                raise HTTPException(status_code=404, detail="Client not found")
            
            recipient_records = await db.activity_recipients.find(
                {"client_id": client_id},
                {"_id": 0, "activity_id": 1}
            ).to_list(1000)
            
            activity_ids = [r['activity_id'] for r in recipient_records]
            query["activity_id"] = {"$in": activity_ids} if activity_ids else {"$in": []}
        
        if unit_id:
            query["unit_id"] = unit_id
        
        # If no filters, show agent's activities (where they are author)
        # This includes both sent activities and drafts
        if not project_id and not client_id and not unit_id:
            query["author_id"] = user['user_id']
    
    if type:
        query["type"] = type
    
    # Get total count
    total = await db.activities.count_documents(query)
    
    # Get activities with pagination (reverse chronological)
    activities = await db.activities.find(
        query, {"_id": 0}
    ).sort("created_at", -1).skip(offset).limit(limit).to_list(limit)
    
    # Enrich each activity
    enriched = []
    for activity in activities:
        enriched.append(await enrich_activity(activity))
    
    return {
        "activities": enriched,
        "total": total,
        "limit": limit,
        "offset": offset
    }

@api_router.post("/activities/mark-seen")
async def mark_activities_seen(user: dict = Depends(get_current_user)):
    """Mark activities as seen - updates last_seen_at timestamp"""
    now = datetime.now(timezone.utc).isoformat()
    
    await db.user_activity_tracking.update_one(
        {"user_id": user['user_id']},
        {"$set": {"last_seen_at": now}},
        upsert=True
    )
    
    return {"last_seen_at": now}

@api_router.get("/activities/unread-count")
async def get_unread_count(user: dict = Depends(get_current_user)):
    """Get count of activities newer than last_seen_at"""
    is_demo = user.get('is_demo', False)
    
    # Get last seen timestamp
    tracking = await db.user_activity_tracking.find_one(
        {"user_id": user['user_id']},
        {"_id": 0}
    )
    last_seen_at = tracking.get('last_seen_at') if tracking else None
    
    if user['role'] == 'buyer':
        # Get buyer's activities
        buyer_context = await get_buyer_unit_context(user)
        if not buyer_context:
            return {"unread_count": 0, "last_seen_at": last_seen_at}
        
        recipient_records = await db.activity_recipients.find(
            {"client_id": buyer_context['client_id']},
            {"_id": 0, "activity_id": 1}
        ).to_list(1000)
        
        activity_ids = [r['activity_id'] for r in recipient_records]
        if not activity_ids:
            return {"unread_count": 0, "last_seen_at": last_seen_at}
        
        query = {"activity_id": {"$in": activity_ids}}
    else:
        # Agent - count their authored activities
        query = {"author_id": user['user_id']}
    
    # Add time filter if last_seen exists
    if last_seen_at:
        query["created_at"] = {"$gt": last_seen_at}
    
    unread_count = await db.activities.count_documents(query)
    
    return {"unread_count": unread_count, "last_seen_at": last_seen_at}

# File serving routes - must be before {activity_id} catch-all
@api_router.get("/activities/files/demo/{filename}")
async def get_demo_activity_file(filename: str, user: dict = Depends(get_current_user)):
    """Serve demo activity file attachments"""
    file_path = ACTIVITY_UPLOAD_DIR / "demo" / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    suffix = file_path.suffix.lower()
    content_type_map = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
        '.webp': 'image/webp', '.gif': 'image/gif', '.pdf': 'application/pdf',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xls': 'application/vnd.ms-excel',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.ppt': 'application/vnd.ms-powerpoint',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
    }
    content_type = content_type_map.get(suffix, 'application/octet-stream')
    
    return FileResponse(file_path, media_type=content_type, filename=filename)

@api_router.get("/activities/files/{filename}")
async def get_activity_file(filename: str, user: dict = Depends(get_current_user)):
    """Serve activity file attachments with access control"""
    file_path = ACTIVITY_UPLOAD_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    suffix = file_path.suffix.lower()
    content_type_map = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
        '.webp': 'image/webp', '.gif': 'image/gif', '.pdf': 'application/pdf',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xls': 'application/vnd.ms-excel',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.ppt': 'application/vnd.ms-powerpoint',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
    }
    content_type = content_type_map.get(suffix, 'application/octet-stream')
    
    return FileResponse(file_path, media_type=content_type, filename=filename)

@api_router.get("/activities/{activity_id}")
async def get_activity(activity_id: str, user: dict = Depends(get_current_user)):
    """
    Get a single activity with full details including replies.
    Role-based access is enforced.
    """
    is_demo = user.get('is_demo', False)
    
    activity = await db.activities.find_one(
        {"activity_id": activity_id},
        {"_id": 0}
    )
    
    if not activity:
        raise HTTPException(status_code=404, detail="Activity not found")
    
    if user['role'] == 'buyer':
        # Verify buyer has access to this activity
        buyer_context = await get_buyer_unit_context(user)
        
        if not buyer_context:
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Check if buyer is a recipient
        recipient = await db.activity_recipients.find_one({
            "activity_id": activity_id,
            "client_id": buyer_context['client_id'],
            "is_demo": is_demo
        })
        
        if not recipient:
            raise HTTPException(status_code=403, detail="Access denied")
    
    else:
        # Agent must own the project or be the author
        if activity['author_id'] != user['user_id']:
            project = await db.projects.find_one({
                "project_id": activity['project_id'],
                "agent_id": user['user_id'],
                "is_demo": is_demo
            })
            if not project:
                raise HTTPException(status_code=403, detail="Access denied")
    
    # Enrich with replies
    activity = await enrich_activity(activity, include_replies=True)
    
    return activity

@api_router.post("/activities/{activity_id}/send")
async def send_draft_activity(
    activity_id: str,
    user: dict = Depends(get_current_agent)
):
    """
    Send a draft activity to recipients.
    This converts a draft to a sent activity and creates recipient records.
    """
    is_demo = user.get('is_demo', False)
    now = datetime.now(timezone.utc).isoformat()
    
    # Find the draft activity
    activity = await db.activities.find_one(
        {"activity_id": activity_id, "author_id": user['user_id']},
        {"_id": 0}
    )
    
    if not activity:
        raise HTTPException(status_code=404, detail="Activity not found")
    
    if not activity.get('is_draft'):
        raise HTTPException(status_code=400, detail="Activity is not a draft")
    
    draft_recipients = activity.get('draft_recipients', [])
    if not draft_recipients:
        raise HTTPException(status_code=400, detail="No recipients specified for this draft")
    
    # Validate recipients
    for client_id in draft_recipients:
        client = await db.clients.find_one(
            {"client_id": client_id, "agent_id": user['user_id']},
            {"_id": 0}
        )
        if not client:
            raise HTTPException(status_code=400, detail=f"Client {client_id} not found")
    
    # Get project for notifications
    project = await db.projects.find_one(
        {"project_id": activity['project_id']},
        {"_id": 0, "name": 1}
    )
    
    # Update activity to mark as sent
    await db.activities.update_one(
        {"activity_id": activity_id},
        {"$set": {"is_draft": False, "updated_at": now}, "$unset": {"draft_recipients": ""}}
    )
    
    # Create recipient links and send notifications
    for client_id in draft_recipients:
        await db.activity_recipients.insert_one({
            "recipient_id": f"rcpt_{uuid.uuid4().hex[:8]}",
            "activity_id": activity_id,
            "client_id": client_id,
            "is_demo": is_demo,
            "created_at": now
        })
        
        # Send email notification to buyer if they have an account
        client = await db.clients.find_one({"client_id": client_id}, {"_id": 0})
        if client and client.get('buyer_id'):
            buyer = await db.users.find_one({"user_id": client['buyer_id']}, {"_id": 0, "email": 1, "name": 1})
            if buyer and buyer.get('email'):
                try:
                    agent = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0, "name": 1, "contact_email": 1, "phone": 1, "company_name": 1})
                    email_data = {
                        "buyer_name": buyer.get('name', client.get('name', 'there')),
                        "agent_name": agent.get('name', 'Your Agent') if agent else 'Your Agent',
                        "agent_email": agent.get('contact_email', '') if agent else '',
                        "agent_phone": agent.get('phone', '') if agent else '',
                        "company_name": agent.get('company_name', '') if agent else '',
                        "project_name": project.get('name', 'Your Project') if project else 'Your Project',
                        "message_preview": activity.get('content', '')[:200] if activity.get('content') else "New update",
                        "link": "buyer/dashboard"
                    }
                    await send_notification_email("feed_update", buyer['email'], email_data)
                    logger.info(f"Sent feed notification email to {buyer['email']}")
                except Exception as e:
                    logger.error(f"Failed to send feed notification email: {e}")
    
    # Return the updated activity
    updated_activity = await db.activities.find_one({"activity_id": activity_id}, {"_id": 0})
    return await enrich_activity(updated_activity)

@api_router.post("/activities/{activity_id}/reply")
async def reply_to_activity(
    activity_id: str,
    reply_data: ActivityReplyCreate,
    user: dict = Depends(get_current_user)
):
    """
    Reply to an activity. Both buyers and agents can reply.
    Access is strictly validated.
    """
    is_demo = user.get('is_demo', False)
    now = datetime.now(timezone.utc).isoformat()
    
    activity = await db.activities.find_one(
        {"activity_id": activity_id},
        {"_id": 0}
    )
    
    if not activity:
        raise HTTPException(status_code=404, detail="Activity not found")
    
    if user['role'] == 'buyer':
        # Verify buyer has access
        buyer_context = await get_buyer_unit_context(user)
        
        if not buyer_context:
            raise HTTPException(status_code=403, detail="Access denied")
        
        recipient = await db.activity_recipients.find_one({
            "activity_id": activity_id,
            "client_id": buyer_context['client_id'],
            "is_demo": is_demo
        })
        
        if not recipient:
            raise HTTPException(status_code=403, detail="Access denied")
    
    else:
        # Agent must own the project
        project = await db.projects.find_one({
            "project_id": activity['project_id'],
            "agent_id": user['user_id'],
            "is_demo": is_demo
        })
        if not project:
            raise HTTPException(status_code=403, detail="Access denied")
    
    # Create reply
    reply_id = f"reply_{uuid.uuid4().hex[:12]}"
    
    reply_doc = {
        "reply_id": reply_id,
        "activity_id": activity_id,
        "author_id": user['user_id'],
        "author_role": user['role'],
        "content": reply_data.content,
        "is_demo": is_demo,
        "created_at": now
    }
    
    await db.activity_replies.insert_one(reply_doc)
    
    # Update activity timestamp
    await db.activities.update_one(
        {"activity_id": activity_id},
        {"$set": {"updated_at": now}}
    )
    
    # Get author name
    author = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0, "name": 1})
    reply_doc['author_name'] = author['name'] if author else 'Unknown'
    
    # Remove MongoDB _id and is_demo before returning
    reply_doc.pop('_id', None)
    reply_doc.pop('is_demo', None)
    
    return reply_doc

@api_router.post("/activities/{activity_id}/request-change")
async def request_activity_change(
    activity_id: str,
    user: dict = Depends(get_current_user)
):
    """
    Request changes on a file-type activity. Buyer only.
    Creates a system reply and updates activity status.
    """
    is_demo = user.get('is_demo', False)
    
    if user['role'] != 'buyer':
        raise HTTPException(status_code=403, detail="Only buyers can request changes")
    
    activity = await db.activities.find_one(
        {"activity_id": activity_id},
        {"_id": 0}
    )
    
    if not activity:
        raise HTTPException(status_code=404, detail="Activity not found")
    
    if activity['type'] != 'file':
        raise HTTPException(status_code=400, detail="Change requests only supported for file-type activities")
    
    # Verify buyer has access
    buyer_context = await get_buyer_unit_context(user)
    if not buyer_context:
        raise HTTPException(status_code=403, detail="Access denied")
    
    recipient = await db.activity_recipients.find_one({
        "activity_id": activity_id,
        "client_id": buyer_context['client_id'],
        "is_demo": is_demo
    })
    
    if not recipient:
        raise HTTPException(status_code=403, detail="Access denied")
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Update activity status
    await db.activities.update_one(
        {"activity_id": activity_id},
        {"$set": {"status": "change_requested", "updated_at": now}}
    )
    
    return {"message": "Change requested", "status": "change_requested"}

@api_router.delete("/activities/{activity_id}")
async def delete_activity(
    activity_id: str,
    user: dict = Depends(get_current_agent)
):
    """
    Delete an activity (agent only).
    Deletes the activity, its recipients, replies, and associated files.
    """
    is_demo = user.get('is_demo', False)
    
    # Find the activity (author_id is the field used for activities)
    activity = await db.activities.find_one(
        {"activity_id": activity_id, "author_id": user['user_id']},
        {"_id": 0}
    )
    
    if not activity:
        raise HTTPException(status_code=404, detail="Activity not found")
    
    # Delete associated file if it exists
    if activity.get('file_path') and os.path.exists(activity['file_path']):
        try:
            os.remove(activity['file_path'])
        except Exception as e:
            logger.warning(f"Failed to delete activity file: {e}")
    
    # Delete replies
    await db.activity_replies.delete_many({"activity_id": activity_id})
    
    # Delete recipients
    await db.activity_recipients.delete_many({"activity_id": activity_id})
    
    # Delete the activity itself
    await db.activities.delete_one({"activity_id": activity_id})
    
    return {"message": "Activity deleted successfully"}

class ActivityUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None

@api_router.put("/activities/{activity_id}")
async def update_activity(
    activity_id: str,
    data: ActivityUpdate,
    user: dict = Depends(get_current_agent)
):
    """
    Update an activity (agent only).
    Can update title and content.
    """
    is_demo = user.get('is_demo', False)
    
    # Find the activity (author_id is the field used for activities)
    activity = await db.activities.find_one(
        {"activity_id": activity_id, "author_id": user['user_id']},
        {"_id": 0}
    )
    
    if not activity:
        raise HTTPException(status_code=404, detail="Activity not found")
    
    # Build update dict
    update_dict = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if data.title is not None:
        update_dict["title"] = data.title
    if data.content is not None:
        update_dict["content"] = data.content
    
    # Update the activity
    await db.activities.update_one(
        {"activity_id": activity_id},
        {"$set": update_dict}
    )
    
    # Return updated activity
    updated = await db.activities.find_one(
        {"activity_id": activity_id},
        {"_id": 0}
    )
    updated.pop('is_demo', None)
    
    return updated

# ==================== TIMELINE/WORKFLOW ENDPOINTS ====================

# Valid status transitions
TIMELINE_STATUS_TRANSITIONS = {
    "pending": ["in_progress", "completed", "approved"],  # Can skip ahead
    "in_progress": ["pending", "completed", "approved"],  # Can go back or skip ahead
    "completed": ["pending", "in_progress", "approved"],  # Can go back or forward
    "approved": ["pending", "in_progress", "completed"]  # Can reopen if needed
}

# ==================== AI TIMELINE EXTRACTION ====================

async def extract_timeline_from_file(file_path: str, filename: str, mime_type: str) -> dict:
    """Extract timeline/phases from a planning document using AI (OpenAI)"""
    if not OPENAI_API_KEY:
        logger.warning("OPENAI_API_KEY not set, returning fallback")
        return {
            "phases": [],
            "project_duration": None,
            "notes": "AI extraction unavailable - please enter timeline manually",
            "confidence": "low",
            "extraction_failed": True
        }
    
    try:
        # Initialize OpenAI client
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        system_prompt = """You are a construction project timeline extraction assistant.

Your task is to analyze planning documents (PDFs, spreadsheets, images) and extract a structured project timeline.

EXTRACTION RULES:

1. PHASES/MILESTONES:
   - Extract distinct project phases or milestones
   - Each phase should have a name and brief description
   - Order phases chronologically
   - Look for: preparation, demolition, construction, installation, finishing, inspection, handover

2. DATES - PRESERVE EXACT WORDING:
   - Copy the date text EXACTLY as written in the document
   - Examples: "March 2026", "Q4 2027", "3 Mar 2026", "Mid-February", "Week 12"
   - Do NOT convert or calculate dates
   - Do NOT use YYYY-MM-DD format unless that's what the document says
   - If no date found, use null

3. KEY DELIVERABLES:
   - What is expected to be completed at each phase

Return ONLY a JSON object with this structure:
{
  "phases": [
    {
      "name": "Phase name",
      "order": 1,
      "planned_date": "exact date text from document or null",
      "description": "Brief description of work in this phase",
      "deliverables": ["key deliverable 1", "key deliverable 2"]
    }
  ],
  "notes": "any important notes about the timeline",
  "confidence": "high/medium/low"
}"""
        
        # Convert file to images for OpenAI vision
        import base64
        image_contents = []
        
        # Check if it's a PDF
        if mime_type == "application/pdf" or file_path.lower().endswith('.pdf'):
            # Convert PDF to images using PyMuPDF
            doc = fitz.open(file_path)
            for page_num in range(min(3, len(doc))):
                page = doc[page_num]
                mat = fitz.Matrix(150/72, 150/72)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                image_contents.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{img_base64}"}
                })
            doc.close()
        else:
            # For images, use directly
            with open(file_path, "rb") as f:
                file_base64 = base64.b64encode(f.read()).decode('utf-8')
            # Map common mime types
            img_mime = mime_type if mime_type in ['image/png', 'image/jpeg', 'image/gif', 'image/webp'] else 'image/png'
            image_contents.append({
                "type": "image_url",
                "image_url": {"url": f"data:{img_mime};base64,{file_base64}"}
            })
        
        # Build message
        message_content = [
            {"type": "text", "text": "Extract the project timeline/phases from this planning document. Identify all phases, milestones, and their durations or dates. Return only valid JSON."}
        ] + image_contents
        
        # Use GPT-4o with vision for document analysis
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": message_content}
            ],
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content
        
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            extracted = json.loads(json_match.group())
        else:
            raise ValueError("No JSON found in response")
        
        # Validate and clean phases - use text-based planned_date
        valid_phases = []
        for idx, phase in enumerate(extracted.get("phases", [])):
            if isinstance(phase, dict) and phase.get("name"):
                valid_phases.append({
                    "name": str(phase.get("name", "")),
                    "order": phase.get("order", idx + 1),
                    "planned_date": phase.get("planned_date") or "",  # Text-based date
                    "description": str(phase.get("description", "")),
                    "deliverables": phase.get("deliverables", [])
                })
        
        return {
            "phases": valid_phases,
            "notes": extracted.get("notes", ""),
            "confidence": extracted.get("confidence", "medium"),
            "extraction_failed": False
        }
        
    except Exception as e:
        logger.error(f"Timeline extraction failed: {str(e)}")
        return {
            "phases": [],
            "notes": f"Extraction error: {str(e)}",
            "confidence": "low",
            "extraction_failed": True
        }

class ManualTimelineCreate(BaseModel):
    project_id: str
    name: str = "Project Timeline"
    steps: List[dict]

@api_router.post("/timeline/create")
async def create_manual_timeline(data: ManualTimelineCreate, user: dict = Depends(get_current_agent)):
    """Create a timeline manually with custom steps"""
    is_demo = user.get('is_demo', False)
    
    # Verify project access
    project = await db.projects.find_one({
        "project_id": data.project_id,
        "agent_id": user['user_id'],
        "is_demo": is_demo
    }, {"_id": 0})
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Check if timeline already exists
    existing = await db.project_timelines.find_one({
        "project_id": data.project_id,
        "is_demo": is_demo
    }, {"_id": 0})
    
    if existing:
        raise HTTPException(status_code=400, detail="Timeline already exists for this project. Delete the existing one first.")
    
    now = datetime.now(timezone.utc).isoformat()
    timeline_id = str(uuid.uuid4())
    
    # Create timeline
    timeline = {
        "timeline_id": timeline_id,
        "project_id": data.project_id,
        "agent_id": user['user_id'],
        "name": data.name,
        "status": "published",
        "is_demo": is_demo,
        "created_at": now,
        "updated_at": now
    }
    await db.project_timelines.insert_one(timeline)
    
    # Create steps
    created_steps = []
    for i, step_data in enumerate(data.steps):
        step = {
            "step_id": str(uuid.uuid4()),
            "project_timeline_id": timeline_id,
            "title": step_data.get('title', f'Step {i+1}'),
            "description": step_data.get('description', ''),
            "planned_date": step_data.get('planned_date'),
            "status": step_data.get('status', 'pending'),
            "order": step_data.get('order', i + 1),
            "is_demo": is_demo,
            "created_at": now
        }
        await db.timeline_steps.insert_one(step)
        created_steps.append({k: v for k, v in step.items() if k != '_id'})
    
    return {
        "message": "Timeline created successfully",
        "timeline": {k: v for k, v in timeline.items() if k != '_id'},
        "steps": created_steps
    }

@api_router.post("/timeline/extract")
async def extract_timeline(
    file: UploadFile = File(...),
    project_id: Optional[str] = Form(None),
    user: dict = Depends(get_current_agent)
):
    """Upload a planning document and extract timeline using AI"""
    is_demo = user.get('is_demo', False)
    
    # Validate file type
    allowed_types = [
        'application/pdf',
        'image/jpeg', 'image/jpg', 'image/png', 'image/webp',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # xlsx
        'application/vnd.ms-excel',  # xls
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Unsupported file type. Allowed: PDF, images, Excel files")
    
    # Read and save file temporarily
    content = await file.read()
    
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File must be less than 20MB")
    
    extraction_id = f"timeline_ext_{uuid.uuid4().hex[:12]}"
    file_ext = file.filename.split('.')[-1] if '.' in file.filename else 'bin'
    temp_filename = f"{extraction_id}.{file_ext}"
    temp_path = UPLOAD_DIR / temp_filename
    
    try:
        with open(temp_path, "wb") as f:
            f.write(content)
        
        # Extract timeline using AI
        result = await extract_timeline_from_file(str(temp_path), file.filename, file.content_type)
        
        # Store the extraction result for later approval
        extraction_doc = {
            "extraction_id": extraction_id,
            "agent_id": user['user_id'],
            "project_id": project_id,
            "original_filename": file.filename,
            "file_path": str(temp_path),
            "extracted_data": result,
            "status": "pending_review",  # pending_review, approved, rejected
            "is_demo": is_demo,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.timeline_extractions.insert_one(extraction_doc)
        extraction_doc.pop('_id', None)
        
        return {
            "extraction_id": extraction_id,
            "status": "pending_review",
            "extracted_data": result
        }
        
    except Exception as e:
        # Clean up temp file on error
        if temp_path.exists():
            temp_path.unlink()
        raise HTTPException(status_code=500, detail=f"Timeline extraction failed: {str(e)}")

@api_router.get("/timeline/extractions")
async def get_timeline_extractions(
    status: Optional[str] = None,
    project_id: Optional[str] = None,
    user: dict = Depends(get_current_agent)
):
    """Get all timeline extractions for the agent"""
    is_demo = user.get('is_demo', False)
    
    query = {"agent_id": user['user_id']}
    if status:
        query["status"] = status
    if project_id:
        query["project_id"] = project_id
    
    extractions = await db.timeline_extractions.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    return extractions

@api_router.get("/timeline/extractions/{extraction_id}")
async def get_timeline_extraction(extraction_id: str, user: dict = Depends(get_current_agent)):
    """Get a specific timeline extraction"""
    is_demo = user.get('is_demo', False)
    
    extraction = await db.timeline_extractions.find_one(
        {"extraction_id": extraction_id, "agent_id": user['user_id']},
        {"_id": 0}
    )
    
    if not extraction:
        raise HTTPException(status_code=404, detail="Extraction not found")
    
    return extraction

@api_router.post("/timeline/extractions/{extraction_id}/approve")
async def approve_timeline_extraction(
    extraction_id: str,
    project_id: str = Form(...),
    phases: str = Form(...),  # JSON string of edited phases
    user: dict = Depends(get_current_agent)
):
    """Approve an extracted timeline and apply it to a project"""
    is_demo = user.get('is_demo', False)
    
    extraction = await db.timeline_extractions.find_one(
        {"extraction_id": extraction_id, "agent_id": user['user_id']}
    )
    
    if not extraction:
        raise HTTPException(status_code=404, detail="Extraction not found")
    
    # Validate project exists
    project = await db.projects.find_one(
        {"project_id": project_id, "agent_id": user['user_id']}
    )
    
    if not project:
        raise HTTPException(status_code=400, detail="Invalid project")
    
    # Parse phases
    try:
        phases_list = json.loads(phases)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid phases format")
    
    # Create a timeline template from the approved extraction
    template_id = f"template_{uuid.uuid4().hex[:12]}"
    template_doc = {
        "template_id": template_id,
        "agent_id": user['user_id'],
        "name": f"Extracted Timeline - {extraction.get('original_filename', 'Unknown')}",
        "is_demo": is_demo,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "source": "ai_extraction",
        "extraction_id": extraction_id
    }
    
    await db.timeline_templates.insert_one(template_doc)
    
    # Create steps from phases
    for idx, phase in enumerate(phases_list):
        step_id = f"step_{uuid.uuid4().hex[:8]}"
        step_doc = {
            "step_id": step_id,
            "template_id": template_id,
            "name": phase.get("name", f"Phase {idx + 1}"),
            "description": phase.get("description", ""),
            "order_index": idx,
            "planned_date": phase.get("planned_date", ""),  # Text-based date
            "deliverables": phase.get("deliverables", []),
            "is_demo": is_demo
        }
        await db.timeline_template_steps.insert_one(step_doc)
    
    # Apply template to project - use timeline_id to match what get_project_timeline expects
    timeline_id = f"ptl_{uuid.uuid4().hex[:12]}"
    project_timeline_doc = {
        "timeline_id": timeline_id,  # This is what get_project_timeline uses
        "project_id": project_id,
        "agent_id": user['user_id'],
        "template_id": template_id,
        "name": template_doc['name'],
        "status": "active",
        "is_demo": is_demo,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.project_timelines.insert_one(project_timeline_doc)
    
    # Create timeline steps - use db.timeline_steps to match what get_project_timeline reads
    for idx, phase in enumerate(phases_list):
        step_id = f"pts_{uuid.uuid4().hex[:8]}"
        step_doc = {
            "step_id": step_id,
            "project_timeline_id": timeline_id,  # References the timeline
            "project_id": project_id,
            "agent_id": user['user_id'],
            "title": phase.get("name", f"Phase {idx + 1}"),  # 'title' matches existing schema
            "description": phase.get("description", ""),
            "order_index": idx,
            "status": "pending",
            "planned_date": phase.get("planned_date", ""),  # Text-based date
            "is_demo": is_demo,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.timeline_steps.insert_one(step_doc)  # Use timeline_steps collection
    
    # Update extraction status
    await db.timeline_extractions.update_one(
        {"extraction_id": extraction_id},
        {"$set": {
            "status": "approved",
            "approved_at": datetime.now(timezone.utc).isoformat(),
            "applied_to_project": project_id,
            "template_id": template_id
        }}
    )
    
    # Clean up temp file
    if extraction.get('file_path') and Path(extraction['file_path']).exists():
        try:
            Path(extraction['file_path']).unlink()
        except OSError:
            pass
    
    return {
        "message": "Timeline approved and applied to project",
        "template_id": template_id,
        "timeline_id": timeline_id,
        "phases_count": len(phases_list)
    }

@api_router.delete("/timeline/extractions/{extraction_id}")
async def delete_timeline_extraction(extraction_id: str, user: dict = Depends(get_current_agent)):
    """Delete a pending timeline extraction"""
    is_demo = user.get('is_demo', False)
    
    extraction = await db.timeline_extractions.find_one(
        {"extraction_id": extraction_id, "agent_id": user['user_id']}
    )
    
    if not extraction:
        raise HTTPException(status_code=404, detail="Extraction not found")
    
    # Clean up temp file
    if extraction.get('file_path') and Path(extraction['file_path']).exists():
        try:
            Path(extraction['file_path']).unlink()
        except OSError:
            pass
    
    await db.timeline_extractions.delete_one({"extraction_id": extraction_id})
    
    return {"message": "Extraction deleted"}

@api_router.get("/timeline/templates")
async def get_timeline_templates(user: dict = Depends(get_current_agent)):
    """Get all timeline templates for agent"""
    is_demo = user.get('is_demo', False)
    
    templates = await db.timeline_templates.find(
        {"$or": [{"agent_id": user['user_id']}, {"is_demo": is_demo}]},
        {"_id": 0}
    ).to_list(100)
    
    # Enrich with steps
    for template in templates:
        steps = await db.timeline_template_steps.find(
            {"template_id": template['template_id']},
            {"_id": 0}
        ).sort("order_index", 1).to_list(50)
        template['steps'] = steps
    
    return templates

@api_router.post("/timeline/templates")
async def create_timeline_template(data: TimelineTemplateCreate, user: dict = Depends(get_current_agent)):
    """Create a new timeline template"""
    is_demo = user.get('is_demo', False)
    template_id = f"tmpl_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    template_doc = {
        "template_id": template_id,
        "agent_id": user['user_id'],
        "name": data.name,
        "is_demo": is_demo,
        "created_at": now
    }
    await db.timeline_templates.insert_one(template_doc)
    
    # Create template steps
    for step in data.steps:
        step_id = f"tmpl_step_{uuid.uuid4().hex[:12]}"
        step_doc = {
            "step_id": step_id,
            "template_id": template_id,
            "title": step.title,
            "description": step.description,
            "order_index": step.order_index
        }
        await db.timeline_template_steps.insert_one(step_doc)
    
    # Return with steps
    template_doc.pop('_id', None)
    steps = await db.timeline_template_steps.find(
        {"template_id": template_id},
        {"_id": 0}
    ).sort("order_index", 1).to_list(50)
    template_doc['steps'] = steps
    
    return template_doc

@api_router.delete("/timeline/templates/{template_id}")
async def delete_timeline_template(template_id: str, user: dict = Depends(get_current_agent)):
    """Delete a timeline template"""
    is_demo = user.get('is_demo', False)
    
    template = await db.timeline_templates.find_one({
        "template_id": template_id,
        "agent_id": user['user_id'],
        "is_demo": is_demo
    })
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    await db.timeline_template_steps.delete_many({"template_id": template_id})
    await db.timeline_templates.delete_one({"template_id": template_id})
    
    return {"message": "Template deleted"}

@api_router.post("/timeline/templates/{template_id}/apply")
async def apply_timeline_template(template_id: str, project_id: str, user: dict = Depends(get_current_agent)):
    """Apply a template to create project timeline"""
    is_demo = user.get('is_demo', False)
    
    # Verify template exists
    template = await db.timeline_templates.find_one({
        "template_id": template_id,
        "is_demo": is_demo
    }, {"_id": 0})
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Verify project exists and agent owns it
    project = await db.projects.find_one({
        "project_id": project_id,
        "agent_id": user['user_id'],
        "is_demo": is_demo
    }, {"_id": 0})
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Check if project already has a timeline
    existing = await db.project_timelines.find_one({
        "project_id": project_id,
        "is_demo": is_demo
    })
    
    if existing:
        raise HTTPException(status_code=400, detail="Project already has a timeline. Delete it first.")
    
    now = datetime.now(timezone.utc).isoformat()
    timeline_id = f"timeline_{uuid.uuid4().hex[:12]}"
    
    # Create project timeline
    timeline_doc = {
        "timeline_id": timeline_id,
        "project_id": project_id,
        "template_id": template_id,
        "is_demo": is_demo,
        "created_at": now
    }
    await db.project_timelines.insert_one(timeline_doc)
    
    # Get template steps and create timeline steps
    template_steps = await db.timeline_template_steps.find(
        {"template_id": template_id},
        {"_id": 0}
    ).sort("order_index", 1).to_list(50)
    
    for tmpl_step in template_steps:
        step_id = f"step_{uuid.uuid4().hex[:12]}"
        step_doc = {
            "step_id": step_id,
            "project_timeline_id": timeline_id,
            "title": tmpl_step['title'],
            "description": tmpl_step.get('description'),
            "status": "pending",
            "order_index": tmpl_step['order_index'],
            "planned_date": None,
            "completed_at": None,
            "is_demo": is_demo,
            "created_at": now,
            "updated_at": now
        }
        await db.timeline_steps.insert_one(step_doc)
    
    return {"message": "Timeline created", "timeline_id": timeline_id}

@api_router.get("/project-timeline")
async def get_project_timeline(
    project_id: Optional[str] = None,
    user: dict = Depends(get_current_user)
):
    """
    Get project construction timeline (workflow stages).
    - Agent: can query any project they own via project_id
    - Buyer: gets timeline for their project automatically
    """
    is_demo = user.get('is_demo', False)
    
    if user['role'] == 'buyer':
        # Get buyer's project
        client = await db.clients.find_one(
            {"buyer_id": user['user_id']},
            {"_id": 0}
        )
        if not client:
            return {"timeline": None, "steps": []}
        project_id = client['project_id']
    else:
        if not project_id:
            raise HTTPException(status_code=400, detail="project_id required for agents")
        
        # Verify agent owns project
        project = await db.projects.find_one({
            "project_id": project_id,
            "agent_id": user['user_id'],
            "is_demo": is_demo
        })
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
    
    # Get project timeline
    timeline = await db.project_timelines.find_one(
        {"project_id": project_id},
        {"_id": 0}
    )
    
    if not timeline:
        return {"timeline": None, "steps": []}
    
    # Handle missing timeline_id (legacy data)
    tl_id = timeline.get('timeline_id') or timeline.get('project_timeline_id')
    if not tl_id:
        return {"timeline": timeline, "steps": []}
    
    # Enrich with template name
    if timeline.get('template_id'):
        template = await db.timeline_templates.find_one(
            {"template_id": timeline['template_id']},
            {"_id": 0, "name": 1}
        )
        timeline['template_name'] = template['name'] if template else None
    
    # Get steps with documents and notes
    steps = await db.timeline_steps.find(
        {"project_timeline_id": tl_id},
        {"_id": 0}
    ).sort("order_index", 1).to_list(100)
    
    # Enrich steps with linked documents and notes
    for step in steps:
        # Get linked documents (activities)
        doc_links = await db.timeline_step_documents.find(
            {"timeline_step_id": step['step_id']},
            {"_id": 0}
        ).to_list(50)
        
        documents = []
        for link in doc_links:
            activity = await db.activities.find_one(
                {"activity_id": link['activity_id']},
                {"_id": 0, "activity_id": 1, "title": 1, "type": 1, "file_name": 1, "file_url": 1, "created_at": 1}
            )
            if activity:
                documents.append(activity)
        step['documents'] = documents
        
        # Get internal notes (agent only)
        if user['role'] == 'agent':
            notes = await db.timeline_step_internal_notes.find(
                {"timeline_step_id": step['step_id']},
                {"_id": 0}
            ).sort("created_at", -1).to_list(50)
            
            # Enrich with author names
            for note in notes:
                author = await db.users.find_one(
                    {"user_id": note['author_id']},
                    {"_id": 0, "name": 1}
                )
                note['author_name'] = author['name'] if author else 'Unknown'
            step['internal_notes'] = notes
        else:
            step['internal_notes'] = []  # Buyer doesn't see notes
    
    timeline['steps'] = steps
    return {"timeline": timeline, "steps": steps}

@api_router.patch("/timeline/steps/{step_id}")
async def update_timeline_step(step_id: str, data: TimelineStepUpdate, user: dict = Depends(get_current_agent)):
    """Update a timeline step (agent only)"""
    is_demo = user.get('is_demo', False)
    
    # Find step and verify ownership
    step = await db.timeline_steps.find_one(
        {"step_id": step_id},
        {"_id": 0}
    )
    
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    
    # Verify agent owns the project via timeline
    timeline = await db.project_timelines.find_one({
        "$or": [
            {"timeline_id": step['project_timeline_id']},
            {"project_timeline_id": step['project_timeline_id']}
        ],
        "is_demo": is_demo
    }, {"_id": 0})
    
    if not timeline:
        raise HTTPException(status_code=404, detail="Timeline not found")
    
    project = await db.projects.find_one(
        {"project_id": timeline['project_id'], "agent_id": user['user_id']}
    )
    
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    update_data = {}
    now = datetime.now(timezone.utc).isoformat()
    
    if data.title is not None:
        update_data['title'] = data.title
    if data.description is not None:
        update_data['description'] = data.description
    if data.planned_date is not None:
        update_data['planned_date'] = data.planned_date
    if data.order_index is not None:
        update_data['order_index'] = data.order_index
    # Extended fields (from project_stages migration)
    if data.planned_start is not None:
        update_data['planned_start'] = data.planned_start
    if data.planned_end is not None:
        update_data['planned_end'] = data.planned_end
    if data.actual_start is not None:
        update_data['actual_start'] = data.actual_start
    if data.progress_percent is not None:
        update_data['progress_percent'] = data.progress_percent
    if data.notes is not None:
        update_data['notes'] = data.notes
    
    # Handle status transition
    if data.status is not None:
        current_status = step['status']
        new_status = data.status
        
        if new_status != current_status:
            allowed = TIMELINE_STATUS_TRANSITIONS.get(current_status, [])
            if new_status not in allowed:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid status transition from '{current_status}' to '{new_status}'"
                )
            
            update_data['status'] = new_status
            
            # Set completed_at when transitioning to completed
            if new_status == 'completed':
                update_data['completed_at'] = now
            elif new_status in ['pending', 'in_progress']:
                update_data['completed_at'] = None
    
    if update_data:
        update_data['updated_at'] = now
        await db.timeline_steps.update_one(
            {"step_id": step_id},
            {"$set": update_data}
        )
        
        # Send milestone notification when status changes to completed
        if update_data.get('status') == 'completed':
            await send_milestone_notification(
                step=step,
                project=project,
                timeline=timeline,
                user=user,
                is_demo=is_demo
            )
    
    updated = await db.timeline_steps.find_one({"step_id": step_id}, {"_id": 0})
    return updated

@api_router.post("/timeline/{timeline_id}/steps")
async def add_timeline_step(timeline_id: str, data: dict, user: dict = Depends(get_current_agent)):
    """Add a new step to an existing timeline"""
    is_demo = user.get('is_demo', False)
    
    # Find timeline (support both timeline_id and project_timeline_id)
    timeline = await db.project_timelines.find_one({
        "$or": [
            {"timeline_id": timeline_id},
            {"project_timeline_id": timeline_id}
        ],
        "is_demo": is_demo
    }, {"_id": 0})
    
    if not timeline:
        raise HTTPException(status_code=404, detail="Timeline not found")
    
    # Verify agent owns the project
    project = await db.projects.find_one(
        {"project_id": timeline['project_id'], "agent_id": user['user_id']}
    )
    
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Get current max order_index
    pipeline = [
        {"$match": {"project_timeline_id": timeline.get('timeline_id') or timeline.get('project_timeline_id')}},
        {"$group": {"_id": None, "max_order": {"$max": "$order_index"}}}
    ]
    result = await db.timeline_steps.aggregate(pipeline).to_list(1)
    max_order = result[0]['max_order'] if result and result[0].get('max_order') is not None else 0
    
    now = datetime.now(timezone.utc).isoformat()
    timeline_id_ref = timeline.get('timeline_id') or timeline.get('project_timeline_id')
    
    step_doc = {
        "step_id": f"step_{uuid.uuid4().hex[:12]}",
        "project_timeline_id": timeline_id_ref,
        "project_id": timeline['project_id'],
        "title": data.get('title', 'New Step'),
        "description": data.get('description', ''),
        "planned_date": data.get('planned_date'),
        "status": "pending",
        "order_index": max_order + 1,
        "documents": [],
        "internal_notes": [],
        "is_demo": is_demo,
        "created_at": now,
        "updated_at": now,
        "completed_at": None
    }
    
    await db.timeline_steps.insert_one(step_doc)
    return await db.timeline_steps.find_one({"step_id": step_doc['step_id']}, {"_id": 0})

@api_router.delete("/timeline/steps/{step_id}")
async def delete_timeline_step(step_id: str, user: dict = Depends(get_current_agent)):
    """Delete a timeline step"""
    is_demo = user.get('is_demo', False)
    
    # Find step
    step = await db.timeline_steps.find_one(
        {"step_id": step_id},
        {"_id": 0}
    )
    
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    
    # Verify agent owns the project via timeline
    timeline = await db.project_timelines.find_one({
        "$or": [
            {"timeline_id": step['project_timeline_id']},
            {"project_timeline_id": step['project_timeline_id']}
        ],
        "is_demo": is_demo
    }, {"_id": 0})
    
    if not timeline:
        raise HTTPException(status_code=404, detail="Timeline not found")
    
    project = await db.projects.find_one(
        {"project_id": timeline['project_id'], "agent_id": user['user_id']}
    )
    
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Delete the step
    await db.timeline_steps.delete_one({"step_id": step_id})
    
    return {"message": "Step deleted", "step_id": step_id}

@api_router.post("/timeline/steps/{step_id}/documents")
async def link_document_to_step(step_id: str, data: TimelineStepDocumentCreate, user: dict = Depends(get_current_agent)):
    """Link an existing activity to a timeline step"""
    is_demo = user.get('is_demo', False)
    
    # Verify step exists and agent has access
    step = await db.timeline_steps.find_one(
        {"step_id": step_id},
        {"_id": 0}
    )
    
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    
    # Verify agent owns the project
    timeline = await db.project_timelines.find_one(
        {"timeline_id": step['project_timeline_id']},
        {"_id": 0}
    )
    
    project = await db.projects.find_one(
        {"project_id": timeline['project_id'], "agent_id": user['user_id']}
    )
    
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Verify activity exists
    activity = await db.activities.find_one(
        {"activity_id": data.activity_id},
        {"_id": 0}
    )
    
    if not activity:
        raise HTTPException(status_code=404, detail="Activity not found")
    
    # Check if already linked
    existing = await db.timeline_step_documents.find_one({
        "timeline_step_id": step_id,
        "activity_id": data.activity_id
    })
    
    if existing:
        raise HTTPException(status_code=400, detail="Document already linked")
    
    now = datetime.now(timezone.utc).isoformat()
    link_id = f"link_{uuid.uuid4().hex[:12]}"
    
    link_doc = {
        "link_id": link_id,
        "timeline_step_id": step_id,
        "activity_id": data.activity_id,
        "created_at": now
    }
    await db.timeline_step_documents.insert_one(link_doc)
    
    return {"message": "Document linked", "link_id": link_id}

@api_router.delete("/timeline/steps/{step_id}/documents/{activity_id}")
async def unlink_document_from_step(step_id: str, activity_id: str, user: dict = Depends(get_current_agent)):
    """Remove a document link from timeline step"""
    is_demo = user.get('is_demo', False)
    
    # Verify ownership chain
    step = await db.timeline_steps.find_one({"step_id": step_id})
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    
    timeline = await db.project_timelines.find_one({"timeline_id": step['project_timeline_id']})
    project = await db.projects.find_one({"project_id": timeline['project_id'], "agent_id": user['user_id']})
    
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    result = await db.timeline_step_documents.delete_one({
        "timeline_step_id": step_id,
        "activity_id": activity_id
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Link not found")
    
    return {"message": "Document unlinked"}

@api_router.post("/timeline/steps/{step_id}/notes")
async def add_internal_note(step_id: str, data: TimelineStepNoteCreate, user: dict = Depends(get_current_agent)):
    """Add an internal note to a timeline step (agent only)"""
    is_demo = user.get('is_demo', False)
    
    # Verify ownership
    step = await db.timeline_steps.find_one({"step_id": step_id})
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    
    timeline = await db.project_timelines.find_one({"timeline_id": step['project_timeline_id']})
    project = await db.projects.find_one({"project_id": timeline['project_id'], "agent_id": user['user_id']})
    
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    now = datetime.now(timezone.utc).isoformat()
    note_id = f"note_{uuid.uuid4().hex[:12]}"
    
    note_doc = {
        "note_id": note_id,
        "timeline_step_id": step_id,
        "author_id": user['user_id'],
        "content": data.content,
        "created_at": now
    }
    await db.timeline_step_internal_notes.insert_one(note_doc)
    
    # Enrich with author name
    note_doc['author_name'] = user.get('name', 'Unknown')
    note_doc.pop('_id', None)
    
    return note_doc

@api_router.delete("/timeline/{timeline_id}")
async def delete_project_timeline(timeline_id: str, user: dict = Depends(get_current_agent)):
    """Delete a project timeline and all its steps"""
    is_demo = user.get('is_demo', False)
    
    # Try to find by timeline_id or project_timeline_id (for legacy data)
    timeline = await db.project_timelines.find_one({
        "$or": [
            {"timeline_id": timeline_id},
            {"project_timeline_id": timeline_id}
        ],
        "is_demo": is_demo
    })
    
    if not timeline:
        raise HTTPException(status_code=404, detail="Timeline not found")
    
    # Get the actual timeline ID from whichever field exists
    actual_timeline_id = timeline.get('timeline_id') or timeline.get('project_timeline_id')
    
    # Verify agent owns the project
    project = await db.projects.find_one({
        "project_id": timeline['project_id'],
        "agent_id": user['user_id'],
        "is_demo": is_demo
    })
    
    if not project:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Delete all related data - try both ID formats
    steps = await db.timeline_steps.find({
        "$or": [
            {"project_timeline_id": actual_timeline_id},
            {"project_timeline_id": timeline_id}
        ]
    }).to_list(100)
    step_ids = [s['step_id'] for s in steps]
    
    if step_ids:
        await db.timeline_step_documents.delete_many({"timeline_step_id": {"$in": step_ids}})
        await db.timeline_step_internal_notes.delete_many({"timeline_step_id": {"$in": step_ids}})
    
    # Delete steps using both possible ID formats
    await db.timeline_steps.delete_many({
        "$or": [
            {"project_timeline_id": actual_timeline_id},
            {"project_timeline_id": timeline_id}
        ]
    })
    
    # Delete the timeline document
    await db.project_timelines.delete_one({"_id": timeline['_id']} if '_id' in timeline else {
        "$or": [
            {"timeline_id": timeline_id},
            {"project_timeline_id": timeline_id}
        ]
    })
    
    return {"message": "Timeline deleted"}

# ==================== DASHBOARD STATS ====================

@api_router.get("/stats/agent")
async def get_agent_stats(user: dict = Depends(get_current_agent)):
    """Get agent dashboard stats"""
    is_demo = user.get('is_demo', False)
    agent_id = user['user_id']
    
    total_clients = await db.clients.count_documents({"agent_id": agent_id})
    
    # Pending quotes (Sent or Change Requested)
    pending_quotes = await db.documents.count_documents({
        "agent_id": agent_id, "type": "quote",
        "status": {"$in": ["Sent", "Change Requested"]}
    })
    
    # Pending invoices (Sent status)
    pending_invoices = await db.documents.count_documents({
        "agent_id": agent_id, "type": "invoice",
        "status": "Sent"
    })
    
    # Revenue (paid invoices)
    paid_invoices = await db.documents.find(
        {"agent_id": agent_id, "type": "invoice", "status": "Paid"},
        {"_id": 0, "amount": 1}
    ).to_list(1000)
    total_revenue = sum(inv['amount'] for inv in paid_invoices)
    
    # Recent documents
    recent_docs = await db.documents.find(
        {"agent_id": agent_id},
        {"_id": 0}
    ).sort("updated_at", -1).limit(5).to_list(5)
    
    # Change requests
    change_requests = await db.documents.find(
        {"agent_id": agent_id, "type": "quote", "status": "Change Requested"},
        {"_id": 0}
    ).to_list(10)
    
    # Approved quotes ready to convert
    approved_quotes = await db.documents.find(
        {"agent_id": agent_id, "type": "quote", "status": "Approved"},
        {"_id": 0}
    ).to_list(10)
    
    return {
        "total_clients": total_clients,
        "pending_quotes": pending_quotes,
        "pending_invoices": pending_invoices,
        "total_revenue": total_revenue,
        "recent_documents": recent_docs,
        "change_requests": change_requests,
        "approved_quotes": approved_quotes
    }

@api_router.get("/stats/buyer")
async def get_buyer_stats(user: dict = Depends(get_current_buyer)):
    """Get buyer dashboard stats"""
    is_demo = user.get('is_demo', False)
    
    clients = await db.clients.find(
        {"buyer_id": user['user_id']},
        {"_id": 0}
    ).to_list(100)
    client_ids = [c['client_id'] for c in clients]
    
    if not client_ids:
        return {
            "pending_quotes": 0,
            "pending_invoices": 0,
            "total_paid": 0,
            "projects": []
        }
    
    pending_quotes = await db.documents.count_documents({
        "client_id": {"$in": client_ids}, "is_demo": is_demo,
        "type": "quote", "status": "Sent"
    })
    
    pending_invoices = await db.documents.count_documents({
        "client_id": {"$in": client_ids}, "is_demo": is_demo,
        "type": "invoice", "status": "Sent"
    })
    
    paid_invoices = await db.documents.find(
        {"client_id": {"$in": client_ids}, "type": "invoice", "status": "Paid"},
        {"_id": 0, "amount": 1}
    ).to_list(1000)
    total_paid = sum(inv['amount'] for inv in paid_invoices)
    
    project_ids = list(set(c['project_id'] for c in clients))
    projects = await db.projects.find({"project_id": {"$in": project_ids}}, {"_id": 0}).to_list(10)
    
    for project in projects:
        client = next((c for c in clients if c['project_id'] == project['project_id']), None)
        if client:
            project['unit_reference'] = client.get('unit_reference', '')
    
    return {
        "pending_quotes": pending_quotes,
        "pending_invoices": pending_invoices,
        "total_paid": total_paid,
        "projects": projects
    }

# ==================== DOCUMENT VAULT ENDPOINTS ====================

# Vault categories for organization
VAULT_CATEGORIES = ["Contracts", "Plans", "Permits", "Reports", "Other"]

class VaultDocumentCreate(BaseModel):
    name: str
    category: str = "Other"
    project_id: Optional[str] = None
    description: Optional[str] = None
    access_level: str = "private"  # private, shared (with specific buyers)
    shared_with_clients: Optional[List[str]] = None  # List of client_ids

class VaultDocumentUpdate(BaseModel):
    name: Optional[str] = None
    category: Optional[str] = None
    project_id: Optional[str] = None
    description: Optional[str] = None
    access_level: Optional[str] = None
    shared_with_clients: Optional[List[str]] = None  # List of client_ids

# Document types for vault
VAULT_DOC_TYPES = ["general", "action_required"]

@api_router.post("/vault/upload")
async def upload_vault_document(
    request: Request,
    file: UploadFile = File(...),
    name: str = Form(...),
    category: str = Form("Other"),
    project_id: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    access_level: str = Form("private"),
    shared_with_clients: Optional[str] = Form(None),  # JSON string of client_ids
    doc_type: str = Form("general"),  # 'general' or 'action_required'
    user: dict = Depends(get_current_agent)
):
    """Upload a document to the vault"""
    # Rate limit: 20 uploads per minute
    rate_limit_check(request, "file_upload")
    
    is_demo = user.get('is_demo', False)
    
    # Validate category
    if category not in VAULT_CATEGORIES:
        category = "Other"
    
    # Validate access level
    if access_level not in ["private", "shared"]:
        access_level = "private"
    
    # Validate doc_type
    if doc_type not in VAULT_DOC_TYPES:
        doc_type = "general"
    
    # Validate file type
    allowed_types = [
        'application/pdf',
        'image/jpeg', 'image/jpg', 'image/png', 'image/webp',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # xlsx
        'application/vnd.ms-excel',  # xls
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # docx
        'application/msword'  # doc
    ]
    
    if file.content_type not in allowed_types:
        # Build user-friendly error message
        allowed_formats = "PDF, JPG, PNG, WEBP, Excel (XLSX/XLS), Word (DOCX/DOC)"
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed formats: {allowed_formats}"
        )
    
    # Read file content
    content = await file.read()
    
    # Check file size (max 20MB)
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File must be less than 20MB")
    
    # Save file
    vault_id = f"vault_{uuid.uuid4().hex[:12]}"
    file_ext = file.filename.split('.')[-1] if '.' in file.filename else 'bin'
    filename = f"{vault_id}_{secure_filename(file.filename)}"
    file_path = UPLOAD_DIR / filename
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    file_url = f"/api/uploads/{filename}"
    
    # Parse shared_with_clients JSON
    shared_clients_list = []
    if shared_with_clients:
        try:
            shared_clients_list = json.loads(shared_with_clients)
            if not isinstance(shared_clients_list, list):
                shared_clients_list = []
        except json.JSONDecodeError:
            shared_clients_list = []
    
    # Create vault document record
    vault_doc = {
        "vault_id": vault_id,
        "agent_id": user['user_id'],
        "name": name,
        "original_filename": file.filename,
        "file_url": file_url,
        "file_path": str(file_path),
        "file_type": file.content_type,
        "file_size": len(content),
        "category": category,
        "project_id": project_id,
        "description": description or "",
        "access_level": access_level,
        "shared_with_clients": shared_clients_list if access_level == "shared" else [],
        "doc_type": doc_type,
        "is_demo": is_demo,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.vault_documents.insert_one(vault_doc)
    
    # Send notifications to shared clients
    if access_level == "shared" and shared_clients_list:
        agent_settings = user.get('settings', {})
        agent_profile = user.get('profile', {})
        agent_name = agent_profile.get('display_name') or user.get('name', 'Your agent')
        company_name = agent_settings.get('company_name', '')
        
        for client_id in shared_clients_list:
            client = await db.clients.find_one({"client_id": client_id}, {"_id": 0})
            if client and client.get('buyer_id'):
                # Create in-app notification
                notification_title = "Action Required" if doc_type == "action_required" else "New Document"
                await create_notification(
                    user_id=client['buyer_id'],
                    title=notification_title,
                    message=f"{agent_name} shared a document: {name}",
                    notification_type="vault_document",
                    link="/buyer/vault",
                    is_demo=is_demo
                )
                
                # Send email notification
                if client.get('email'):
                    await send_notification_email(
                        "document_sent",
                        client['email'],
                        {
                            "doc_type": "document",
                            "buyer_name": client.get('name', 'there'),
                            "agent_name": agent_name,
                            "company_name": company_name,
                            "title": name,
                            "summary": description or ("This document requires your attention" if doc_type == "action_required" else "A new document has been shared with you"),
                            "currency": "CHF",
                            "amount": 0,
                            "project_name": "",
                            "unit_reference": client.get('unit_reference', '')
                        }
                    )
    
    # Remove MongoDB _id before returning
    vault_doc.pop('_id', None)
    
    return vault_doc

@api_router.get("/vault", response_model=List[VaultDocumentResponse])
async def get_vault_documents(
    category: Optional[str] = None,
    project_id: Optional[str] = None,
    user: dict = Depends(get_current_agent)
):
    """Get all vault documents for the agent, optionally filtered - ownership scoped"""
    query = {"agent_id": user['user_id']}
    
    if category and category in VAULT_CATEGORIES:
        query["category"] = category
    
    if project_id:
        query["project_id"] = project_id
    
    docs = await db.vault_documents.find(query, {"_id": 0}).sort("created_at", -1).to_list(500)
    
    return docs


@api_router.get("/vault/buyer", response_model=List[VaultDocumentResponse])
async def get_buyer_vault_documents_v2(user: dict = Depends(get_current_user)):
    """Get vault documents shared with the buyer based on their client_id - ownership scoped"""
    if user['role'] != 'buyer':
        raise HTTPException(status_code=403, detail="Buyers only")
    
    # Get the buyer's client record (includes client_id and agent_id)
    client = await db.clients.find_one(
        {"buyer_id": user['user_id']},
        {"_id": 0, "client_id": 1, "agent_id": 1}
    )
    
    if not client:
        return []
    
    client_id = client.get('client_id')
    agent_id = client.get('agent_id')
    
    if not client_id or not agent_id:
        return []
    
    # Find shared documents where this client is in shared_with_clients array
    # Note: Using agent_id from client relationship for ownership scoping
    query = {
        "agent_id": agent_id,
        "access_level": "shared",
        "shared_with_clients": client_id
    }
    
    docs = await db.vault_documents.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    
    return docs


@api_router.get("/vault/categories/list")
async def get_vault_categories(user: dict = Depends(get_current_agent)):
    """Get available vault categories"""
    return VAULT_CATEGORIES


@api_router.get("/vault/{vault_id}", response_model=VaultDocumentResponse)
async def get_vault_document(vault_id: str, user: dict = Depends(get_current_user)):
    """
    Get a specific vault document - ownership scoped.
    - Agents can access their own documents
    - Buyers can access shared documents from their agent
    """
    if user['role'] == 'agent':
        # Agent: must own the document
        doc = await db.vault_documents.find_one(
            {"vault_id": vault_id, "agent_id": user['user_id']},
            {"_id": 0}
        )
    elif user['role'] == 'buyer':
        # Buyer: must have sharing access via client relationship
        client = await db.clients.find_one(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1, "agent_id": 1}
        )
        
        if not client:
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Document must belong to buyer's agent and be shared with them
        doc = await db.vault_documents.find_one(
            {
                "vault_id": vault_id,
                "agent_id": client.get('agent_id'),
                "access_level": "shared",
                "shared_with_clients": client.get('client_id')
            },
            {"_id": 0}
        )
    else:
        raise HTTPException(status_code=403, detail="Access denied")
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return doc

@api_router.put("/vault/{vault_id}", response_model=VaultDocumentResponse)
async def update_vault_document(
    vault_id: str,
    data: VaultDocumentUpdate,
    user: dict = Depends(get_current_agent)
):
    """Update vault document metadata - ownership scoped"""
    
    doc = await db.vault_documents.find_one(
        {"vault_id": vault_id, "agent_id": user['user_id']}
    )
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
    
    if data.name is not None:
        update_data["name"] = data.name
    if data.category is not None and data.category in VAULT_CATEGORIES:
        update_data["category"] = data.category
    if data.project_id is not None:
        update_data["project_id"] = data.project_id
    if data.description is not None:
        update_data["description"] = data.description
    if data.access_level is not None and data.access_level in ["private", "shared"]:
        update_data["access_level"] = data.access_level
        # If switching to private, clear shared_with_clients
        if data.access_level == "private":
            update_data["shared_with_clients"] = []
    if data.shared_with_clients is not None:
        update_data["shared_with_clients"] = data.shared_with_clients
    
    await db.vault_documents.update_one(
        {"vault_id": vault_id},
        {"$set": update_data}
    )
    
    return {"message": "Document updated"}

@api_router.delete("/vault/{vault_id}")
async def delete_vault_document(vault_id: str, user: dict = Depends(get_current_agent)):
    """Delete a vault document"""
    is_demo = user.get('is_demo', False)
    
    doc = await db.vault_documents.find_one(
        {"vault_id": vault_id, "agent_id": user['user_id']}
    )
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete physical file
    if doc.get('file_path'):
        try:
            Path(doc['file_path']).unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"Failed to delete vault file: {e}")
    
    await db.vault_documents.delete_one({"vault_id": vault_id})
    
    return {"message": "Document deleted"}

@api_router.get("/vault/{vault_id}/download")
async def download_vault_document(vault_id: str, user: dict = Depends(get_current_user)):
    """Download a vault document - accessible by agent owner or buyers with sharing access"""
    is_demo = user.get('is_demo', False)
    
    # Get the document
    doc = await db.vault_documents.find_one(
        {"vault_id": vault_id},
        {"_id": 0}
    )
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check access permissions
    if user['role'] == 'agent':
        # Agent must own the document
        if doc.get('agent_id') != user['user_id']:
            raise HTTPException(status_code=403, detail="Access denied")
    elif user['role'] == 'buyer':
        # Get buyer's client record to verify access
        client = await db.clients.find_one(
            {"buyer_id": user['user_id']},
            {"_id": 0, "client_id": 1, "agent_id": 1}
        )
        
        if not client:
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Document must belong to buyer's agent
        if doc.get('agent_id') != client.get('agent_id'):
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Check if document is shared with this buyer's client_id
        client_id = client.get('client_id')
        if not client_id or client_id not in doc.get('shared_with_clients', []):
            raise HTTPException(status_code=403, detail="Document not shared with you")
    else:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Get file path and serve
    file_path = doc.get('file_path')
    if not file_path or not Path(file_path).exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        file_path,
        filename=doc.get('filename', 'document'),
        media_type='application/octet-stream'
    )

# ==================== ANALYTICS ENDPOINTS ====================

@api_router.get("/analytics")
async def get_analytics(period: str = "month", user: dict = Depends(get_current_agent)):
    """
    Get comprehensive analytics for the agent dashboard.
    Period can be: week, month, quarter, year, all
    """
    is_demo = user.get('is_demo', False)
    agent_id = user['user_id']
    
    # Calculate date filter based on period
    now = datetime.now(timezone.utc)
    date_filter = None
    
    if period == "week":
        date_filter = (now - timedelta(days=7)).isoformat()
    elif period == "month":
        date_filter = (now - timedelta(days=30)).isoformat()
    elif period == "quarter":
        date_filter = (now - timedelta(days=90)).isoformat()
    elif period == "year":
        date_filter = (now - timedelta(days=365)).isoformat()
    # "all" = no date filter
    
    # Build base query
    base_query = {"agent_id": agent_id}
    if date_filter:
        base_query["created_at"] = {"$gte": date_filter}
    
    # Count total documents by type
    total_quotes = await db.documents.count_documents({**base_query, "type": "quote"})
    total_invoices = await db.documents.count_documents({**base_query, "type": "invoice"})
    
    # Quote status breakdown
    quote_approved = await db.documents.count_documents({**base_query, "type": "quote", "status": "Approved"})
    quote_sent = await db.documents.count_documents({**base_query, "type": "quote", "status": "Sent"})
    quote_rejected = await db.documents.count_documents({**base_query, "type": "quote", "status": "Rejected"})
    quote_draft = await db.documents.count_documents({**base_query, "type": "quote", "status": "Draft"})
    quote_change_requested = await db.documents.count_documents({**base_query, "type": "quote", "status": "Change Requested"})
    
    # Invoice status breakdown
    invoice_paid = await db.documents.count_documents({**base_query, "type": "invoice", "status": "Paid"})
    invoice_sent = await db.documents.count_documents({**base_query, "type": "invoice", "status": "Sent"})
    invoice_draft = await db.documents.count_documents({**base_query, "type": "invoice", "status": "Draft"})
    
    # Revenue from paid invoices (no date filter for total revenue)
    revenue_query = {"agent_id": agent_id, "type": "invoice", "status": "Paid"}
    if date_filter:
        revenue_query["paid_date"] = {"$gte": date_filter}
    
    paid_invoices = await db.documents.find(
        revenue_query,
        {"_id": 0, "amount": 1}
    ).to_list(10000)
    total_revenue = sum(inv.get('amount', 0) for inv in paid_invoices)
    
    # Pending amount (invoices sent but not paid)
    pending_invoices = await db.documents.find(
        {**base_query, "type": "invoice", "status": "Sent"},
        {"_id": 0, "amount": 1}
    ).to_list(10000)
    pending_amount = sum(inv.get('amount', 0) for inv in pending_invoices)
    
    # Client count (all time, regardless of period)
    total_clients = await db.clients.count_documents({"agent_id": agent_id})
    
    # Project count (all time)
    total_projects = await db.projects.count_documents({"agent_id": agent_id})
    
    return {
        "totalQuotes": total_quotes,
        "totalInvoices": total_invoices,
        "totalClients": total_clients,
        "totalProjects": total_projects,
        "totalRevenue": total_revenue,
        "pendingAmount": pending_amount,
        "quoteStats": {
            "approved": quote_approved,
            "sent": quote_sent + quote_change_requested,  # Group pending states
            "rejected": quote_rejected,
            "draft": quote_draft
        },
        "invoiceStats": {
            "paid": invoice_paid,
            "sent": invoice_sent,
            "draft": invoice_draft
        }
    }



# ==================== TEST ENDPOINTS ====================

@api_router.get("/test/email-template/{template_type}")
async def test_email_template(template_type: str):
    """Test endpoint to preview email template HTML (for CTA button validation)"""
    valid_templates = ["document_sent", "quote_approved", "invoice_paid", "payment_confirmed", "project_update", "welcome", "milestone_completed", "change_requested", "new_message"]
    if template_type not in valid_templates:
        raise HTTPException(status_code=400, detail=f"Invalid template. Valid: {valid_templates}")
    
    # Sample data for testing
    test_data = {
        "buyer_name": "Test Buyer",
        "agent_name": "Test Agent",
        "company_name": "Test Company SA",
        "doc_type": "quote",
        "title": "Kitchen Renovation",
        "summary": "Complete kitchen remodel with premium appliances",
        "currency": "CHF",
        "amount": 15000.00,
        "project_name": "Alpine View Residence",
        "unit_reference": "Unit A-101",
        "document_number": "QT-2024-0001",
        "document_id": "test_doc_001",
        "link": "buyer/dashboard",
        # Milestone notification specific data
        "milestone_name": "Foundation Complete",
        "milestone_description": "The foundation has been poured and cured successfully.",
        "progress_percent": 35,
        # Change request specific data
        "comment": "Please review the pricing on item 3.",
        # New message specific data
        "sender_name": "Test Sender",
        "message_preview": "Hello, I wanted to discuss the project timeline..."
    }
    
    subject, html = get_email_template(template_type, test_data)
    
    # Validate CTA button has proper inline styles
    cta_checks = {
        "has_bg_color": "background-color: #2563EB" in html,
        "has_text_color": "color: #FFFFFF" in html,
        "has_inline_style_on_a_tag": 'style="display: inline-block; background-color: #2563EB; color: #FFFFFF' in html
    }
    
    return {
        "template_type": template_type,
        "subject": subject,
        "cta_validation": cta_checks,
        "cta_valid": all(cta_checks.values()),
        "html_preview": html[:2000] + "..." if len(html) > 2000 else html
    }

# ==================== DEMO DATA SEEDING ====================

@api_router.get("/demo/seed")
async def seed_demo_data_get():
    """GET version for easy browser access"""
    return await seed_demo_data()

@api_router.post("/demo/seed")
async def seed_demo_data():
    """Seed comprehensive demo data using unified documents collection"""
    # Clear existing demo data
    await db.users.delete_many({"is_demo": True})
    await db.projects.delete_many({"is_demo": True})
    await db.clients.delete_many({"is_demo": True})
    await db.documents.delete_many({"is_demo": True})
    await db.notifications.delete_many({"is_demo": True})
    await db.project_stages.delete_many({"is_demo": True})
    
    # Create demo agent
    demo_agent_id = "demo_agent_001"
    demo_agent = {
        "user_id": demo_agent_id,
        "email": "demo.agent@upgradeflow.com",
        "name": "Marc Dubois",
        "password_hash": bcrypt.hashpw("demo123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8'),
        "role": "agent",
        "picture": None,
        "is_demo": True,
        "created_at": datetime.now(timezone.utc).isoformat(),
        # Demo agent has Pro subscription for showcasing all features
        "subscription_plan": "pro",
        "subscription_status": "active",
        "settings": {
            "language": "en",
            "currency": "CHF",
            "company_name": "Dubois Immobilier"
        }
    }
    await db.users.insert_one(demo_agent)
    
    # Create demo logo for the agent using PIL
    try:
        from PIL import Image, ImageDraw
        logo_filename = f"logo_{demo_agent_id}_demo.png"
        logo_path = UPLOAD_DIR / logo_filename
        
        # Create a 200x200 image with transparency
        img = Image.new('RGBA', (200, 200), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        
        # Primary color (blue)
        primary = (37, 99, 235, 255)
        white = (255, 255, 255, 255)
        
        # Draw rounded rectangle helper
        def rounded_rectangle(draw, xy, radius, fill):
            x1, y1, x2, y2 = xy
            draw.rectangle([x1 + radius, y1, x2 - radius, y2], fill=fill)
            draw.rectangle([x1, y1 + radius, x2, y2 - radius], fill=fill)
            draw.ellipse([x1, y1, x1 + radius * 2, y1 + radius * 2], fill=fill)
            draw.ellipse([x2 - radius * 2, y1, x2, y1 + radius * 2], fill=fill)
            draw.ellipse([x1, y2 - radius * 2, x1 + radius * 2, y2], fill=fill)
            draw.ellipse([x2 - radius * 2, y2 - radius * 2, x2, y2], fill=fill)
        
        # Draw blue rounded background
        rounded_rectangle(draw, [10, 10, 190, 190], 24, primary)
        
        # Draw house icon (white)
        draw.rectangle([55, 100, 145, 165], fill=white)  # House body
        draw.polygon([(45, 100), (100, 50), (155, 100)], fill=white)  # Roof
        draw.rectangle([85, 125, 115, 165], fill=primary)  # Door
        draw.rectangle([62, 110, 78, 125], fill=primary)  # Window left
        draw.rectangle([122, 110, 138, 125], fill=primary)  # Window right
        
        # Save logo
        img.save(str(logo_path), 'PNG')
        
        # Update agent with logo URL
        await db.users.update_one(
            {"user_id": demo_agent_id},
            {"$set": {"company_logo_url": f"/api/uploads/{logo_filename}"}}
        )
    except Exception as e:
        print(f"Warning: Could not create demo logo: {e}")
    
    # Create demo buyers
    demo_buyer1_id = "demo_buyer_001"
    demo_buyer1 = {
        "user_id": demo_buyer1_id,
        "email": "sophie.mueller@example.com",
        "name": "Sophie Müller",
        "role": "buyer",
        "picture": None,
        "is_demo": True,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(demo_buyer1)
    
    demo_buyer2_id = "demo_buyer_002"
    demo_buyer2 = {
        "user_id": demo_buyer2_id,
        "email": "thomas.weber@example.com",
        "name": "Thomas Weber",
        "role": "buyer",
        "picture": None,
        "is_demo": True,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(demo_buyer2)
    
    # Create demo project
    demo_project_id = "demo_proj_001"
    demo_project = {
        "project_id": demo_project_id,
        "agent_id": demo_agent_id,
        "name": "Residenza Lago Vista",
        "address": "Via del Sole 15, 6900 Lugano, Switzerland",
        "description": "Luxury lakefront apartments with panoramic views of Lake Lugano.",
        "is_demo": True,
        "created_at": (datetime.now(timezone.utc) - timedelta(days=90)).isoformat()
    }
    await db.projects.insert_one(demo_project)
    
    # Create demo clients
    demo_client1_id = "demo_client_001"
    demo_client1 = {
        "client_id": demo_client1_id,
        "agent_id": demo_agent_id,
        "buyer_id": demo_buyer1_id,
        "name": "Sophie Müller",
        "email": "sophie.mueller@example.com",
        "phone": "+41 79 123 45 67",
        "project_id": demo_project_id,
        "unit_id": "demo_unit_001",  # Link to unit
        "unit_reference": "Unit A-301",
        "is_demo": True,
        "created_at": (datetime.now(timezone.utc) - timedelta(days=60)).isoformat()
    }
    await db.clients.insert_one(demo_client1)
    
    demo_client2_id = "demo_client_002"
    demo_client2 = {
        "client_id": demo_client2_id,
        "agent_id": demo_agent_id,
        "buyer_id": demo_buyer2_id,
        "name": "Thomas Weber",
        "email": "thomas.weber@example.com",
        "phone": "+41 78 987 65 43",
        "project_id": demo_project_id,
        "unit_id": "demo_unit_002",  # Link to unit
        "unit_reference": "Unit B-502",
        "is_demo": True,
        "created_at": (datetime.now(timezone.utc) - timedelta(days=45)).isoformat()
    }
    await db.clients.insert_one(demo_client2)
    
    now = datetime.now(timezone.utc)
    
    # Create demo DOCUMENTS (unified)
    demo_documents = [
        # Quote 1: Approved (ready to convert to invoice)
        {
            "document_id": "doc_demo_001",
            "document_number": "QT-2024-0001",
            "type": "quote",
            "status": "Approved",
            "agent_id": demo_agent_id,
            "client_id": demo_client1_id,
            "buyer_id": demo_buyer1_id,
            "project_id": demo_project_id,
            "unit_reference": "Unit A-301",
            "title": "Premium Kitchen Upgrade Package",
            "summary": "Transform your kitchen with Gaggenau appliances, custom oak cabinetry, and premium Silestone countertops.",
            "hero_image_url": None,
            "hero_image_path": None,
            "amount": 43500.00,
            "items": [
                {"description": "Gaggenau Oven Set", "quantity": 1, "unit_price": 12500.00, "total": 12500.00},
                {"description": "Custom Oak Cabinetry", "quantity": 1, "unit_price": 18000.00, "total": 18000.00},
                {"description": "Silestone Countertops", "quantity": 1, "unit_price": 8500.00, "total": 8500.00},
                {"description": "Installation & Labor", "quantity": 1, "unit_price": 4500.00, "total": 4500.00}
            ],
            "currency": "CHF",
            "supplier_name": "Kitchen Solutions AG",
            "notes": "Premium package includes 5-year warranty.",
            "change_request_comment": None,
            "pdf_filename": None,
            "pdf_path": None,
            "ai_extraction_confidence": None,
            "parent_document_id": None,
            "due_date": None,
            "paid_date": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=50)).isoformat(),
            "updated_at": (now - timedelta(days=20)).isoformat()
        },
        # Quote 2: Sent (awaiting buyer response) - for Sophie
        {
            "document_id": "doc_demo_002",
            "document_number": "QT-2024-0002",
            "type": "quote",
            "status": "Sent",
            "agent_id": demo_agent_id,
            "client_id": demo_client1_id,
            "buyer_id": demo_buyer1_id,
            "project_id": demo_project_id,
            "unit_reference": "Unit A-301",
            "title": "Smart Home Automation System",
            "summary": "Complete home automation with KNX controller, smart lighting, climate control, and integrated security.",
            "hero_image_url": None,
            "hero_image_path": None,
            "amount": 22800.00,
            "items": [
                {"description": "KNX Controller Hub", "quantity": 1, "unit_price": 3500.00, "total": 3500.00},
                {"description": "Smart Lighting System (12 zones)", "quantity": 1, "unit_price": 6800.00, "total": 6800.00},
                {"description": "Climate Control Integration", "quantity": 1, "unit_price": 4200.00, "total": 4200.00},
                {"description": "Security System Integration", "quantity": 1, "unit_price": 5500.00, "total": 5500.00},
                {"description": "Programming & Configuration", "quantity": 1, "unit_price": 2800.00, "total": 2800.00}
            ],
            "currency": "CHF",
            "supplier_name": "Smart Living GmbH",
            "notes": "System demonstration included upon completion.",
            "change_request_comment": None,
            "pdf_filename": None,
            "pdf_path": None,
            "ai_extraction_confidence": None,
            "parent_document_id": None,
            "due_date": None,
            "paid_date": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=5)).isoformat(),
            "updated_at": (now - timedelta(days=4)).isoformat()
        },
        # Quote 3: Sent (awaiting buyer response) - for Thomas
        {
            "document_id": "doc_demo_003",
            "document_number": "QT-2024-0003",
            "type": "quote",
            "status": "Sent",
            "agent_id": demo_agent_id,
            "client_id": demo_client2_id,
            "buyer_id": demo_buyer2_id,
            "project_id": demo_project_id,
            "unit_reference": "Unit B-502",
            "title": "Bathroom Wellness Upgrade",
            "summary": "Luxury spa experience with Hansgrohe rain shower, heated floors, and natural stone finishes.",
            "hero_image_url": None,
            "hero_image_path": None,
            "amount": 27800.00,
            "items": [
                {"description": "Hansgrohe Rain Shower System", "quantity": 1, "unit_price": 4800.00, "total": 4800.00},
                {"description": "Heated Floor System", "quantity": 1, "unit_price": 3200.00, "total": 3200.00},
                {"description": "Premium Duravit Fixtures", "quantity": 1, "unit_price": 6500.00, "total": 6500.00},
                {"description": "Natural Stone Tiles", "quantity": 1, "unit_price": 7800.00, "total": 7800.00},
                {"description": "Installation", "quantity": 1, "unit_price": 5500.00, "total": 5500.00}
            ],
            "currency": "CHF",
            "supplier_name": "Bathroom Design Studio",
            "notes": "Estimated completion: 3 weeks from approval.",
            "change_request_comment": None,
            "pdf_filename": None,
            "pdf_path": None,
            "ai_extraction_confidence": None,
            "parent_document_id": None,
            "due_date": None,
            "paid_date": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=3)).isoformat(),
            "updated_at": (now - timedelta(days=2)).isoformat()
        },
        # Quote 4: Change Requested - for Thomas
        {
            "document_id": "doc_demo_004",
            "document_number": "QT-2024-0004",
            "type": "quote",
            "status": "Change Requested",
            "agent_id": demo_agent_id,
            "client_id": demo_client2_id,
            "buyer_id": demo_buyer2_id,
            "project_id": demo_project_id,
            "unit_reference": "Unit B-502",
            "title": "Terrace Extension & Landscaping",
            "summary": "Expand your outdoor living with a 25 sqm terrace extension, outdoor kitchen, and professional landscaping.",
            "hero_image_url": None,
            "hero_image_path": None,
            "amount": 30000.00,
            "items": [
                {"description": "Terrace Extension (25 sqm)", "quantity": 1, "unit_price": 15000.00, "total": 15000.00},
                {"description": "Outdoor Kitchen Station", "quantity": 1, "unit_price": 8500.00, "total": 8500.00},
                {"description": "Professional Landscaping", "quantity": 1, "unit_price": 6500.00, "total": 6500.00}
            ],
            "currency": "CHF",
            "supplier_name": "Outdoor Living AG",
            "notes": None,
            "change_request_comment": "I'd like to explore composite decking instead of natural wood for the terrace flooring - easier maintenance. Can you provide an updated quote with this alternative? Also, could we add a pergola (~3x4m) to the outdoor kitchen area?",
            "pdf_filename": None,
            "pdf_path": None,
            "ai_extraction_confidence": None,
            "parent_document_id": None,
            "due_date": None,
            "paid_date": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=8)).isoformat(),
            "updated_at": (now - timedelta(days=2)).isoformat()
        },
        # Invoice 1: Sent (awaiting payment) - for Sophie
        {
            "document_id": "doc_demo_005",
            "document_number": "INV-2024-0001",
            "type": "invoice",
            "status": "Sent",
            "agent_id": demo_agent_id,
            "client_id": demo_client1_id,
            "buyer_id": demo_buyer1_id,
            "project_id": demo_project_id,
            "unit_reference": "Unit A-301",
            "title": "Walk-in Closet Installation",
            "summary": "Custom wardrobe system with integrated LED lighting and jewelry drawers.",
            "hero_image_url": None,
            "hero_image_path": None,
            "amount": 20000.00,
            "items": [
                {"description": "Custom Wardrobe System", "quantity": 1, "unit_price": 12000.00, "total": 12000.00},
                {"description": "Integrated LED Lighting", "quantity": 1, "unit_price": 2500.00, "total": 2500.00},
                {"description": "Jewelry & Accessory Drawers", "quantity": 1, "unit_price": 3500.00, "total": 3500.00},
                {"description": "Installation & Assembly", "quantity": 1, "unit_price": 2000.00, "total": 2000.00}
            ],
            "currency": "CHF",
            "supplier_name": "Closet Masters",
            "notes": None,
            "change_request_comment": None,
            "pdf_filename": None,
            "pdf_path": None,
            "ai_extraction_confidence": None,
            "parent_document_id": "doc_demo_006",  # Link to source quote
            "due_date": (now + timedelta(days=20)).isoformat(),
            "paid_date": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=7)).isoformat(),
            "updated_at": (now - timedelta(days=7)).isoformat()
        },
        # Invoice 2: Paid - for Sophie
        {
            "document_id": "doc_demo_007",
            "document_number": "INV-2024-0002",
            "type": "invoice",
            "status": "Paid",
            "agent_id": demo_agent_id,
            "client_id": demo_client1_id,
            "buyer_id": demo_buyer1_id,
            "project_id": demo_project_id,
            "unit_reference": "Unit A-301",
            "title": "Wine Cellar Installation",
            "summary": "Climate-controlled wine cellar with custom racking and ambient lighting.",
            "hero_image_url": None,
            "hero_image_path": None,
            "amount": 16200.00,
            "items": [
                {"description": "Climate Control Unit", "quantity": 1, "unit_price": 4500.00, "total": 4500.00},
                {"description": "Custom Wine Racking", "quantity": 1, "unit_price": 8000.00, "total": 8000.00},
                {"description": "LED Lighting System", "quantity": 1, "unit_price": 1200.00, "total": 1200.00},
                {"description": "Installation", "quantity": 1, "unit_price": 2500.00, "total": 2500.00}
            ],
            "currency": "CHF",
            "supplier_name": "Wine Cellar Design",
            "notes": None,
            "change_request_comment": None,
            "pdf_filename": None,
            "pdf_path": None,
            "ai_extraction_confidence": None,
            "parent_document_id": None,
            "due_date": (now - timedelta(days=10)).isoformat(),
            "paid_date": (now - timedelta(days=15)).isoformat(),
            "is_demo": True,
            "created_at": (now - timedelta(days=40)).isoformat(),
            "updated_at": (now - timedelta(days=15)).isoformat()
        },
    ]
    
    for doc in demo_documents:
        await db.documents.insert_one(doc)
    
    
    # ==================== SEED DEMO ACTIVITIES ====================
    # Clear existing demo activities
    await db.activities.delete_many({"is_demo": True})
    await db.activity_recipients.delete_many({"is_demo": True})
    await db.activity_replies.delete_many({"is_demo": True})
    
    # Create demo units for activity context
    demo_unit1_id = "demo_unit_001"
    demo_unit2_id = "demo_unit_002"
    
    await db.project_units.delete_many({"is_demo": True})
    await db.project_units.insert_many([
        {
            "unit_id": demo_unit1_id,
            "project_id": demo_project_id,
            "unit_reference": "A-301",
            "client_id": demo_client1_id,
            "is_demo": True,
            "created_at": now.isoformat()
        },
        {
            "unit_id": demo_unit2_id,
            "project_id": demo_project_id,
            "unit_reference": "B-502",
            "client_id": demo_client2_id,
            "is_demo": True,
            "created_at": now.isoformat()
        }
    ])
    
    demo_activities = [
        # Activity 1: Construction update with image (to both clients)
        {
            "activity_id": "demo_act_001",
            "type": "image",
            "title": "Foundation Complete",
            "content": "Great news! The foundation work for your unit has been completed successfully. The concrete curing process is underway and structural work will begin next week.",
            "file_url": "/api/activities/files/demo/foundation_complete.jpg",
            "file_name": "foundation_complete.jpg",
            "file_size": 19735,
            "file_type": "image",
            "author_id": demo_agent_id,
            "author_role": "agent",
            "project_id": demo_project_id,
            "unit_id": None,  # Applies to whole project
            "is_demo": True,
            "created_at": (now - timedelta(days=30)).isoformat(),
            "updated_at": (now - timedelta(days=30)).isoformat()
        },
        # Activity 2: PDF document - floor plan (to Sophie)
        {
            "activity_id": "demo_act_002",
            "type": "file",
            "title": "Updated Floor Plan - Unit A-301",
            "content": "Please find attached the updated floor plan reflecting your requested modifications to the living room layout. The changes include an expanded balcony access and repositioned kitchen island.",
            "file_url": "/api/activities/files/demo/floor_plan_a301.pdf",
            "file_name": "floor_plan_a301.pdf",
            "file_size": 2077,
            "file_type": "pdf",
            "author_id": demo_agent_id,
            "author_role": "agent",
            "project_id": demo_project_id,
            "unit_id": demo_unit1_id,
            "is_demo": True,
            "created_at": (now - timedelta(days=20)).isoformat(),
            "updated_at": (now - timedelta(days=20)).isoformat()
        },
        # Activity 3: Contract document (to Sophie)
        {
            "activity_id": "demo_act_003",
            "type": "file",
            "title": "Premium Kitchen Upgrade Contract",
            "content": "Attached is the contract for your premium kitchen upgrade package. Please review the terms and let me know if you have any questions before signing.",
            "file_url": "/api/activities/files/demo/contract_upgrade_package.pdf",
            "file_name": "contract_upgrade_package.pdf",
            "file_size": 2152,
            "file_type": "pdf",
            "author_id": demo_agent_id,
            "author_role": "agent",
            "project_id": demo_project_id,
            "unit_id": demo_unit1_id,
            "is_demo": True,
            "created_at": (now - timedelta(days=15)).isoformat(),
            "updated_at": (now - timedelta(days=15)).isoformat()
        },
        # Activity 4: Status update - electrical (to Thomas)
        {
            "activity_id": "demo_act_004",
            "type": "status",
            "title": "Electrical Rough-In Scheduled",
            "content": "The electrical rough-in for your unit is scheduled for next week. Our technician will begin installing wiring for the smart home system you selected.",
            "file_url": None,
            "file_name": None,
            "file_size": None,
            "file_type": None,
            "author_id": demo_agent_id,
            "author_role": "agent",
            "project_id": demo_project_id,
            "unit_id": demo_unit2_id,
            "is_demo": True,
            "created_at": (now - timedelta(days=10)).isoformat(),
            "updated_at": (now - timedelta(days=10)).isoformat()
        },
        # Activity 5: Message - tile selection (to Sophie)
        {
            "activity_id": "demo_act_005",
            "type": "message",
            "title": "Action Required: Tile Selection Deadline",
            "content": "Please confirm your bathroom tile selection by Friday. The options we discussed are:\n\n1. Carrara Marble Look (Premium)\n2. Terrazzo Effect (Modern)\n3. Large Format Porcelain (Contemporary)\n\nLet me know your preference and we'll proceed with the order.",
            "file_url": None,
            "file_name": None,
            "file_size": None,
            "file_type": None,
            "author_id": demo_agent_id,
            "author_role": "agent",
            "project_id": demo_project_id,
            "unit_id": demo_unit1_id,
            "is_demo": True,
            "created_at": (now - timedelta(days=5)).isoformat(),
            "updated_at": (now - timedelta(days=4)).isoformat()
        },
        # Activity 6: Monthly report with PDF (to both)
        {
            "activity_id": "demo_act_006",
            "type": "file",
            "title": "March Progress Report",
            "content": "Here's your monthly construction progress report. All milestones are on track and we remain on schedule for the planned completion date.",
            "file_url": "/api/activities/files/demo/progress_report_march.pdf",
            "file_name": "progress_report_march.pdf",
            "file_size": 2088,
            "file_type": "pdf",
            "author_id": demo_agent_id,
            "author_role": "agent",
            "project_id": demo_project_id,
            "unit_id": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=2)).isoformat(),
            "updated_at": (now - timedelta(days=2)).isoformat()
        }
    ]
    
    for activity in demo_activities:
        await db.activities.insert_one(activity)
    
    # Create activity recipients
    demo_recipients = [
        # Activity 1: Both clients (Foundation image)
        {"recipient_id": "rcpt_001", "activity_id": "demo_act_001", "client_id": demo_client1_id, "is_demo": True, "created_at": (now - timedelta(days=30)).isoformat()},
        {"recipient_id": "rcpt_002", "activity_id": "demo_act_001", "client_id": demo_client2_id, "is_demo": True, "created_at": (now - timedelta(days=30)).isoformat()},
        # Activity 2: Sophie only (Floor plan)
        {"recipient_id": "rcpt_003", "activity_id": "demo_act_002", "client_id": demo_client1_id, "is_demo": True, "created_at": (now - timedelta(days=20)).isoformat()},
        # Activity 3: Sophie only (Contract)
        {"recipient_id": "rcpt_004", "activity_id": "demo_act_003", "client_id": demo_client1_id, "is_demo": True, "created_at": (now - timedelta(days=15)).isoformat()},
        # Activity 4: Thomas only (Electrical status)
        {"recipient_id": "rcpt_005", "activity_id": "demo_act_004", "client_id": demo_client2_id, "is_demo": True, "created_at": (now - timedelta(days=10)).isoformat()},
        # Activity 5: Sophie only (Tile selection)
        {"recipient_id": "rcpt_006", "activity_id": "demo_act_005", "client_id": demo_client1_id, "is_demo": True, "created_at": (now - timedelta(days=5)).isoformat()},
        # Activity 6: Both clients (Progress report)
        {"recipient_id": "rcpt_007", "activity_id": "demo_act_006", "client_id": demo_client1_id, "is_demo": True, "created_at": (now - timedelta(days=2)).isoformat()},
        {"recipient_id": "rcpt_008", "activity_id": "demo_act_006", "client_id": demo_client2_id, "is_demo": True, "created_at": (now - timedelta(days=2)).isoformat()}
    ]
    
    for recipient in demo_recipients:
        await db.activity_recipients.insert_one(recipient)
    
    # Create demo buyer reply (Sophie replied to tile selection)
    demo_reply = {
        "reply_id": "reply_001",
        "activity_id": "demo_act_005",
        "author_id": demo_buyer1_id,
        "author_role": "buyer",
        "content": "Thank you for the reminder. I'd like to go with option 1 - the Carrara Marble Look. It matches the kitchen design perfectly.",
        "is_demo": True,
        "created_at": (now - timedelta(days=4)).isoformat()
    }
    await db.activity_replies.insert_one(demo_reply)
    
    # ==================== SEED DEMO TEAM MEMBERS ====================
    await db.team_members.delete_many({"is_demo": True})
    
    demo_team = [
        {
            "member_id": "member_001",
            "project_id": demo_project_id,
            "agent_id": demo_agent_id,
            "company_name": "SaniTech SA",
            "contact_name": "Pierre Dupont",
            "role": "Plumber",
            "email": "pierre.dupont@sanitech.ch",
            "phone": "+41 76 555 0101",
            "website": "https://sanitech.ch",
            "notes": "Main plumbing contractor for bathrooms and kitchens",
            "is_demo": True,
            "created_at": now.isoformat()
        },
        {
            "member_id": "member_002",
            "project_id": demo_project_id,
            "agent_id": demo_agent_id,
            "company_name": "ElecPro Sàrl",
            "contact_name": "Marie Fontaine",
            "role": "Electrician",
            "email": "m.fontaine@elecpro.ch",
            "phone": "+41 76 555 0202",
            "website": "https://elecpro.ch",
            "notes": "Smart home specialist, handles all electrical installations",
            "is_demo": True,
            "created_at": now.isoformat()
        },
        {
            "member_id": "member_003",
            "project_id": demo_project_id,
            "agent_id": demo_agent_id,
            "company_name": "Meier Architekten AG",
            "contact_name": "Hans Meier",
            "role": "Architect",
            "email": "hans@meier-architekten.ch",
            "phone": "+41 76 555 0303",
            "website": "https://meier-architekten.ch",
            "notes": "Lead architect for the project",
            "is_demo": True,
            "created_at": now.isoformat()
        },
        {
            "member_id": "member_004",
            "project_id": demo_project_id,
            "agent_id": demo_agent_id,
            "name": "Anna Keller",
            "role": "Interior Designer",
            "email": "anna@kelldesign.ch",
            "phone": "+41 76 555 0404",
            "website": None,
            "notes": "Custom interior solutions and material selection",
            "is_demo": True,
            "created_at": now.isoformat()
        }
    ]
    
    for member in demo_team:
        await db.team_members.insert_one(member)
    
    # ==================== SEED DEMO TIMELINE ====================
    await db.timeline_templates.delete_many({"is_demo": True})
    await db.timeline_template_steps.delete_many({})
    await db.project_timelines.delete_many({"is_demo": True})
    await db.timeline_steps.delete_many({"is_demo": True})
    await db.timeline_step_documents.delete_many({})
    await db.timeline_step_internal_notes.delete_many({})
    
    # Create demo template
    demo_template_id = "tmpl_demo_001"
    demo_template = {
        "template_id": demo_template_id,
        "agent_id": demo_agent_id,
        "name": "Standard Construction",
        "is_demo": True,
        "created_at": now.isoformat()
    }
    await db.timeline_templates.insert_one(demo_template)
    
    # Template steps
    template_steps = [
        {"step_id": "tmpl_step_001", "template_id": demo_template_id, "title": "Site Preparation", "description": "Clear and prepare construction site", "order_index": 1},
        {"step_id": "tmpl_step_002", "template_id": demo_template_id, "title": "Excavation", "description": "Excavate foundation area", "order_index": 2},
        {"step_id": "tmpl_step_003", "template_id": demo_template_id, "title": "Foundation", "description": "Pour and cure foundation concrete", "order_index": 3},
        {"step_id": "tmpl_step_004", "template_id": demo_template_id, "title": "Structure", "description": "Build structural framework and walls", "order_index": 4},
        {"step_id": "tmpl_step_005", "template_id": demo_template_id, "title": "Finishes", "description": "Interior and exterior finishing work", "order_index": 5}
    ]
    for step in template_steps:
        await db.timeline_template_steps.insert_one(step)
    
    # Create project timeline instance
    demo_timeline_id = "timeline_demo_001"
    demo_timeline = {
        "timeline_id": demo_timeline_id,
        "project_id": demo_project_id,
        "template_id": demo_template_id,
        "is_demo": True,
        "created_at": (now - timedelta(days=60)).isoformat()
    }
    await db.project_timelines.insert_one(demo_timeline)
    
    # Create actual timeline steps with realistic progress
    timeline_steps = [
        {
            "step_id": "step_demo_001",
            "project_timeline_id": demo_timeline_id,
            "title": "Site Preparation",
            "description": "Clear vegetation, mark boundaries, set up site office and safety perimeter",
            "status": "completed",
            "order_index": 1,
            "planned_date": (now - timedelta(days=45)).strftime("%Y-%m-%d"),
            "completed_at": (now - timedelta(days=42)).isoformat(),
            "is_demo": True,
            "created_at": (now - timedelta(days=60)).isoformat(),
            "updated_at": (now - timedelta(days=42)).isoformat()
        },
        {
            "step_id": "step_demo_002",
            "project_timeline_id": demo_timeline_id,
            "title": "Excavation",
            "description": "Excavate foundation trenches, install drainage, prepare for concrete",
            "status": "completed",
            "order_index": 2,
            "planned_date": (now - timedelta(days=35)).strftime("%Y-%m-%d"),
            "completed_at": (now - timedelta(days=30)).isoformat(),
            "is_demo": True,
            "created_at": (now - timedelta(days=60)).isoformat(),
            "updated_at": (now - timedelta(days=30)).isoformat()
        },
        {
            "step_id": "step_demo_003",
            "project_timeline_id": demo_timeline_id,
            "title": "Foundation",
            "description": "Reinforcement installation, concrete pour, waterproofing membrane",
            "status": "in_progress",
            "order_index": 3,
            "planned_date": (now - timedelta(days=15)).strftime("%Y-%m-%d"),
            "completed_at": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=60)).isoformat(),
            "updated_at": (now - timedelta(days=5)).isoformat()
        },
        {
            "step_id": "step_demo_004",
            "project_timeline_id": demo_timeline_id,
            "title": "Structure",
            "description": "Steel framework, load-bearing walls, floor slabs for each level",
            "status": "pending",
            "order_index": 4,
            "planned_date": (now + timedelta(days=15)).strftime("%Y-%m-%d"),
            "completed_at": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=60)).isoformat(),
            "updated_at": (now - timedelta(days=60)).isoformat()
        },
        {
            "step_id": "step_demo_005",
            "project_timeline_id": demo_timeline_id,
            "title": "Finishes",
            "description": "Plastering, painting, flooring, fixtures, final inspections",
            "status": "pending",
            "order_index": 5,
            "planned_date": (now + timedelta(days=60)).strftime("%Y-%m-%d"),
            "completed_at": None,
            "is_demo": True,
            "created_at": (now - timedelta(days=60)).isoformat(),
            "updated_at": (now - timedelta(days=60)).isoformat()
        }
    ]
    
    for step in timeline_steps:
        await db.timeline_steps.insert_one(step)
    
    # Link the foundation activity to Foundation step
    demo_doc_link = {
        "link_id": "link_demo_001",
        "timeline_step_id": "step_demo_003",  # Foundation step
        "activity_id": "demo_act_001",  # Foundation complete image
        "created_at": (now - timedelta(days=5)).isoformat()
    }
    await db.timeline_step_documents.insert_one(demo_doc_link)
    
    # Add an internal note to the Foundation step
    demo_note = {
        "note_id": "note_demo_001",
        "timeline_step_id": "step_demo_003",
        "author_id": demo_agent_id,
        "content": "Waiting for concrete test results. Expected by end of week.",
        "created_at": (now - timedelta(days=3)).isoformat()
    }
    await db.timeline_step_internal_notes.insert_one(demo_note)
    
    return {
        "message": "Demo data seeded successfully",
        "demo_credentials": {
            "agent": {"email": "demo.agent@upgradeflow.com", "password": "demo123"},
            "buyer1": {"name": "Sophie Müller"},
            "buyer2": {"name": "Thomas Weber"}
        }
    }

@api_router.post("/demo/reset")
async def reset_demo_data():
    """Reset demo data to fresh state"""
    return await seed_demo_data()


@api_router.post("/admin/migrate-clients")
async def migrate_client_data(user: dict = Depends(get_current_agent)):
    """
    Migration endpoint to fix clients with missing fields.
    Adds default values for unit_reference and status.
    """
    # Find clients with missing unit_reference
    clients_without_unit_ref = await db.clients.find(
        {"unit_reference": {"$exists": False}},
        {"_id": 0, "client_id": 1}
    ).to_list(1000)
    
    # Find clients with missing status
    clients_without_status = await db.clients.find(
        {"status": {"$exists": False}},
        {"_id": 0, "client_id": 1}
    ).to_list(1000)
    
    # Update clients with missing unit_reference
    if clients_without_unit_ref:
        await db.clients.update_many(
            {"unit_reference": {"$exists": False}},
            {"$set": {"unit_reference": "General"}}
        )
    
    # Update clients with missing status
    if clients_without_status:
        await db.clients.update_many(
            {"status": {"$exists": False}},
            {"$set": {"status": "active"}}
        )
    
    return {
        "migrated": True,
        "clients_fixed_unit_reference": len(clients_without_unit_ref),
        "clients_fixed_status": len(clients_without_status)
    }


@api_router.get("/admin/data-health")
async def check_data_health(user: dict = Depends(get_current_agent)):
    """
    Check data integrity across collections.
    Returns counts of records with missing required fields.
    """
    # Check clients
    clients_missing_unit_ref = await db.clients.count_documents({"unit_reference": {"$exists": False}})
    clients_missing_status = await db.clients.count_documents({"status": {"$exists": False}})
    
    # Check documents
    docs_missing_client = await db.documents.count_documents({"client_id": {"$exists": False}})
    docs_missing_project = await db.documents.count_documents({"project_id": {"$exists": False}})
    
    # Check buyers without client linkage
    all_buyers = await db.users.find({"role": "buyer"}, {"_id": 0, "user_id": 1}).to_list(1000)
    buyer_ids = [b['user_id'] for b in all_buyers]
    linked_buyers = await db.clients.distinct("buyer_id", {"buyer_id": {"$in": buyer_ids}})
    orphan_buyers = len(buyer_ids) - len(linked_buyers)
    
    issues = []
    if clients_missing_unit_ref > 0:
        issues.append(f"{clients_missing_unit_ref} clients missing unit_reference")
    if clients_missing_status > 0:
        issues.append(f"{clients_missing_status} clients missing status")
    if docs_missing_client > 0:
        issues.append(f"{docs_missing_client} documents missing client_id")
    if docs_missing_project > 0:
        issues.append(f"{docs_missing_project} documents missing project_id")
    if orphan_buyers > 0:
        issues.append(f"{orphan_buyers} buyers without client linkage")
    
    return {
        "healthy": len(issues) == 0,
        "issues": issues,
        "details": {
            "clients_missing_unit_reference": clients_missing_unit_ref,
            "clients_missing_status": clients_missing_status,
            "documents_missing_client_id": docs_missing_client,
            "documents_missing_project_id": docs_missing_project,
            "buyers_without_client_linkage": orphan_buyers
        }
    }

# ==================== BILLING / SUBSCRIPTION ENDPOINTS ====================

STRIPE_API_KEY = os.environ.get('STRIPE_API_KEY', '')

async def get_agent_unit_count(agent_id: str, is_demo: bool = False) -> int:
    """Count total units across all projects owned by an agent"""
    # Get all projects for this agent
    projects = await db.projects.find({
        "agent_id": agent_id,
        "is_demo": is_demo
    }, {"project_id": 1}).to_list(None)
    
    if not projects:
        return 0
    
    project_ids = [p['project_id'] for p in projects]
    
    # Count units in the project_units collection
    total_units = await db.project_units.count_documents({
        "project_id": {"$in": project_ids},
        "is_demo": is_demo
    })
    
    # If no units in project_units collection, count projects as units
    # (for backwards compatibility with projects that don't have explicit units)
    if total_units == 0:
        total_units = len(projects)
    
    return total_units

async def get_agent_subscription_data(user: dict) -> dict:
    """Get subscription data for an agent from database"""
    agent_id = user['user_id']
    is_demo = user.get('is_demo', False)
    
    # Get user's subscription info from database
    user_doc = await db.users.find_one({"user_id": agent_id}, {"_id": 0})
    
    # Default to free plan if no subscription
    plan_id = user_doc.get('subscription_plan', 'free') if user_doc else 'free'
    subscription_status = user_doc.get('subscription_status', 'active')
    
    # Get plan details
    plan = SUBSCRIPTION_PLANS.get(plan_id, SUBSCRIPTION_PLANS['free'])
    
    # Count units (not projects)
    unit_count = await get_agent_unit_count(agent_id, is_demo)
    
    # Determine if can create more units
    unit_limit = plan.get('property_limit')  # Still called property_limit in plans config
    can_create = unit_limit is None or unit_count < unit_limit
    
    # Calculate usage percentage for warnings
    usage_percent = 0
    near_limit = False
    if unit_limit and unit_limit > 0:
        usage_percent = (unit_count / unit_limit) * 100
        near_limit = usage_percent >= 80 and can_create  # 80% threshold warning
    
    return {
        "plan_id": plan_id,
        "plan_name": plan['name'],
        "unit_limit": unit_limit,  # Renamed for clarity
        "unit_usage": unit_count,  # Renamed for clarity
        "property_limit": unit_limit,  # Keep for backwards compatibility
        "property_usage": unit_count,  # Keep for backwards compatibility
        "usage_percent": usage_percent,
        "near_limit": near_limit,  # True if at 80%+ usage
        "can_create_property": can_create and subscription_status == 'active',
        "can_create_unit": can_create and subscription_status == 'active',
        "stripe_subscription_id": user_doc.get('stripe_subscription_id') if user_doc else None,
        "stripe_customer_id": user_doc.get('stripe_customer_id') if user_doc else None,
        "subscription_status": subscription_status,
        "current_period_end": user_doc.get('subscription_period_end') if user_doc else None
    }

@api_router.get("/billing/plans")
async def get_available_plans(user: dict = Depends(get_current_agent)):
    """Get all available subscription plans"""
    plans = []
    for plan_id, plan_data in SUBSCRIPTION_PLANS.items():
        plans.append({
            "plan_id": plan_id,
            "name": plan_data['name'],
            "price": plan_data['price'],
            "currency": plan_data['currency'],
            "property_limit": plan_data['property_limit'],
            "features": plan_data['features'],
            "is_enterprise": plan_id == 'enterprise'
        })
    return plans

@api_router.get("/billing/status")
async def get_subscription_status(user: dict = Depends(get_current_agent)):
    """Get current subscription status for the agent"""
    return await get_agent_subscription_data(user)

@api_router.post("/billing/create-checkout-session")
async def create_checkout_session(data: CreateCheckoutRequest, user: dict = Depends(get_current_agent)):
    """Create a Stripe checkout session for subscribing to a plan"""
    if not STRIPE_API_KEY:
        raise HTTPException(status_code=500, detail="Stripe is not configured")
    
    plan_id = data.plan_id
    if plan_id not in SUBSCRIPTION_PLANS or plan_id in ['free', 'enterprise']:
        raise HTTPException(status_code=400, detail="Invalid plan selected")
    
    plan = SUBSCRIPTION_PLANS[plan_id]
    
    try:
        # Use Stripe SDK directly
        stripe.api_key = STRIPE_API_KEY
        
        # Amount in cents for Stripe
        amount_in_cents = int(plan['price'] * 100)
        
        session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price_data': {
                    'currency': 'chf',
                    'unit_amount': amount_in_cents,
                    'product_data': {
                        'name': f"Evohome {plan['name']} Plan",
                        'description': f"Monthly subscription - {plan['property_limit']} units"
                    },
                },
                'quantity': 1,
            }],
            mode='payment',
            success_url=f"{data.origin_url}/agent/billing?session_id={{CHECKOUT_SESSION_ID}}&success=true",
            cancel_url=f"{data.origin_url}/agent/billing?canceled=true",
            metadata={
                "agent_id": user['user_id'],
                "plan_id": plan_id,
                "plan_name": plan['name']
            }
        )
        
        return {
            "checkout_url": session.url,
            "session_id": session.id
        }
        
    except Exception as e:
        # Capture Stripe errors for monitoring
        capture_payment_error(e, user_id=user['user_id'], operation="create_checkout")
        logger.error(f"Stripe checkout error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to create checkout session. Please try again.")

@api_router.post("/billing/verify-session")
async def verify_checkout_session(data: CheckoutStatusRequest, user: dict = Depends(get_current_agent)):
    """Verify a checkout session and update subscription status"""
    if not STRIPE_API_KEY:
        raise HTTPException(status_code=500, detail="Stripe is not configured")
    
    try:
        # Use Stripe SDK directly
        stripe.api_key = STRIPE_API_KEY
        session = stripe.checkout.Session.retrieve(data.session_id)
        
        logger.info(f"Checkout status: {session.status}, payment: {session.payment_status}")
        
        if session.status == "complete" and session.payment_status == "paid":
            # Get plan_id from metadata (set during checkout creation)
            metadata = session.metadata or {}
            plan_id = metadata.get('plan_id', 'starter')
            
            # Fallback: determine plan from amount if metadata missing
            if not plan_id and session.amount_total:
                amount = session.amount_total / 100
                if amount >= 79:
                    plan_id = "pro"
                elif amount >= 29:
                    plan_id = "starter"
                else:
                    plan_id = "free"
            
            # Update user's subscription in database
            update_data = {
                "subscription_plan": plan_id,
                "subscription_status": "active",
                "subscription_updated_at": datetime.now(timezone.utc).isoformat()
            }
            
            # Store session ID for reference
            update_data["last_checkout_session"] = data.session_id
            
            await db.users.update_one(
                {"user_id": user['user_id']},
                {"$set": update_data}
            )
            
            logger.info(f"Subscription updated for {user['user_id']}: plan={plan_id}")
            
            return {
                "success": True,
                "plan_id": plan_id,
                "subscription_status": "active"
            }
        else:
            return {
                "success": False,
                "status": session.status,
                "payment_status": session.payment_status
            }
            
    except Exception as e:
        logger.error(f"Stripe session verification error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to verify session: {str(e)}")

@api_router.post("/billing/webhook")
async def stripe_webhook(request: Request):
    """Handle Stripe webhook events for subscription updates"""
    if not STRIPE_API_KEY:
        raise HTTPException(status_code=500, detail="Stripe is not configured")
    
    try:
        payload = await request.body()
        event_data = json.loads(payload)
        event_type = event_data.get('type', '')
        
        logger.info(f"Received Stripe webhook: {event_type}")
        
        if event_type == 'checkout.session.completed':
            session = event_data.get('data', {}).get('object', {})
            metadata = session.get('metadata', {})
            agent_id = metadata.get('agent_id')
            plan_id = metadata.get('plan_id', 'starter')
            
            if agent_id:
                await db.users.update_one(
                    {"user_id": agent_id},
                    {"$set": {
                        "stripe_customer_id": session.get('customer'),
                        "stripe_subscription_id": session.get('subscription'),
                        "subscription_plan": plan_id,
                        "subscription_status": "active"
                    }}
                )
                logger.info(f"Updated subscription for agent {agent_id} to plan {plan_id}")
        
        elif event_type == 'customer.subscription.updated':
            subscription = event_data.get('data', {}).get('object', {})
            customer_id = subscription.get('customer')
            status = subscription.get('status')
            
            # Find user by customer ID and update status
            await db.users.update_many(
                {"stripe_customer_id": customer_id},
                {"$set": {
                    "subscription_status": status,
                    "subscription_period_end": subscription.get('current_period_end')
                }}
            )
            logger.info(f"Updated subscription status to {status} for customer {customer_id}")
        
        elif event_type == 'customer.subscription.deleted':
            subscription = event_data.get('data', {}).get('object', {})
            customer_id = subscription.get('customer')
            
            # Downgrade to free plan
            await db.users.update_many(
                {"stripe_customer_id": customer_id},
                {"$set": {
                    "subscription_plan": "free",
                    "subscription_status": "canceled",
                    "stripe_subscription_id": None
                }}
            )
            logger.info(f"Subscription canceled for customer {customer_id}, downgraded to free")
        
        elif event_type == 'invoice.payment_failed':
            invoice = event_data.get('data', {}).get('object', {})
            customer_id = invoice.get('customer')
            
            await db.users.update_many(
                {"stripe_customer_id": customer_id},
                {"$set": {"subscription_status": "past_due"}}
            )
            logger.info(f"Payment failed for customer {customer_id}")
        
        return {"received": True}
        
    except Exception as e:
        # Capture webhook errors for monitoring
        capture_payment_error(e, operation="webhook_processing")
        logger.error(f"Webhook processing error: {str(e)}")
        raise HTTPException(status_code=400, detail="Webhook processing failed")

@api_router.post("/billing/cancel")
async def cancel_subscription(user: dict = Depends(get_current_agent)):
    """Cancel the current subscription (keeps active until period end)"""
    user_doc = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0})
    
    if not user_doc or not user_doc.get('stripe_subscription_id'):
        raise HTTPException(status_code=400, detail="No active subscription found")
    
    # Note: In production, you would call Stripe API to cancel
    # For now, we'll mark it as canceled locally
    await db.users.update_one(
        {"user_id": user['user_id']},
        {"$set": {"subscription_status": "canceling"}}
    )
    
    return {"message": "Subscription will be canceled at period end"}

@api_router.post("/billing/sync")
async def sync_subscription_from_stripe(user: dict = Depends(get_current_agent)):
    """Sync subscription status - re-verify the last checkout session"""
    if not STRIPE_API_KEY:
        raise HTTPException(status_code=500, detail="Stripe is not configured")
    
    user_doc = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    
    last_session = user_doc.get('last_checkout_session')
    
    if not last_session:
        # No checkout session found - check if they have a plan already
        current_plan = user_doc.get('subscription_plan', 'free')
        return {
            "message": "No checkout session to sync",
            "synced": False,
            "current_plan": current_plan
        }
    
    try:
        # Use Stripe SDK directly
        stripe.api_key = STRIPE_API_KEY
        session = stripe.checkout.Session.retrieve(last_session)
        
        logger.info(f"Sync - Checkout status for {user['user_id']}: {session.status}, payment: {session.payment_status}")
        
        if session.status == "complete" and session.payment_status == "paid":
            # Get plan_id from metadata
            metadata = session.metadata or {}
            plan_id = metadata.get('plan_id')
            
            # Fallback: determine plan from amount
            if not plan_id and session.amount_total:
                amount = session.amount_total / 100
                if amount >= 79:
                    plan_id = "pro"
                elif amount >= 29:
                    plan_id = "starter"
                else:
                    plan_id = "free"
            
            if not plan_id:
                plan_id = "starter"  # Default
            
            # Update subscription
            update_data = {
                "subscription_plan": plan_id,
                "subscription_status": "active",
                "subscription_updated_at": datetime.now(timezone.utc).isoformat()
            }
            
            await db.users.update_one(
                {"user_id": user['user_id']},
                {"$set": update_data}
            )
            
            logger.info(f"Synced subscription for {user['user_id']}: plan={plan_id}")
            
            return {
                "message": "Subscription synced successfully",
                "synced": True,
                "plan_id": plan_id,
                "status": "active"
            }
        else:
            return {
                "message": f"Checkout session not complete (status: {session.status}, payment: {session.payment_status})",
                "synced": False,
                "current_plan": user_doc.get('subscription_plan', 'free')
            }
            
    except Exception as e:
        logger.error(f"Failed to sync subscription: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to sync: {str(e)}")

@api_router.post("/billing/portal")
async def create_billing_portal(request: Request, user: dict = Depends(get_current_agent)):
    """Create a Stripe billing portal session for subscription management"""
    if not STRIPE_API_KEY:
        raise HTTPException(status_code=500, detail="Stripe is not configured")
    
    user_doc = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0})
    
    if not user_doc or not user_doc.get('stripe_customer_id'):
        raise HTTPException(status_code=400, detail="No Stripe customer found. Please subscribe first.")
    
    try:
        # Get the return URL from request body
        body = await request.json()
        return_url = body.get('return_url', '')
        
        # Use Stripe API directly for billing portal (emergentintegrations may not have this)
        import stripe
        stripe.api_key = STRIPE_API_KEY
        
        portal_session = stripe.billing_portal.Session.create(
            customer=user_doc['stripe_customer_id'],
            return_url=return_url or f"{os.environ.get('FRONTEND_URL', '')}/agent/billing"
        )
        
        return {
            "portal_url": portal_session.url
        }
        
    except Exception as e:
        logger.error(f"Billing portal error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create billing portal: {str(e)}")


# ==================== TEAM INVITATIONS ENDPOINTS ====================

class TeamInviteCreate(BaseModel):
    email: EmailStr
    role: str = "member"  # "admin" or "member"
    message: Optional[str] = None

class TeamInviteResponse(BaseModel):
    invitation_id: str
    email: str
    role: str
    status: str  # "pending", "accepted", "declined", "expired"
    invited_by: str
    invited_by_name: str
    created_at: str
    expires_at: str

@api_router.post("/team/invitations")
async def create_team_invitation(data: TeamInviteCreate, user: dict = Depends(get_current_agent)):
    """Invite a new team member to the agent's workspace"""
    is_demo = user.get('is_demo', False)
    
    # Check if user already exists as an agent
    existing_agent = await db.users.find_one({"email": data.email, "role": "agent"}, {"_id": 0})
    if existing_agent:
        # Check if already part of this workspace
        if existing_agent.get('workspace_owner_id') == user['user_id']:
            raise HTTPException(status_code=400, detail="This user is already part of your team")
        if existing_agent.get('workspace_owner_id'):
            raise HTTPException(status_code=400, detail="This user is already part of another workspace")
    
    # Check for existing pending invitation
    existing_invite = await db.team_invitations.find_one({
        "email": data.email,
        "invited_by": user['user_id'],
        "status": "pending",
        "is_demo": is_demo
    })
    if existing_invite:
        raise HTTPException(status_code=400, detail="An invitation has already been sent to this email")
    
    # Create invitation
    invitation_id = f"invite_{uuid.uuid4().hex[:12]}"
    invitation_token = secrets.token_urlsafe(32)
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    
    invitation_doc = {
        "invitation_id": invitation_id,
        "email": data.email.lower(),
        "role": data.role,
        "message": data.message,
        "status": "pending",
        "invited_by": user['user_id'],
        "invited_by_name": user.get('name', 'Unknown'),
        "invitation_token": invitation_token,
        "is_demo": is_demo,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": expires_at.isoformat()
    }
    
    await db.team_invitations.insert_one(invitation_doc)
    
    # Send invitation email
    frontend_url = os.environ.get('REACT_APP_BACKEND_URL', 'https://evohome.ch').replace('/api', '')
    invite_link = f"{frontend_url}/team/accept?token={invitation_token}"
    
    company_name = user.get('settings', {}).get('company_name') or user.get('name', 'Evohome')
    
    subject = f"You've been invited to join {company_name} on Evohome"
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <h2 style="color: #2563EB;">Team Invitation</h2>
        <p>Hi,</p>
        <p><strong>{user.get('name', 'A team member')}</strong> has invited you to join <strong>{company_name}</strong> on Evohome as a <strong>{data.role}</strong>.</p>
        {f'<p style="color: #666; font-style: italic;">"{data.message}"</p>' if data.message else ''}
        <p style="text-align: center; margin: 30px 0;">
            <a href="{invite_link}" style="background-color: #2563EB; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block;">
                Accept Invitation
            </a>
        </p>
        <p style="color: #666; font-size: 14px;">This invitation will expire in 7 days.</p>
        <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;">
        <p style="color: #999; font-size: 12px;">Evohome - Real Estate Management</p>
    </div>
    """
    
    await send_email_async(data.email, subject, html_content)
    
    logger.info(f"Team invitation sent to {data.email} by {user['user_id']}")
    
    return {
        "invitation_id": invitation_id,
        "email": data.email,
        "role": data.role,
        "status": "pending",
        "invited_by": user['user_id'],
        "invited_by_name": user.get('name', 'Unknown'),
        "created_at": invitation_doc['created_at'],
        "expires_at": invitation_doc['expires_at']
    }

@api_router.get("/team/invitations")
async def list_team_invitations(user: dict = Depends(get_current_agent)):
    """List all team invitations sent by this agent"""
    is_demo = user.get('is_demo', False)
    
    invitations = await db.team_invitations.find(
        {"invited_by": user['user_id']},
        {"_id": 0, "invitation_token": 0}
    ).sort("created_at", -1).to_list(100)
    
    return invitations

@api_router.delete("/team/invitations/{invitation_id}")
async def cancel_team_invitation(invitation_id: str, user: dict = Depends(get_current_agent)):
    """Cancel a pending team invitation"""
    is_demo = user.get('is_demo', False)
    
    invitation = await db.team_invitations.find_one({
        "invitation_id": invitation_id,
        "invited_by": user['user_id'],
        "is_demo": is_demo
    })
    
    if not invitation:
        raise HTTPException(status_code=404, detail="Invitation not found")
    
    if invitation['status'] != 'pending':
        raise HTTPException(status_code=400, detail="Can only cancel pending invitations")
    
    await db.team_invitations.delete_one({"invitation_id": invitation_id})
    
    return {"message": "Invitation cancelled"}

@api_router.get("/team/members")
async def list_team_members(user: dict = Depends(get_current_agent)):
    """List all team members in the agent's workspace"""
    is_demo = user.get('is_demo', False)
    
    # Get the workspace owner - either self or owner if this is a team member
    workspace_owner_id = user.get('workspace_owner_id') or user['user_id']
    
    # Find all agents in this workspace
    team_members = await db.users.find(
        {
            "$or": [
                {"user_id": workspace_owner_id},
                {"workspace_owner_id": workspace_owner_id}
            ],
            "role": "agent",
            "is_demo": is_demo
        },
        {"_id": 0, "password_hash": 0, "password_reset_token": 0, "password_reset_expires": 0}
    ).to_list(100)
    
    # Enrich with role info
    for member in team_members:
        if member['user_id'] == workspace_owner_id:
            member['team_role'] = 'owner'
        else:
            member['team_role'] = member.get('workspace_role', 'member')
    
    return team_members

@api_router.post("/team/accept")
async def accept_team_invitation(token: str, response: Response):
    """Accept a team invitation (public endpoint, creates/links account)"""
    
    invitation = await db.team_invitations.find_one({
        "invitation_token": token,
        "status": "pending"
    })
    
    if not invitation:
        raise HTTPException(status_code=404, detail="Invitation not found or already used")
    
    # Check expiry
    expires_at = datetime.fromisoformat(invitation['expires_at'].replace('Z', '+00:00'))
    if datetime.now(timezone.utc) > expires_at:
        await db.team_invitations.update_one(
            {"invitation_id": invitation['invitation_id']},
            {"$set": {"status": "expired"}}
        )
        raise HTTPException(status_code=400, detail="Invitation has expired")
    
    is_demo = invitation.get('is_demo', False)
    email = invitation['email']
    
    # Check if user already exists
    existing_user = await db.users.find_one({"email": email, "role": "agent"}, {"_id": 0})
    
    if existing_user:
        # Link existing user to workspace
        if existing_user.get('workspace_owner_id'):
            raise HTTPException(status_code=400, detail="You are already part of another workspace")
        
        await db.users.update_one(
            {"user_id": existing_user['user_id']},
            {"$set": {
                "workspace_owner_id": invitation['invited_by'],
                "workspace_role": invitation['role']
            }}
        )
        user_id = existing_user['user_id']
    else:
        # Return info for registration - user needs to complete signup
        return {
            "status": "needs_registration",
            "email": email,
            "invited_by_name": invitation['invited_by_name'],
            "role": invitation['role'],
            "token": token
        }
    
    # Mark invitation as accepted
    await db.team_invitations.update_one(
        {"invitation_id": invitation['invitation_id']},
        {"$set": {
            "status": "accepted",
            "accepted_at": datetime.now(timezone.utc).isoformat(),
            "accepted_by": user_id
        }}
    )
    
    # Create session token
    token = create_jwt_token(user_id, "agent", is_demo)
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    return {
        "status": "accepted",
        "user_id": user_id,
        "workspace_owner_id": invitation['invited_by'],
        "token": token
    }

@api_router.post("/team/register-invited")
async def register_invited_user(
    email: str = Form(...),
    password: str = Form(...),
    name: str = Form(...),
    token: str = Form(...),
    response: Response = None
):
    """Register a new user from a team invitation"""
    
    invitation = await db.team_invitations.find_one({
        "invitation_token": token,
        "status": "pending",
        "email": email.lower()
    })
    
    if not invitation:
        raise HTTPException(status_code=404, detail="Invalid invitation token")
    
    # Check expiry
    expires_at = datetime.fromisoformat(invitation['expires_at'].replace('Z', '+00:00'))
    if datetime.now(timezone.utc) > expires_at:
        raise HTTPException(status_code=400, detail="Invitation has expired")
    
    is_demo = invitation.get('is_demo', False)
    
    # Check if user already exists
    existing = await db.users.find_one({"email": email.lower(), "role": "agent"})
    if existing:
        raise HTTPException(status_code=400, detail="Account already exists. Please login and accept the invitation.")
    
    # Create new user
    hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
    user_id = f"agent_{uuid.uuid4().hex[:12]}"
    
    user_doc = {
        "user_id": user_id,
        "email": email.lower(),
        "name": name,
        "password_hash": hashed_password.decode('utf-8'),
        "role": "agent",
        "picture": None,
        "is_demo": is_demo,
        "workspace_owner_id": invitation['invited_by'],
        "workspace_role": invitation['role'],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    # Mark invitation as accepted
    await db.team_invitations.update_one(
        {"invitation_id": invitation['invitation_id']},
        {"$set": {
            "status": "accepted",
            "accepted_at": datetime.now(timezone.utc).isoformat(),
            "accepted_by": user_id
        }}
    )
    
    # Create session
    jwt_token = create_jwt_token(user_id, "agent", is_demo)
    
    response.set_cookie(
        key="session_token",
        value=jwt_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    logger.info(f"New team member registered: {user_id} joined workspace of {invitation['invited_by']}")
    
    return {
        "user_id": user_id,
        "email": email,
        "name": name,
        "role": "agent",
        "workspace_owner_id": invitation['invited_by'],
        "workspace_role": invitation['role'],
        "token": jwt_token
    }

@api_router.delete("/team/members/{member_id}")
async def remove_team_member(member_id: str, user: dict = Depends(get_current_agent)):
    """Remove a team member from the workspace (owner only)"""
    is_demo = user.get('is_demo', False)
    
    # Only workspace owner can remove members
    if user.get('workspace_owner_id'):
        raise HTTPException(status_code=403, detail="Only workspace owner can remove team members")
    
    # Find the member
    member = await db.users.find_one({
        "user_id": member_id,
        "workspace_owner_id": user['user_id'],
        "is_demo": is_demo
    })
    
    if not member:
        raise HTTPException(status_code=404, detail="Team member not found")
    
    # Remove from workspace (don't delete the user)
    await db.users.update_one(
        {"user_id": member_id},
        {"$unset": {"workspace_owner_id": "", "workspace_role": ""}}
    )
    
    logger.info(f"Team member {member_id} removed from workspace by {user['user_id']}")
    
    return {"message": "Team member removed from workspace"}


# ==================== AGENT SETTINGS ENDPOINTS ====================

class BillingSettings(BaseModel):
    iban: Optional[str] = None
    company_name: Optional[str] = None
    address: Optional[str] = None
    postal_code: Optional[str] = None
    city: Optional[str] = None

class AgentProfileUpdate(BaseModel):
    display_name: Optional[str] = None
    contact_email: Optional[str] = None
    contact_phone: Optional[str] = None

class AgentSettingsUpdate(BaseModel):
    language: Optional[str] = None
    currency: Optional[str] = None
    company_name: Optional[str] = None
    billing: Optional[BillingSettings] = None
    profile: Optional[AgentProfileUpdate] = None

@api_router.get("/settings")
async def get_agent_settings(user: dict = Depends(get_current_agent)):
    """Get agent settings including profile info"""
    user_doc = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0})
    settings = user_doc.get('settings', {}) if user_doc else {}
    profile = user_doc.get('profile', {}) if user_doc else {}
    
    return {
        "language": settings.get('language', 'en'),
        "currency": settings.get('currency', 'CHF'),
        "company_name": settings.get('company_name', ''),
        "company_logo_url": user_doc.get('company_logo_url') if user_doc else None,
        "billing": settings.get('billing', {}),
        "profile": {
            "display_name": profile.get('display_name', user_doc.get('name', '') if user_doc else ''),
            "contact_email": profile.get('contact_email', user_doc.get('email', '') if user_doc else ''),
            "contact_phone": profile.get('contact_phone', '')
        }
    }

@api_router.get("/branding")
async def get_agent_branding_for_buyer(user: dict = Depends(get_current_user)):
    """Get agent branding info for buyer view - returns the agent's branding who manages this buyer"""
    if user['role'] == 'agent':
        # Agents see their own branding
        user_doc = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0})
        return {
            "company_name": user_doc.get('settings', {}).get('company_name', '') if user_doc else '',
            "company_logo_url": user_doc.get('company_logo_url') if user_doc else None,
            "language": user_doc.get('settings', {}).get('language', 'en') if user_doc else 'en',
            "currency": user_doc.get('settings', {}).get('currency', 'CHF') if user_doc else 'CHF'
        }
    
    # For buyers, find their associated agent through client relationship
    client = await db.clients.find_one({"buyer_id": user['user_id']}, {"_id": 0})
    if not client:
        # No client relationship - return defaults
        return {
            "company_name": "",
            "company_logo_url": None,
            "language": "en",
            "currency": "CHF"
        }
    
    # Find the project to get agent_id
    project = await db.projects.find_one({"project_id": client.get('project_id')}, {"_id": 0})
    if not project:
        return {
            "company_name": "",
            "company_logo_url": None,
            "language": "en",
            "currency": "CHF"
        }
    
    agent_id = project.get('agent_id')
    agent = await db.users.find_one({"user_id": agent_id}, {"_id": 0})
    
    if not agent:
        return {
            "company_name": "",
            "company_logo_url": None,
            "language": "en",
            "currency": "CHF"
        }
    
    return {
        "company_name": agent.get('settings', {}).get('company_name', ''),
        "company_logo_url": agent.get('company_logo_url'),
        "language": agent.get('settings', {}).get('language', 'en'),
        "currency": agent.get('settings', {}).get('currency', 'CHF')
    }

@api_router.put("/settings")
async def update_agent_settings(data: AgentSettingsUpdate, user: dict = Depends(get_current_agent)):
    """Update agent settings including profile info"""
    update_data = {}
    
    if data.language:
        if data.language not in ['en', 'de', 'fr', 'it']:
            raise HTTPException(status_code=400, detail="Invalid language")
        update_data['settings.language'] = data.language
    
    if data.currency:
        if data.currency not in ['CHF', 'EUR', 'USD']:
            raise HTTPException(status_code=400, detail="Invalid currency")
        update_data['settings.currency'] = data.currency
    
    if data.company_name is not None:
        update_data['settings.company_name'] = data.company_name
    
    if data.billing:
        billing_dict = data.billing.dict(exclude_none=True)
        for key, value in billing_dict.items():
            update_data[f'settings.billing.{key}'] = value
    
    # Handle profile updates
    if data.profile:
        if data.profile.display_name is not None:
            update_data['profile.display_name'] = data.profile.display_name
        if data.profile.contact_email is not None:
            update_data['profile.contact_email'] = data.profile.contact_email
        if data.profile.contact_phone is not None:
            update_data['profile.contact_phone'] = data.profile.contact_phone
    
    if update_data:
        await db.users.update_one(
            {"user_id": user['user_id']},
            {"$set": update_data}
        )
    
    return {"message": "Settings updated"}

@api_router.post("/settings/logo")
async def upload_company_logo(
    file: UploadFile = File(...),
    user: dict = Depends(get_current_agent)
):
    """Upload company logo (Pro plan required)"""
    # Check if user has Pro plan or above
    subscription_data = await get_agent_subscription_data(user)
    if subscription_data['plan_id'] not in ['pro', 'enterprise']:
        raise HTTPException(status_code=403, detail="Logo upload requires Pro plan or higher")
    
    # Validate file
    allowed_types = ['image/jpeg', 'image/jpg', 'image/png', 'image/webp']
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Only JPEG, PNG, and WebP images are allowed")
    
    # Read file content
    content = await file.read()
    
    # Check file size (max 2MB)
    if len(content) > 2 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Image must be less than 2MB")
    
    # Delete old logo if exists
    user_doc = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0})
    if user_doc and user_doc.get('company_logo_path'):
        old_path = Path(user_doc['company_logo_path'])
        if old_path.exists():
            try:
                old_path.unlink()
            except OSError:
                pass
    
    # Save new logo
    file_ext = file.filename.split('.')[-1] if '.' in file.filename else 'png'
    logo_filename = f"logo_{user['user_id']}_{uuid.uuid4().hex[:8]}.{file_ext}"
    logo_path = UPLOAD_DIR / logo_filename
    
    with open(logo_path, "wb") as f:
        f.write(content)
    
    logo_url = f"/api/uploads/{logo_filename}"
    
    await db.users.update_one(
        {"user_id": user['user_id']},
        {"$set": {
            "company_logo_url": logo_url,
            "company_logo_path": str(logo_path),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "logo_url": logo_url,
        "message": "Logo uploaded successfully"
    }

# ==================== ADMIN / DIAGNOSTIC ENDPOINTS ====================

@api_router.get("/admin/diagnose-buyer/{email}")
async def diagnose_buyer_account(email: str, user: dict = Depends(get_current_agent)):
    """
    Diagnose buyer account linkage issues.
    Returns detailed information about buyer account, client records, and their linkage.
    """
    is_demo = user.get('is_demo', False)
    
    # Find buyer user
    buyer_user = await db.users.find_one(
        {"email": email, "role": "buyer"},
        {"_id": 0, "password_hash": 0}
    )
    
    # Find client records with this email
    client_records = await db.clients.find(
        {"email": email},
        {"_id": 0}
    ).to_list(100)
    
    # Find client records linked to this buyer (if buyer exists)
    linked_clients = []
    if buyer_user:
        linked_clients = await db.clients.find(
            {"buyer_id": buyer_user['user_id']},
            {"_id": 0}
        ).to_list(100)
    
    # Find activities for these clients
    all_client_ids = list(set([c['client_id'] for c in client_records] + [c['client_id'] for c in linked_clients]))
    
    activity_recipients = await db.activity_recipients.find(
        {"client_id": {"$in": all_client_ids}},
        {"_id": 0}
    ).to_list(500)
    
    # Diagnose issues
    issues = []
    
    if not buyer_user:
        issues.append({
            "type": "NO_BUYER_ACCOUNT",
            "message": f"No buyer user account exists for email {email}. Buyer needs to register first."
        })
    
    if not client_records:
        issues.append({
            "type": "NO_CLIENT_RECORD",
            "message": f"No client record exists for email {email}. Agent needs to create a client."
        })
    
    # Check for unlinked clients (buyer exists but client.buyer_id is null)
    unlinked_clients = [c for c in client_records if not c.get('buyer_id') and buyer_user]
    if unlinked_clients:
        issues.append({
            "type": "UNLINKED_CLIENTS",
            "message": f"Found {len(unlinked_clients)} client record(s) not linked to buyer account",
            "client_ids": [c['client_id'] for c in unlinked_clients]
        })
    
    # Check for mismatched buyer_id
    if buyer_user:
        mismatched = [c for c in client_records if c.get('buyer_id') and c['buyer_id'] != buyer_user['user_id']]
        if mismatched:
            issues.append({
                "type": "BUYER_ID_MISMATCH",
                "message": f"Found {len(mismatched)} client(s) linked to different buyer account",
                "details": [(c['client_id'], c['buyer_id']) for c in mismatched]
            })
    
    if not activity_recipients:
        issues.append({
            "type": "NO_ACTIVITIES",
            "message": "No feed activities have been posted to this client"
        })
    
    return {
        "email": email,
        "buyer_user": buyer_user,
        "client_records": client_records,
        "linked_via_buyer_id": linked_clients,
        "activity_recipient_count": len(activity_recipients),
        "issues": issues,
        "can_auto_fix": len(unlinked_clients) > 0 and buyer_user is not None
    }

@api_router.post("/admin/fix-buyer-linkage/{email}")
async def fix_buyer_client_linkage(email: str, user: dict = Depends(get_current_agent)):
    """
    Fix buyer-client linkage for a specific email.
    Links all client records with this email to the corresponding buyer account.
    """
    is_demo = user.get('is_demo', False)
    
    # Find buyer user
    buyer_user = await db.users.find_one(
        {"email": email, "role": "buyer"},
        {"_id": 0}
    )
    
    if not buyer_user:
        raise HTTPException(
            status_code=404, 
            detail=f"No buyer account found for {email}. Buyer must register first."
        )
    
    # Find and update unlinked client records
    result = await db.clients.update_many(
        {"email": email, "buyer_id": None},
        {"$set": {"buyer_id": buyer_user['user_id']}}
    )
    
    # Also check for clients with matching email but no buyer_id set
    result2 = await db.clients.update_many(
        {"email": email, "buyer_id": {"$exists": False}},
        {"$set": {"buyer_id": buyer_user['user_id']}}
    )
    
    total_fixed = result.modified_count + result2.modified_count
    
    logger.info(f"Fixed buyer linkage for {email}: {total_fixed} client records updated")
    
    return {
        "message": f"Successfully linked {total_fixed} client record(s) to buyer account",
        "email": email,
        "buyer_id": buyer_user['user_id'],
        "records_fixed": total_fixed
    }

@api_router.get("/admin/email-status")
async def check_email_configuration(user: dict = Depends(get_current_agent)):
    """
    Check email system configuration status.
    Returns configuration status and allows sending test emails.
    """
    return {
        "resend_api_key_configured": bool(RESEND_API_KEY),
        "sender_email_configured": bool(SENDER_EMAIL),
        "sender_email": SENDER_EMAIL if SENDER_EMAIL else "NOT SET",
        "frontend_url": FRONTEND_URL,
        "status": "ready" if (RESEND_API_KEY and SENDER_EMAIL) else "not_configured",
        "issues": [
            issue for issue in [
                "RESEND_API_KEY not set" if not RESEND_API_KEY else None,
                "SENDER_EMAIL not set" if not SENDER_EMAIL else None,
            ] if issue
        ]
    }

@api_router.post("/admin/test-email")
async def send_test_email(to_email: str, user: dict = Depends(get_current_agent)):
    """
    Send a test email to verify email system is working.
    """
    if not RESEND_API_KEY:
        raise HTTPException(status_code=503, detail="RESEND_API_KEY not configured")
    if not SENDER_EMAIL:
        raise HTTPException(status_code=503, detail="SENDER_EMAIL not configured")
    
    subject = "Evohome Test Email"
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head><meta charset="utf-8"></head>
    <body style="font-family: Arial, sans-serif; padding: 20px;">
        <h1 style="color: #2563EB;">Test Email from Evohome</h1>
        <p>This is a test email to verify your email configuration is working correctly.</p>
        <p><strong>Configuration:</strong></p>
        <ul>
            <li>Sender: {SENDER_EMAIL}</li>
            <li>Frontend URL: {FRONTEND_URL}</li>
            <li>Sent by: {user.get('name', user['user_id'])}</li>
        </ul>
        <p>If you received this email, your Resend integration is working!</p>
    </body>
    </html>
    """
    
    result = await send_email_async(to_email, subject, html_content)
    
    if result.get("status") == "error":
        raise HTTPException(status_code=500, detail=f"Email failed: {result.get('error')}")
    
    return {
        "message": f"Test email sent to {to_email}",
        "result": result
    }

# ==================== END ADMIN ENDPOINTS ====================

@api_router.delete("/settings/logo")
async def delete_company_logo(user: dict = Depends(get_current_agent)):
    """Delete company logo"""
    user_doc = await db.users.find_one({"user_id": user['user_id']}, {"_id": 0})
    
    if user_doc and user_doc.get('company_logo_path'):
        logo_path = Path(user_doc['company_logo_path'])
        if logo_path.exists():
            try:
                logo_path.unlink()
            except OSError:
                pass
    
    await db.users.update_one(
        {"user_id": user['user_id']},
        {"$unset": {"company_logo_url": "", "company_logo_path": ""}}
    )
    
    return {"message": "Logo deleted"}

# ==================== COMMAND SERVICE ENDPOINTS ====================
# Phase 2: Agent Command Workspace
# All execution creates drafts first, then confirms

from services.command_service import (
    CommandInterpreter,
    CommandExecutor,
    CommandContext,
    CommandPlan,
    CommandDraft,
    CommandIntent,
    TOOL_REGISTRY
)

# Initialize command services
command_interpreter = CommandInterpreter(db)
command_executor = CommandExecutor(db)

# Request models for command endpoints
class CommandInterpretRequest(BaseModel):
    command: str
    context: Optional[dict] = None
    
class CommandExecuteRequest(BaseModel):
    draft_id: str
    confirmed: bool = True


@api_router.post("/command/interpret")
async def interpret_command(
    command: str = Form(...),
    context: str = Form("{}"),
    user: dict = Depends(get_current_agent)
):
    """
    Interpret a command into a structured plan.
    
    This endpoint:
    - Classifies the intent (create_quote, create_invoice, create_message)
    - Extracts fields from the command text
    - Validates required fields
    - Returns a structured plan for confirmation
    
    NO side effects. NO database writes (except logging).
    """
    try:
        # Parse context
        ctx_dict = json.loads(context) if context else {}
        cmd_context = CommandContext(
            project_id=ctx_dict.get("project_id"),
            client_id=ctx_dict.get("client_id"),
            unit_id=ctx_dict.get("unit_id")
        )
        
        # Interpret the command
        plan = await command_interpreter.interpret(
            command_text=command,
            context=cmd_context,
            user_id=user['user_id']
        )
        
        # Convert to dict for response
        plan_dict = plan.model_dump()
        
        # Convert enum values to strings
        plan_dict['intent'] = plan.intent.value
        plan_dict['fields'] = [
            {**f.model_dump(), 'confidence': f.confidence}
            for f in plan.fields
        ]
        plan_dict['missing_fields'] = [
            f.model_dump() for f in plan.missing_fields
        ]
        
        # Convert datetime to ISO string
        plan_dict['created_at'] = plan.created_at.isoformat()
        
        return plan_dict
        
    except Exception as e:
        logging.error(f"Command interpretation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Interpretation failed: {str(e)}")


@api_router.post("/command/draft")
async def create_command_draft(
    plan: dict,
    user: dict = Depends(get_current_agent)
):
    """
    Create a draft from a validated plan.
    
    This creates a draft object that can be:
    - Reviewed and edited
    - Confirmed for execution
    - Cancelled
    
    Does NOT create the actual object yet.
    """
    try:
        # Reconstruct the CommandPlan
        from services.command_service import ExtractedField, MissingField
        
        cmd_plan = CommandPlan(
            plan_id=plan.get('plan_id', f"plan_{uuid.uuid4().hex[:12]}"),
            intent=CommandIntent(plan['intent']),
            intent_confidence=plan.get('intent_confidence', 1.0),
            entities=plan.get('entities', {}),
            fields=[
                ExtractedField(**f) for f in plan.get('fields', [])
            ],
            missing_fields=[
                MissingField(**f) for f in plan.get('missing_fields', [])
            ],
            is_valid=plan.get('is_valid', False),
            validation_errors=plan.get('validation_errors', []),
            requires_confirmation=True,
            can_execute=plan.get('can_execute', False),
            raw_command=plan.get('raw_command', '')
        )
        
        # Create draft
        draft = await command_executor.create_draft(cmd_plan, user['user_id'])
        
        # Return draft info
        return {
            "draft_id": draft.draft_id,
            "plan_id": draft.plan_id,
            "intent": draft.intent.value,
            "status": draft.status.value,
            "draft_data": draft.draft_data,
            "created_at": draft.created_at.isoformat()
        }
        
    except Exception as e:
        logging.error(f"Draft creation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Draft creation failed: {str(e)}")


@api_router.post("/command/execute")
async def execute_command(
    request: CommandExecuteRequest,
    user: dict = Depends(get_current_agent)
):
    """
    Execute a confirmed draft.
    
    This:
    - Validates the draft exists and is pending
    - Checks user authorization
    - Creates the actual object through existing services
    - Logs the execution
    - Prevents duplicate execution (idempotent)
    
    If confirmed=false, cancels the draft instead.
    """
    try:
        # Check if draft was already executed (idempotency)
        existing_draft = await db.command_drafts.find_one(
            {"draft_id": request.draft_id},
            {"_id": 0, "status": 1, "result_id": 1, "result_type": 1}
        )
        
        if existing_draft and existing_draft.get("status") == "executed":
            # Already executed - return the existing result
            logger.info(f"Returning cached execution result for draft: {request.draft_id}")
            return {
                "status": "executed",
                "draft_id": request.draft_id,
                "result": {
                    "type": existing_draft.get("result_type"),
                    "id": existing_draft.get("result_id"),
                    "already_executed": True
                }
            }
        
        result = await command_executor.execute_draft(
            draft_id=request.draft_id,
            user_id=user['user_id'],
            confirmed=request.confirmed,
            is_demo=user.get('is_demo', False)
        )
        
        return result
        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logging.error(f"Command execution failed: {e}")
        raise HTTPException(status_code=500, detail=f"Execution failed: {str(e)}")


@api_router.get("/command/draft/{draft_id}")
async def get_draft(
    draft_id: str,
    user: dict = Depends(get_current_agent)
):
    """Get a command draft by ID"""
    draft = await db.command_drafts.find_one(
        {"draft_id": draft_id, "created_by": user['user_id']},
        {"_id": 0}
    )
    
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    
    return draft


@api_router.get("/command/drafts")
async def list_drafts(
    status: Optional[str] = None,
    user: dict = Depends(get_current_agent)
):
    """List command drafts for the current user"""
    query = {"created_by": user['user_id']}
    if status:
        query["status"] = status
    
    drafts = await db.command_drafts.find(
        query,
        {"_id": 0}
    ).sort("created_at", -1).limit(50).to_list(50)
    
    return drafts


@api_router.get("/command/tools")
async def list_tools(user: dict = Depends(get_current_agent)):
    """
    List available command tools and their definitions.
    Useful for UI hints and validation.
    """
    tools = []
    for intent, tool in TOOL_REGISTRY.items():
        tools.append({
            "intent": intent.value,
            "name": tool.name,
            "description": tool.description,
            "required_fields": tool.required_fields,
            "optional_fields": tool.optional_fields
        })
    return tools


@api_router.get("/command/logs")
async def get_command_logs(
    draft_id: Optional[str] = None,
    limit: int = 50,
    user: dict = Depends(get_current_agent)
):
    """Get command execution logs"""
    query = {"user_id": user['user_id']}
    if draft_id:
        query["draft_id"] = draft_id
    
    logs = await db.command_logs.find(
        query,
        {"_id": 0}
    ).sort("timestamp", -1).limit(limit).to_list(limit)
    
    return logs


@api_router.get("/command/history")
async def get_command_history(
    limit: int = 20,
    user: dict = Depends(get_current_agent)
):
    """
    Get recent command history for the user.
    Includes both executed drafts and recent extractions.
    """
    is_demo = user.get('is_demo', False)
    
    # Get recent drafts
    drafts = await db.command_drafts.find(
        {"created_by": user['user_id']},
        {"_id": 0}
    ).sort("created_at", -1).limit(limit).to_list(limit)
    
    # Get recent extraction cache entries
    extractions = await db.extraction_cache.find(
        {"user_id": user['user_id']},
        {"_id": 0, "idempotency_key": 1, "created_at": 1, "result.document_type": 1, "result.intent": 1}
    ).sort("created_at", -1).limit(limit).to_list(limit)
    
    return {
        "drafts": drafts,
        "recent_extractions": extractions
    }


@api_router.post("/command/draft/auto-save")
async def auto_save_draft(
    plan_id: str = Form(...),
    intent: str = Form(...),
    draft_data: str = Form("{}"),
    user: dict = Depends(get_current_agent)
):
    """
    Auto-save a draft in progress.
    Called periodically by the frontend to persist work.
    """
    try:
        data = json.loads(draft_data) if draft_data else {}
        
        # Upsert the auto-save draft
        await db.command_drafts_autosave.update_one(
            {"plan_id": plan_id, "created_by": user['user_id']},
            {"$set": {
                "plan_id": plan_id,
                "intent": intent,
                "draft_data": data,
                "created_by": user['user_id'],
                "updated_at": datetime.now(timezone.utc).isoformat()
            }},
            upsert=True
        )
        
        return {"status": "saved", "plan_id": plan_id}
    except Exception as e:
        logger.error(f"Auto-save failed: {e}")
        raise HTTPException(status_code=500, detail="Auto-save failed")


@api_router.get("/command/draft/auto-save/{plan_id}")
async def get_auto_saved_draft(
    plan_id: str,
    user: dict = Depends(get_current_agent)
):
    """
    Retrieve an auto-saved draft.
    Used to recover work after browser refresh.
    """
    draft = await db.command_drafts_autosave.find_one(
        {"plan_id": plan_id, "created_by": user['user_id']},
        {"_id": 0}
    )
    
    if not draft:
        raise HTTPException(status_code=404, detail="No auto-saved draft found")
    
    return draft


# ==================== DOCUMENT EXTRACTION ENDPOINTS (Phase 3) ====================

from services.command_service import (
    classify_document,
    DocumentType,
)


@api_router.post("/command/classify-document")
async def classify_uploaded_document(
    file: UploadFile = File(...),
    user: dict = Depends(get_current_agent)
):
    """
    Classify an uploaded document type.
    Returns document type with confidence score.
    Agent can override the classification.
    
    Supports: PDF files and images (jpg, jpeg, png, webp)
    """
    try:
        # Validate file type - PDFs and images supported
        file_ext = os.path.splitext(file.filename)[1].lower()
        supported_extensions = [".pdf", ".jpg", ".jpeg", ".png", ".webp"]
        
        if file_ext not in supported_extensions:
            raise HTTPException(
                status_code=400, 
                detail="Unsupported file type. Supported formats: PDF, JPG, PNG, WEBP"
            )
        
        # Save file temporarily
        file_id = f"temp_{uuid.uuid4().hex[:12]}"
        file_path = f"/tmp/{file_id}{file_ext}"
        
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # Extract text for classification
        text_content = ""
        extraction_method = "none"
        
        if file_ext == ".pdf":
            try:
                doc = fitz.open(file_path)
                for page_num in range(min(3, len(doc))):  # First 3 pages
                    text_content += doc[page_num].get_text()
                doc.close()
                extraction_method = "pdf_text"
            except Exception as e:
                logger.warning(f"PDF text extraction failed: {e}")
                if os.path.exists(file_path):
                    os.remove(file_path)
                raise HTTPException(
                    status_code=400,
                    detail="Could not read PDF file. The file may be corrupted or password-protected."
                )
        else:
            # Image file - use OCR via OpenAI Vision
            if OPENAI_API_KEY:
                try:
                    import base64
                    with open(file_path, "rb") as img_file:
                        image_data = base64.b64encode(img_file.read()).decode('utf-8')
                    
                    # Determine mime type
                    mime_types = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png", ".webp": "image/webp"}
                    mime_type = mime_types.get(file_ext, "image/jpeg")
                    
                    client = openai.OpenAI(api_key=OPENAI_API_KEY)
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "Extract all text from this document image. Include numbers, dates, and any structured data. Return only the extracted text, no commentary."},
                                {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_data}"}}
                            ]
                        }],
                        max_tokens=2000
                    )
                    text_content = response.choices[0].message.content
                    extraction_method = "ocr_vision"
                except Exception as e:
                    logger.warning(f"OCR extraction failed: {e}")
                    extraction_method = "ocr_failed"
            else:
                extraction_method = "no_api_key"
        
        # Classify document
        doc_type, confidence = classify_document(text_content, file.filename)
        
        # Adjust confidence based on extraction method
        if extraction_method in ["ocr_failed", "no_api_key", "none"]:
            confidence = min(confidence, 0.3)  # Lower confidence if we couldn't extract text
        
        return {
            "file_id": file_id,
            "file_path": file_path,
            "filename": file.filename,
            "document_type": doc_type.value,
            "confidence": round(confidence, 2),
            "extraction_method": extraction_method,
            "can_override": True,
            "available_types": [t.value for t in DocumentType if t != DocumentType.UNKNOWN]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document classification failed: {e}")
        raise HTTPException(status_code=500, detail="Document classification failed. Please try again.")


@api_router.post("/command/extract-document")
async def extract_document_data(
    file_path: str = Form(...),
    document_type: str = Form(...),
    context: str = Form("{}"),
    idempotency_key: str = Form(None),
    user: dict = Depends(get_current_agent)
):
    """
    Extract structured data from a classified document.
    Returns extracted fields that populate the draft form.
    
    Supports: PDF and image files (jpg, jpeg, png, webp)
    
    Features:
    - Idempotency key to prevent duplicate extractions
    - Validation of extracted amounts (must be positive, reasonable range)
    - Retry-safe: same idempotency_key returns cached result
    """
    try:
        ctx_dict = json.loads(context) if context else {}
        is_demo = user.get('is_demo', False)
        
        # Check idempotency - return cached result if exists
        if idempotency_key:
            cached = await db.extraction_cache.find_one({
                "idempotency_key": idempotency_key,
                "user_id": user['user_id']
            }, {"_id": 0})
            if cached and cached.get("result"):
                logger.info(f"Returning cached extraction for key: {idempotency_key}")
                return cached["result"]
        
        # Validate file exists
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=400, 
                detail="Document file not found. Please re-upload the document."
            )
        
        # Validate file type - PDFs and images supported
        file_ext = os.path.splitext(file_path)[1].lower()
        supported_extensions = [".pdf", ".jpg", ".jpeg", ".png", ".webp"]
        if file_ext not in supported_extensions:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Supported: PDF, JPG, PNG, WEBP"
            )
        
        # Validate document type
        try:
            doc_type = DocumentType(document_type)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid document type: {document_type}")
        
        # Map document type to extraction intent
        intent_map = {
            DocumentType.QUOTE: CommandIntent.EXTRACT_QUOTE,
            DocumentType.INVOICE: CommandIntent.EXTRACT_INVOICE,
            DocumentType.TIMELINE: CommandIntent.EXTRACT_TIMELINE,
            DocumentType.CONTACTS: CommandIntent.EXTRACT_CONTACTS,
        }
        
        intent = intent_map.get(doc_type)
        if not intent:
            raise HTTPException(status_code=400, detail=f"Cannot extract from document type: {document_type}")
        
        # Extract data using appropriate method based on file type
        extracted_data = {}
        extraction_confidence = 0.5
        extraction_warnings = []
        
        if doc_type in [DocumentType.QUOTE, DocumentType.INVOICE]:
            # Use existing PDF extraction for PDFs, or image extraction for images
            if file_ext == ".pdf":
                extraction_result = await extract_document_from_pdf(file_path, os.path.basename(file_path))
            else:
                # Image extraction using Vision API
                extraction_result = await _extract_document_from_image(file_path, doc_type.value)
            
            extracted_data = {
                "supplier_name": extraction_result.get("supplier_name"),
                "total_amount": extraction_result.get("amount"),
                "currency": "CHF",
                "description": extraction_result.get("description"),
                "title": extraction_result.get("title"),
                "line_items": extraction_result.get("items", []),
                "reference_number": extraction_result.get("reference"),
            }
            
            # Validate extracted amount
            amount = extracted_data.get("total_amount")
            if amount is not None:
                try:
                    amount = float(amount)
                    if amount < 0:
                        extraction_warnings.append("Negative amount detected - please verify")
                        extracted_data["_amount_warning"] = "negative"
                    elif amount > 10000000:  # 10M threshold
                        extraction_warnings.append("Unusually large amount - please verify")
                        extracted_data["_amount_warning"] = "large"
                    elif amount == 0:
                        extraction_warnings.append("Zero amount detected - please verify")
                        extracted_data["_amount_warning"] = "zero"
                except (ValueError, TypeError):
                    extraction_warnings.append("Could not parse amount - manual entry required")
                    extracted_data["total_amount"] = None
            
            confidence_map = {"high": 0.9, "medium": 0.7, "low": 0.4}
            extraction_confidence = confidence_map.get(extraction_result.get("confidence", "low"), 0.4)
            
            if extraction_result.get("extraction_failed"):
                extraction_confidence = 0.1
        
        elif doc_type == DocumentType.TIMELINE:
            # Check if project already has a timeline
            project_id = ctx_dict.get("project_id")
            if project_id:
                existing_timeline = await db.project_timelines.find_one({
                    "project_id": project_id,
                    "agent_id": user['user_id'],
                    "is_demo": user.get('is_demo', False)
                }, {"_id": 0})
                
                if existing_timeline:
                    # Timeline already exists - return info about existing timeline
                    return {
                        "plan_id": f"plan_{uuid.uuid4().hex[:12]}",
                        "intent": "view_timeline",  # Change intent to view instead of create
                        "document_type": "timeline",
                        "can_execute": False,
                        "timeline_exists": True,
                        "existing_timeline": {
                            "timeline_id": existing_timeline['timeline_id'],
                            "name": existing_timeline.get('name', 'Project Timeline'),
                            "created_at": existing_timeline.get('created_at')
                        },
                        "fields": [
                            {"name": "existing_timeline_id", "value": existing_timeline['timeline_id'], "confidence": 1.0, "source": "database"},
                            {"name": "timeline_name", "value": existing_timeline.get('name', 'Project Timeline'), "confidence": 1.0, "source": "database"},
                            {"name": "project_id", "value": project_id, "confidence": 1.0, "source": "context"}
                        ],
                        "missing_fields": [],
                        "message": f"This project already has a timeline: '{existing_timeline.get('name', 'Project Timeline')}'. Would you like to view or update it?",
                        "available_actions": [
                            {"action": "view", "label": "View Existing Timeline", "path": f"/agent/timeline?project={project_id}"},
                            {"action": "replace", "label": "Replace Timeline", "warning": "This will delete the existing timeline"}
                        ],
                        "source_file": file_path,
                        "created_at": datetime.now(timezone.utc).isoformat()
                    }
            
            # No existing timeline - proceed with extraction
            extracted_data = await _extract_timeline_stages(file_path)
            extraction_confidence = extracted_data.get("confidence", 0.5)
        
        elif doc_type == DocumentType.CONTACTS:
            # Extract contacts
            extracted_data = await _extract_contacts_list(file_path)
            extraction_confidence = extracted_data.get("confidence", 0.5)
        
        # Build the command plan
        fields = []
        for key, value in extracted_data.items():
            if value is not None and key != "confidence":
                conf = 0.9 if value else 0.3
                fields.append({
                    "name": key,
                    "value": value,
                    "confidence": conf,
                    "source": "ai_extraction"
                })
        
        # Add context fields
        if ctx_dict.get("project_id"):
            fields.append({
                "name": "project_id",
                "value": ctx_dict["project_id"],
                "confidence": 1.0,
                "source": "context"
            })
        
        if ctx_dict.get("client_id"):
            fields.append({
                "name": "client_id",
                "value": ctx_dict["client_id"],
                "confidence": 1.0,
                "source": "context"
            })
        
        # Determine missing required fields
        missing_fields = []
        if not ctx_dict.get("project_id"):
            missing_fields.append({
                "name": "project_id",
                "description": "Select a project",
                "required": True
            })
        
        if doc_type in [DocumentType.QUOTE, DocumentType.INVOICE] and not ctx_dict.get("client_id"):
            missing_fields.append({
                "name": "client_id",
                "description": "Select a client",
                "required": False  # Can be added later
            })
        
        # Check if we can execute
        can_execute = len([m for m in missing_fields if m.get("required")]) == 0
        
        # Build result
        result = {
            "plan_id": f"plan_{uuid.uuid4().hex[:12]}",
            "intent": intent.value,
            "intent_confidence": extraction_confidence,
            "document_type": doc_type.value,
            "entities": {
                "project_id": ctx_dict.get("project_id"),
                "client_id": ctx_dict.get("client_id"),
            },
            "fields": fields,
            "extracted_data": extracted_data,
            "missing_fields": missing_fields,
            "is_valid": can_execute,
            "validation_errors": [],
            "extraction_warnings": extraction_warnings if 'extraction_warnings' in dir() else [],
            "requires_confirmation": True,
            "can_execute": can_execute,
            "source_file": file_path,
            "interpretation_log": [
                f"Document type: {doc_type.value}",
                f"Extraction confidence: {extraction_confidence:.0%}",
                f"Fields extracted: {len(fields)}",
            ]
        }
        
        # Cache the result for idempotency
        if idempotency_key:
            await db.extraction_cache.update_one(
                {"idempotency_key": idempotency_key, "user_id": user['user_id']},
                {"$set": {
                    "result": result,
                    "created_at": datetime.now(timezone.utc)
                }},
                upsert=True
            )
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")


# ==================== WORKFLOW ENDPOINTS (Phase 4) ====================

from services.workflow_service import (
    get_workflow_service, 
    WorkflowTemplate, 
    WorkflowExecution,
    WorkflowStatus,
    WorkflowStepStatus
)

class WorkflowExecuteRequest(BaseModel):
    template_id: str
    context: Dict[str, Any] = {}
    mode: str = "automatic"  # "automatic" or "step_by_step"

class WorkflowConfirmRequest(BaseModel):
    skip_step: bool = False

@api_router.get("/workflows/templates")
async def get_workflow_templates(
    category: Optional[str] = None,
    user: dict = Depends(get_current_agent)
):
    """Get all available workflow templates"""
    service = get_workflow_service(db)
    templates = service.get_templates(category)
    
    return {
        "templates": [
            {
                "template_id": t.template_id,
                "name": t.name,
                "description": t.description,
                "category": t.category,
                "icon": t.icon,
                "estimated_duration": t.estimated_duration,
                "required_context": t.required_context,
                "ui_selectors": t.ui_selectors,
                "steps_count": len(t.steps),
                "steps_preview": [
                    {"name": s["name"], "description": s["description"], "optional": s.get("optional", False)}
                    for s in t.steps
                ]
            }
            for t in templates
        ]
    }

@api_router.get("/workflows/templates/{template_id}")
async def get_workflow_template(
    template_id: str,
    user: dict = Depends(get_current_agent)
):
    """Get details of a specific workflow template"""
    service = get_workflow_service(db)
    template = service.get_template(template_id)
    
    if not template:
        raise HTTPException(status_code=404, detail="Workflow template not found")
    
    return {
        "template_id": template.template_id,
        "name": template.name,
        "description": template.description,
        "category": template.category,
        "icon": template.icon,
        "estimated_duration": template.estimated_duration,
        "required_context": template.required_context,
        "steps": [
            {
                "name": s["name"],
                "description": s["description"],
                "action": s["action"],
                "optional": s.get("optional", False),
                "requires_confirmation": s.get("requires_confirmation", False)
            }
            for s in template.steps
        ]
    }

@api_router.post("/workflows/execute")
async def execute_workflow(
    request: WorkflowExecuteRequest,
    user: dict = Depends(get_current_agent)
):
    """Start executing a workflow"""
    service = get_workflow_service(db)
    is_demo = user.get('is_demo', False)
    
    # Get agent profile for email sending
    agent_profile = await db.agent_profiles.find_one(
        {"user_id": user['user_id']},
        {"_id": 0}
    )
    agent_name = agent_profile.get('display_name', user.get('name', 'Your Agent')) if agent_profile else 'Your Agent'
    agent_email = agent_profile.get('contact_email', user.get('email')) if agent_profile else user.get('email')
    
    try:
        # Enrich context with document/step details if needed
        enriched_context = dict(request.context)
        
        # If document_id provided, fetch document details
        if enriched_context.get('document_id'):
            doc = await db.documents.find_one(
                {"document_id": enriched_context['document_id']},
                {"_id": 0}
            )
            if doc:
                enriched_context['document_type'] = doc.get('type', 'Document')
                enriched_context['document_title'] = doc.get('title', 'Untitled')
                enriched_context['amount'] = doc.get('total_amount', 0)
                enriched_context['client_id'] = doc.get('client_id')
                enriched_context['project_id'] = doc.get('project_id')
        
        # If client_id provided, fetch client details
        if enriched_context.get('client_id'):
            client = await db.clients.find_one(
                {"client_id": enriched_context['client_id']},
                {"_id": 0}
            )
            if client:
                enriched_context['client_name'] = client.get('name', 'Client')
                enriched_context['client_email'] = client.get('email')
        
        # If project_id provided, fetch project details
        if enriched_context.get('project_id'):
            project = await db.projects.find_one(
                {"project_id": enriched_context['project_id']},
                {"_id": 0}
            )
            if project:
                enriched_context['project_name'] = project.get('name', 'Project')
        
        # If step_id provided, fetch step details
        if enriched_context.get('step_id'):
            step = await db.timeline_steps.find_one(
                {"step_id": enriched_context['step_id']},
                {"_id": 0}
            )
            if step:
                enriched_context['step_name'] = step.get('name', 'Milestone')
                enriched_context['project_id'] = step.get('project_id') or enriched_context.get('project_id')
                # Get client from project
                if enriched_context.get('project_id') and not enriched_context.get('client_id'):
                    project = await db.projects.find_one(
                        {"project_id": enriched_context['project_id']},
                        {"_id": 0}
                    )
                    if project and project.get('client_id'):
                        enriched_context['client_id'] = project['client_id']
                        client = await db.clients.find_one(
                            {"client_id": project['client_id']},
                            {"_id": 0}
                        )
                        if client:
                            enriched_context['client_name'] = client.get('name')
                            enriched_context['client_email'] = client.get('email')
        
        # Create execution with enriched context
        execution = await service.create_execution(
            template_id=request.template_id,
            agent_id=user['user_id'],
            context=enriched_context,
            mode=request.mode,
            is_demo=is_demo
        )
        
        # Track step warnings for non-fatal issues
        step_warnings = {}
        
        # Helper to safely send email with error handling
        async def safe_send_email(to_email: str, subject: str, html_content: str, step_name: str) -> dict:
            """Send email with graceful error handling"""
            if not to_email:
                return {"sent": False, "warning": "No recipient email address"}
            
            if not RESEND_API_KEY:
                logger.warning(f"RESEND_API_KEY not configured, skipping email for {step_name}")
                return {"sent": False, "warning": "Email service not configured"}
            
            try:
                await send_email_async(to_email, subject, html_content)
                return {"sent": True}
            except Exception as e:
                logger.warning(f"Email failed for {step_name}: {e}")
                return {"sent": False, "warning": f"Email delivery failed: {str(e)[:100]}"}
        
        # Define action executor with real email sending
        async def action_executor(action: str, params: Dict, context: Dict) -> Dict:
            """Execute a workflow action and return results"""
            result = {}
            
            if action == "create_client":
                client_data = {
                    "client_id": f"cl_{uuid.uuid4().hex[:8]}",
                    "agent_id": user['user_id'],
                    "project_id": params.get("project_id") or context.get("project_id"),
                    "name": params.get("client_name") or context.get("client_name", "New Client"),
                    "email": params.get("client_email") or context.get("client_email"),
                    "phone": params.get("client_phone") or context.get("client_phone"),
                    "status": "active",
                    "is_demo": is_demo,
                    "created_at": datetime.now(timezone.utc).isoformat()
                }
                await db.clients.insert_one(client_data)
                result["client_id"] = client_data["client_id"]
                result["client_email"] = client_data["email"]
                result["client_name"] = client_data["name"]
                logger.info(f"Workflow created client: {client_data['client_id']}")
                
            elif action == "send_welcome_email":
                client_email = params.get("client_email") or context.get("client_email")
                client_name = params.get("client_name") or context.get("client_name", "")
                project_name = params.get("project_name") or context.get("project_name", "your project")
                
                subject = f"Welcome to {project_name}!"
                html_content = f"""
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                    <h2 style="color: #3b82f6;">Welcome, {client_name}!</h2>
                    <p>You've been added to <strong>{project_name}</strong> by {agent_name}.</p>
                    <p>You can now:</p>
                    <ul>
                        <li>View project updates and timeline</li>
                        <li>Review and approve quotes</li>
                        <li>Track invoice payments</li>
                        <li>Access shared documents</li>
                    </ul>
                    <p style="margin-top: 24px;">
                        <a href="{os.environ.get('FRONTEND_URL', 'https://evohome.app')}/login" 
                           style="background-color: #3b82f6; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px;">
                            Access Your Portal
                        </a>
                    </p>
                    <p style="margin-top: 24px; color: #6b7280; font-size: 14px;">
                        If you have any questions, contact {agent_name} at {agent_email}.
                    </p>
                </div>
                """
                email_result = await safe_send_email(client_email, subject, html_content, "send_welcome_email")
                result["email_sent"] = email_result["sent"]
                if email_result.get("warning"):
                    result["_warning"] = email_result["warning"]
                    step_warnings["send_welcome_email"] = email_result["warning"]
                logger.info(f"Workflow welcome email: sent={email_result['sent']}, to={client_email}")
                
            elif action == "update_document_status":
                doc_id = params.get("document_id") or context.get("document_id")
                new_status = params.get("status", "Sent")
                if doc_id:
                    update_data = {"status": new_status, "updated_at": datetime.now(timezone.utc).isoformat()}
                    if new_status == "Paid":
                        update_data["paid_at"] = datetime.now(timezone.utc).isoformat()
                    elif new_status == "Sent":
                        update_data["sent_at"] = datetime.now(timezone.utc).isoformat()
                    
                    await db.documents.update_one(
                        {"document_id": doc_id},
                        {"$set": update_data}
                    )
                    result["status"] = new_status
                    logger.info(f"Workflow updated document {doc_id} to {new_status}")
                    
            elif action == "send_payment_confirmation_email":
                client_email = params.get("client_email") or context.get("client_email")
                client_name = params.get("client_name") or context.get("client_name", "")
                amount = params.get("amount") or context.get("amount", 0)
                doc_id = params.get("document_id") or context.get("document_id")
                
                subject = "Payment Received - Thank You!"
                html_content = f"""
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                    <h2 style="color: #10b981;">Payment Received!</h2>
                    <p>Dear {client_name},</p>
                    <p>We have received your payment of <strong>CHF {amount:,.2f}</strong>.</p>
                    <p>Your invoice has been marked as paid. Thank you for your prompt payment!</p>
                    <p style="margin-top: 24px; color: #6b7280; font-size: 14px;">
                        Reference: {doc_id}<br>
                        Processed by: {agent_name}
                    </p>
                </div>
                """
                email_result = await safe_send_email(client_email, subject, html_content, "send_payment_confirmation_email")
                result["email_sent"] = email_result["sent"]
                if email_result.get("warning"):
                    result["_warning"] = email_result["warning"]
                    step_warnings["send_payment_confirmation_email"] = email_result["warning"]
                    
            elif action == "complete_timeline_step":
                step_id = params.get("step_id") or context.get("step_id")
                if step_id:
                    await db.timeline_steps.update_one(
                        {"step_id": step_id},
                        {"$set": {"status": "completed", "completed_at": datetime.now(timezone.utc).isoformat()}}
                    )
                    result["step_status"] = "completed"
                    logger.info(f"Workflow completed step: {step_id}")
                    
            elif action == "send_milestone_email":
                client_email = params.get("client_email") or context.get("client_email")
                client_name = params.get("client_name") or context.get("client_name", "")
                step_name = params.get("step_name") or context.get("step_name", "a milestone")
                project_name = params.get("project_name") or context.get("project_name", "your project")
                
                subject = f"Milestone Completed: {step_name}"
                html_content = f"""
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                    <h2 style="color: #f59e0b;">Milestone Completed!</h2>
                    <p>Dear {client_name},</p>
                    <p>Great news! The following milestone has been completed for <strong>{project_name}</strong>:</p>
                    <div style="background-color: #fef3c7; padding: 16px; border-radius: 8px; margin: 16px 0;">
                        <strong>{step_name}</strong>
                    </div>
                    <p>View your project timeline for full details and upcoming milestones.</p>
                    <p style="margin-top: 24px;">
                        <a href="{os.environ.get('FRONTEND_URL', 'https://evohome.app')}/buyer/timeline" 
                           style="background-color: #f59e0b; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px;">
                            View Timeline
                        </a>
                    </p>
                    <p style="margin-top: 24px; color: #6b7280; font-size: 14px;">
                        Best regards,<br>{agent_name}
                    </p>
                </div>
                """
                email_result = await safe_send_email(client_email, subject, html_content, "send_milestone_email")
                result["email_sent"] = email_result["sent"]
                if email_result.get("warning"):
                    result["_warning"] = email_result["warning"]
                    step_warnings["send_milestone_email"] = email_result["warning"]
                    
            elif action == "send_document_email":
                client_email = params.get("client_email") or context.get("client_email")
                client_name = params.get("client_name") or context.get("client_name", "")
                doc_type = params.get("document_type") or context.get("document_type", "Document")
                doc_title = params.get("document_title") or context.get("document_title", "")
                doc_id = params.get("document_id") or context.get("document_id")
                
                subject = f"New {doc_type}: {doc_title}"
                html_content = f"""
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                    <h2 style="color: #3b82f6;">New {doc_type} Available</h2>
                    <p>Dear {client_name},</p>
                    <p>{agent_name} has sent you a new {doc_type.lower()}:</p>
                    <div style="background-color: #eff6ff; padding: 16px; border-radius: 8px; margin: 16px 0;">
                        <strong>{doc_title}</strong>
                    </div>
                    <p>Please review and respond at your earliest convenience.</p>
                    <p style="margin-top: 24px;">
                        <a href="{os.environ.get('FRONTEND_URL', 'https://evohome.app')}/buyer/documents/{doc_id}" 
                           style="background-color: #3b82f6; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px;">
                            View {doc_type}
                        </a>
                    </p>
                    <p style="margin-top: 24px; color: #6b7280; font-size: 14px;">
                        Best regards,<br>{agent_name}
                    </p>
                </div>
                """
                email_result = await safe_send_email(client_email, subject, html_content, "send_document_email")
                result["email_sent"] = email_result["sent"]
                if email_result.get("warning"):
                    result["_warning"] = email_result["warning"]
                    step_warnings["send_document_email"] = email_result["warning"]
                    
            elif action == "create_announcement":
                activity_data = {
                    "activity_id": f"act_{uuid.uuid4().hex[:8]}",
                    "author_id": user['user_id'],
                    "agent_id": user['user_id'],
                    "project_id": params.get("project_id") or context.get("project_id"),
                    "type": "announcement",
                    "title": params.get("message_title") or context.get("message_title", "Announcement"),
                    "content": params.get("message_content") or context.get("message_content", ""),
                    "is_draft": False,
                    "is_demo": is_demo,
                    "created_at": datetime.now(timezone.utc).isoformat()
                }
                await db.activities.insert_one(activity_data)
                result["activity_id"] = activity_data["activity_id"]
                
            elif action == "send_project_announcement_email":
                project_id = params.get("project_id") or context.get("project_id")
                message_title = params.get("message_title") or context.get("message_title", "Project Update")
                message_content = params.get("message_content") or context.get("message_content", "")
                
                if project_id:
                    # Get all clients in project
                    clients = await db.clients.find(
                        {"project_id": project_id, "is_demo": is_demo, "email": {"$exists": True, "$ne": None}},
                        {"_id": 0, "email": 1, "name": 1}
                    ).to_list(100)
                    
                    emails_sent = 0
                    email_failures = []
                    
                    if not RESEND_API_KEY:
                        result["emails_sent"] = 0
                        result["total_clients"] = len(clients)
                        result["_warning"] = "Email service not configured"
                        step_warnings["send_project_announcement_email"] = "Email service not configured"
                    else:
                        for client in clients:
                            if client.get('email'):
                                subject = f"Project Update: {message_title}"
                                html_content = f"""
                                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                                    <h2 style="color: #8b5cf6;">{message_title}</h2>
                                    <p>Dear {client.get('name', 'Valued Client')},</p>
                                    <div style="background-color: #f5f3ff; padding: 16px; border-radius: 8px; margin: 16px 0;">
                                        {message_content}
                                    </div>
                                    <p style="margin-top: 24px; color: #6b7280; font-size: 14px;">
                                        Best regards,<br>{agent_name}
                                    </p>
                                </div>
                                """
                                try:
                                    await send_email_async(client['email'], subject, html_content)
                                    emails_sent += 1
                                except Exception as e:
                                    email_failures.append(client['email'])
                                    logger.warning(f"Failed to send announcement to {client['email']}: {e}")
                        
                        result["emails_sent"] = emails_sent
                        result["total_clients"] = len(clients)
                        
                        if email_failures:
                            result["_warning"] = f"Failed to send to {len(email_failures)} recipient(s)"
                            step_warnings["send_project_announcement_email"] = result["_warning"]
                        
                        logger.info(f"Workflow sent announcement to {emails_sent}/{len(clients)} clients")
                else:
                    result["emails_sent"] = 0
                    result["_warning"] = "No project_id provided"
                    step_warnings["send_project_announcement_email"] = "No project_id provided"
                    
            else:
                logger.warning(f"Unknown workflow action: {action}")
            
            return result
        
        # Run the workflow
        execution = await service.run_workflow(execution, action_executor)
        
        # Apply warnings to steps
        for step in execution.steps:
            if step.action in step_warnings:
                step.warning = step_warnings[step.action]
                if step.status == WorkflowStepStatus.COMPLETED:
                    step.status = WorkflowStepStatus.COMPLETED_WITH_WARNING
        
        # Update final status if there are warnings
        has_warnings = any(s.status == WorkflowStepStatus.COMPLETED_WITH_WARNING for s in execution.steps)
        if has_warnings and execution.status == WorkflowStatus.COMPLETED:
            execution.status = WorkflowStatus.COMPLETED_WITH_WARNINGS
        
        await service.update_execution(execution)
        
        # Get summary for response
        summary = service.get_execution_summary(execution)
        
        return {
            "success": execution.status in [WorkflowStatus.COMPLETED, WorkflowStatus.COMPLETED_WITH_WARNINGS],
            "execution": summary
        }
        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Workflow execution failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/workflows/executions/{execution_id}")
async def get_workflow_execution(
    execution_id: str,
    user: dict = Depends(get_current_agent)
):
    """Get the status of a workflow execution"""
    service = get_workflow_service(db)
    execution = await service.get_execution(execution_id)
    
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")
    
    if execution.agent_id != user['user_id']:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    return service.get_execution_summary(execution)

@api_router.get("/workflows/history")
async def get_workflow_history(
    limit: int = 20,
    user: dict = Depends(get_current_agent)
):
    """Get recent workflow executions for the agent"""
    service = get_workflow_service(db)
    is_demo = user.get('is_demo', False)
    
    executions = await service.get_agent_executions(user['user_id'], is_demo, limit)
    
    return {
        "executions": [service.get_execution_summary(e) for e in executions]
    }

@api_router.get("/workflows/selectors")
async def get_workflow_selectors(
    selector_type: str,
    project_id: Optional[str] = None,
    user: dict = Depends(get_current_agent)
):
    """Get items for workflow selectors (documents, timeline steps, etc.)"""
    is_demo = user.get('is_demo', False)
    
    if selector_type == "document":
        # Get recent documents (invoices and quotes)
        query = {"agent_id": user['user_id']}
        if project_id:
            query["project_id"] = project_id
        
        docs = await db.documents.find(
            query,
            {"_id": 0, "document_id": 1, "type": 1, "title": 1, "status": 1, "client_id": 1, "total_amount": 1, "created_at": 1}
        ).sort("created_at", -1).limit(50).to_list(50)
        
        # Enrich with client names
        for doc in docs:
            if doc.get('client_id'):
                client = await db.clients.find_one(
                    {"client_id": doc['client_id']},
                    {"_id": 0, "name": 1}
                )
                doc['client_name'] = client.get('name', 'Unknown') if client else 'Unknown'
        
        return {"items": docs}
    
    elif selector_type == "timeline_step":
        # Get timeline steps (ownership enforced via project/timeline chain)
        query = {}
        if project_id:
            # Get timeline for this project
            timeline = await db.project_timelines.find_one(
                {"project_id": project_id},
                {"_id": 0, "timeline_id": 1}
            )
            if timeline:
                query["project_timeline_id"] = timeline['timeline_id']
            else:
                return {"items": []}
        else:
            # Get steps from agent's projects
            projects = await db.projects.find(
                {"agent_id": user['user_id']},
                {"_id": 0, "project_id": 1}
            ).to_list(100)
            project_ids = [p['project_id'] for p in projects]
            
            # Get timelines for these projects
            timelines = await db.project_timelines.find(
                {"project_id": {"$in": project_ids}},
                {"_id": 0, "timeline_id": 1, "project_id": 1}
            ).to_list(100)
            timeline_ids = [t['timeline_id'] for t in timelines]
            timeline_project_map = {t['timeline_id']: t['project_id'] for t in timelines}
            
            if not timeline_ids:
                return {"items": []}
            query["project_timeline_id"] = {"$in": timeline_ids}
        
        # Use correct field names: title (not name), order_index (not order)
        steps = await db.timeline_steps.find(
            query,
            {"_id": 0, "step_id": 1, "title": 1, "status": 1, "project_timeline_id": 1, "planned_date": 1}
        ).sort("order_index", 1).limit(100).to_list(100)
        
        # Enrich with project names
        for step in steps:
            # Map title to name for UI consistency
            step['name'] = step.pop('title', 'Untitled')
            
            timeline_id = step.get('project_timeline_id')
            if timeline_id:
                # Get project_id from timeline
                timeline = await db.project_timelines.find_one(
                    {"timeline_id": timeline_id},
                    {"_id": 0, "project_id": 1}
                )
                if timeline:
                    project = await db.projects.find_one(
                        {"project_id": timeline['project_id']},
                        {"_id": 0, "name": 1}
                    )
                    step['project_name'] = project.get('name', 'Unknown') if project else 'Unknown'
        
        return {"items": steps}
    
    elif selector_type == "client":
        # Get clients
        query = {"agent_id": user['user_id']}
        if project_id:
            query["project_id"] = project_id
        
        clients = await db.clients.find(
            query,
            {"_id": 0, "client_id": 1, "name": 1, "email": 1, "project_id": 1}
        ).sort("name", 1).limit(100).to_list(100)
        
        return {"items": clients}
    
    else:
        raise HTTPException(status_code=400, detail=f"Unknown selector type: {selector_type}")

@api_router.post("/workflows/executions/{execution_id}/confirm")
async def confirm_workflow_step(
    execution_id: str,
    request: WorkflowConfirmRequest,
    user: dict = Depends(get_current_agent)
):
    """Confirm and continue a paused workflow step"""
    service = get_workflow_service(db)
    execution = service.get_execution(execution_id)
    
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")
    
    if execution.agent_id != user['user_id']:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    # This would need the action_executor - simplified for now
    raise HTTPException(status_code=501, detail="Step confirmation not yet implemented")

@api_router.post("/workflows/executions/{execution_id}/cancel")
async def cancel_workflow(
    execution_id: str,
    user: dict = Depends(get_current_agent)
):
    """Cancel a workflow execution"""
    service = get_workflow_service(db)
    
    try:
        execution = await service.cancel_execution(execution_id, user['user_id'])
        return service.get_execution_summary(execution)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

@api_router.post("/workflows/executions/{execution_id}/steps/{step_index}/retry")
async def retry_workflow_step(
    execution_id: str,
    step_index: int,
    user: dict = Depends(get_current_agent)
):
    """Retry a failed or warning step in a workflow execution"""
    service = get_workflow_service(db)
    execution = await service.get_execution(execution_id)
    
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")
    
    if execution.agent_id != user['user_id']:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    is_demo = user.get('is_demo', False)
    
    # Get agent settings for email
    agent_settings = await db.agent_settings.find_one(
        {"agent_id": user['user_id']},
        {"_id": 0}
    ) or {}
    
    agent_name = agent_settings.get('profile', {}).get('display_name') or user.get('name', 'Your Agent')
    agent_email = agent_settings.get('profile', {}).get('contact_email') or user.get('email', '')
    
    # Helper to safely send email with error handling
    async def safe_send_email(to_email: str, subject: str, html_content: str, step_name: str) -> dict:
        """Send email with graceful error handling"""
        if not to_email:
            return {"sent": False, "warning": "No recipient email address"}
        
        if not RESEND_API_KEY:
            logger.warning(f"RESEND_API_KEY not configured, skipping email for {step_name}")
            return {"sent": False, "warning": "Email service not configured"}
        
        try:
            await send_email_async(to_email, subject, html_content)
            return {"sent": True}
        except Exception as e:
            logger.warning(f"Email failed for {step_name}: {e}")
            return {"sent": False, "warning": f"Email delivery failed: {str(e)[:100]}"}
    
    # Re-create the action executor with the same context
    async def action_executor(action: str, params: dict, context: dict) -> dict:
        """Execute a workflow action and return results"""
        result = {}
        
        if action == "complete_timeline_step":
            step_id = params.get("step_id") or context.get("step_id")
            if step_id:
                await db.timeline_steps.update_one(
                    {"step_id": step_id},
                    {"$set": {"status": "completed", "completed_at": datetime.now(timezone.utc).isoformat()}}
                )
                result["step_status"] = "completed"
                logger.info(f"Retry: completed step {step_id}")
                
        elif action == "send_milestone_email":
            client_email = params.get("client_email") or context.get("client_email")
            client_name = params.get("client_name") or context.get("client_name", "")
            step_name = params.get("step_name") or context.get("step_name", "a milestone")
            project_name = params.get("project_name") or context.get("project_name", "your project")
            
            subject = f"Milestone Completed: {step_name}"
            html_content = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <h2 style="color: #f59e0b;">Milestone Completed!</h2>
                <p>Dear {client_name},</p>
                <p>Great news! The following milestone has been completed for <strong>{project_name}</strong>:</p>
                <div style="background-color: #fef3c7; padding: 16px; border-radius: 8px; margin: 16px 0;">
                    <strong>{step_name}</strong>
                </div>
                <p>Best regards,<br>{agent_name}</p>
            </div>
            """
            email_result = await safe_send_email(client_email, subject, html_content, "send_milestone_email")
            result["email_sent"] = email_result["sent"]
            if email_result.get("warning"):
                result["_warning"] = email_result["warning"]
                
        elif action == "update_document_status":
            doc_id = params.get("document_id") or context.get("document_id")
            new_status = params.get("status", "Sent")
            if doc_id:
                update_data = {"status": new_status, "updated_at": datetime.now(timezone.utc).isoformat()}
                if new_status == "Paid":
                    update_data["paid_at"] = datetime.now(timezone.utc).isoformat()
                await db.documents.update_one(
                    {"document_id": doc_id},
                    {"$set": update_data}
                )
                result["status"] = new_status
                logger.info(f"Retry: updated document {doc_id} to {new_status}")
                
        elif action == "send_payment_confirmation_email":
            client_email = params.get("client_email") or context.get("client_email")
            client_name = params.get("client_name") or context.get("client_name", "")
            amount = params.get("amount") or context.get("amount", 0)
            doc_id = params.get("document_id") or context.get("document_id")
            
            subject = "Payment Received - Thank You!"
            html_content = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <h2 style="color: #10b981;">Payment Received!</h2>
                <p>Dear {client_name},</p>
                <p>We have received your payment of <strong>CHF {amount:,.2f}</strong>.</p>
                <p>Your invoice has been marked as paid.</p>
            </div>
            """
            email_result = await safe_send_email(client_email, subject, html_content, "send_payment_confirmation_email")
            result["email_sent"] = email_result["sent"]
            if email_result.get("warning"):
                result["_warning"] = email_result["warning"]
        else:
            logger.warning(f"Retry: Unknown action {action}")
        
        return result
    
    try:
        execution = await service.retry_step(execution, step_index, action_executor)
        return {
            "success": True,
            "execution": service.get_execution_summary(execution),
            "retried_step": step_index
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Retry workflow step failed: {e}")
        # Still return the updated execution state
        updated_execution = await service.get_execution(execution_id)
        return {
            "success": False,
            "error": str(e),
            "execution": service.get_execution_summary(updated_execution) if updated_execution else None
        }


async def _extract_document_from_image(file_path: str, doc_type: str) -> dict:
    """Extract invoice/quote data from image using Vision API"""
    if not OPENAI_API_KEY:
        return {"extraction_failed": True, "confidence": "low"}
    
    try:
        import base64
        with open(file_path, "rb") as img_file:
            image_data = base64.b64encode(img_file.read()).decode('utf-8')
        
        file_ext = os.path.splitext(file_path)[1].lower()
        mime_types = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png", ".webp": "image/webp"}
        mime_type = mime_types.get(file_ext, "image/jpeg")
        
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        prompt = f"""Analyze this {doc_type} image and extract the following information:
- supplier_name: Company or person issuing the document
- amount: Total amount (number only)
- title: Brief description of what this is for
- description: Detailed description if available
- reference: Document/invoice number
- items: List of line items if visible [{{"description": "...", "quantity": 1, "unit_price": 0, "total": 0}}]

Return ONLY a JSON object with these fields. Use null for missing fields."""
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_data}"}}
                ]
            }],
            max_tokens=1500
        )
        
        result_text = response.choices[0].message.content
        
        # Parse JSON from response
        import re
        json_match = re.search(r'\{[\s\S]*\}', result_text)
        if json_match:
            result = json.loads(json_match.group())
            result["confidence"] = "medium"
            return result
        
        return {"extraction_failed": True, "confidence": "low"}
        
    except Exception as e:
        logger.error(f"Image extraction failed: {e}")
        return {"extraction_failed": True, "confidence": "low", "error": str(e)}


async def _extract_timeline_stages(file_path: str) -> dict:
    """Extract timeline/schedule data from document using AI"""
    if not OPENAI_API_KEY:
        return {"stages": [], "confidence": 0.1, "extraction_failed": True}
    
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        # Extract text from PDF or image
        text_content = ""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            doc = fitz.open(file_path)
            for page in doc:
                text_content += page.get_text()
            doc.close()
        elif file_ext in ['.jpg', '.jpeg', '.png', '.webp']:
            # Use Vision API for images
            import base64
            with open(file_path, "rb") as img_file:
                image_data = base64.b64encode(img_file.read()).decode('utf-8')
            
            mime_types = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png", ".webp": "image/webp"}
            mime_type = mime_types.get(file_ext, "image/jpeg")
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract all text from this document, especially any timeline, schedule, or phase information."},
                        {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_data}"}}
                    ]
                }],
                max_tokens=2000
            )
            text_content = response.choices[0].message.content
        
        if not text_content:
            return {"stages": [], "confidence": 0.1, "extraction_failed": True}
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": """You extract project timeline/schedule information from documents.
                    
IMPORTANT: Keep all dates as RAW TEXT. Do not convert or normalize them.
Examples of valid date formats: "Q4 2027", "Spring 2026", "Week 12-15", "March 2025", "6 months"

Return ONLY a JSON object:
{
  "stages": [
    {"name": "Phase name", "date_text": "Raw date as written", "description": "Brief description"}
  ],
  "project_duration": "Total duration as text",
  "confidence": 0.0-1.0
}"""
                },
                {
                    "role": "user",
                    "content": f"Extract timeline stages from this document:\n\n{text_content[:8000]}"
                }
            ],
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            return json.loads(json_match.group())
        
        return {"stages": [], "confidence": 0.3}
        
    except Exception as e:
        logger.error(f"Timeline extraction failed: {e}")
        return {"stages": [], "confidence": 0.1, "extraction_failed": True}


async def _extract_contacts_list(file_path: str) -> dict:
    """Extract contact information from document using AI"""
    if not OPENAI_API_KEY:
        return {"contacts": [], "confidence": 0.1, "extraction_failed": True}
    
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        # Extract text from PDF
        text_content = ""
        if file_path.lower().endswith('.pdf'):
            doc = fitz.open(file_path)
            for page in doc:
                text_content += page.get_text()
            doc.close()
        
        if not text_content:
            return {"contacts": [], "confidence": 0.1, "extraction_failed": True}
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": """You extract contact information from documents.
                    
Extract names, roles, companies, emails, and phone numbers.
Remove duplicates. Only include contacts with at least a name.

Return ONLY a JSON object:
{
  "contacts": [
    {"name": "Full Name", "role": "Job title or role", "company": "Company name", "email": "email@example.com", "phone": "+41 XX XXX XX XX"}
  ],
  "confidence": 0.0-1.0
}"""
                },
                {
                    "role": "user",
                    "content": f"Extract contacts from this document:\n\n{text_content[:8000]}"
                }
            ],
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            result = json.loads(json_match.group())
            # Deduplicate by name
            seen_names = set()
            unique_contacts = []
            for contact in result.get("contacts", []):
                name = contact.get("name", "").strip().lower()
                if name and name not in seen_names:
                    seen_names.add(name)
                    unique_contacts.append(contact)
            result["contacts"] = unique_contacts
            return result
        
        return {"contacts": [], "confidence": 0.3}
        
    except Exception as e:
        logger.error(f"Contact extraction failed: {e}")
        return {"contacts": [], "confidence": 0.1, "extraction_failed": True}


# ==================== ROOT ENDPOINTS ====================

@api_router.get("/")
async def root():
    return {"message": "Evohome API", "version": "3.0.0"}

@api_router.get("/health")
async def health():
    return {"status": "healthy"}

# Also add root-level health for cases where prefix is stripped
@app.get("/health")
async def health_root():
    return {"status": "healthy"}

@app.get("/api/health")
async def health_api():
    return {"status": "healthy"}

# Root-level demo seed for DigitalOcean (prefix might be stripped)
@app.get("/demo/seed")
async def seed_demo_root():
    return await seed_demo_data()

# Include router
app.include_router(api_router)

# Mount static files for uploads (logos, images) under /api/uploads
# This ensures proper routing through the ingress
app.mount("/api/uploads", StaticFiles(directory=str(UPLOAD_DIR)), name="uploads")

# CORS - using validated config (no wildcard in production)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=app_config.CORS_ORIGINS,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

@app.on_event("startup")
async def create_indexes():
    """Create database indexes on startup"""
    try:
        # TTL index for workflow executions - auto-delete after 30 days
        # Uses created_at_date which is a BSON Date, not created_at which is a string
        await db.workflow_executions.create_index(
            "created_at_date",
            expireAfterSeconds=30 * 24 * 60 * 60  # 30 days
        )
        logger.info("Database indexes created/verified")
    except Exception as e:
        logger.warning(f"Index creation warning (may already exist): {e}")
